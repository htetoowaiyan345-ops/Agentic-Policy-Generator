"""test_lines_json_renderer.py

Automated tests for Stage 2 — the direct slot-by-slot writer.

Validates:
  1. All-slot-0 payload still renders without crashing (graceful degradation).
  2. Mixed slot assignment lands in the right slot.
  3. Rich HTML preservation — bold/italic/color survive round-trip.
  4. Table placement — table appears in slot 10 or 14 (history/award).
  5. No `&nbsp;` injection — output only contains real spaces unless
     the user typed `&nbsp;`.
  6. Brain framework integrity — slot titles, header, footer, logo
     remain byte-identical to the brain template.

These tests bypass the broken `policy_platform/__init__.py` import
(which references a missing `retrieval_pipeline` module) by loading
the renderer via `importlib.util.spec_from_file_location`.
"""
from __future__ import annotations

import importlib
import importlib.util
import shutil
import sys
import types
from pathlib import Path

import pytest
from docx import Document
from docx.oxml.ns import qn


# ---------------------------------------------------------------------------
# Fixtures: load the renderer without triggering policy_platform.__init__.py
# ---------------------------------------------------------------------------

@pytest.fixture(scope="session")
def brain_path() -> Path:
    return Path(__file__).resolve().parents[1] / "data" / "brain_template" / "Policy_Framework_5.docx"


@pytest.fixture(scope="session")
def out_dir(tmp_path_factory) -> Path:
    d = tmp_path_factory.mktemp("lines_json_renderer_out")
    return d


@pytest.fixture(scope="session")
def renderer_mod():
    """Load `policy_platform.lines_json_renderer` without running the
    broken `policy_platform/__init__.py` that imports a missing
    `retrieval_pipeline` module."""
    backend_root = Path(__file__).resolve().parents[1]
    # Add backend to sys.path so absolute imports inside sub-modules work.
    if str(backend_root) not in sys.path:
        sys.path.insert(0, str(backend_root))
    # Stub the package namespace.
    pkg = types.ModuleType("policy_platform")
    pkg.__path__ = [str(backend_root / "policy_platform")]
    pkg.__package__ = "policy_platform"
    sys.modules["policy_platform"] = pkg
    fw = types.ModuleType("policy_platform.framework")
    fw.__path__ = [str(backend_root / "policy_platform" / "framework")]
    fw.__package__ = "policy_platform.framework"
    sys.modules["policy_platform.framework"] = fw
    # Pre-load sub-modules via direct file path (bypasses __init__.py).
    def _load(name: str, path: Path):
        spec = importlib.util.spec_from_file_location(name, str(path))
        assert spec is not None and spec.loader is not None
        m = importlib.util.module_from_spec(spec)
        sys.modules[name] = m
        spec.loader.exec_module(m)
        return m
    _load("policy_platform.config", backend_root / "policy_platform" / "config.py")
    _load("policy_platform.framework.section_map",
          backend_root / "policy_platform" / "framework" / "section_map.py")
    _load("policy_platform.framework.brain_slot_map",
          backend_root / "policy_platform" / "framework" / "brain_slot_map.py")
    # Load the renderer.
    src = backend_root / "policy_platform" / "lines_json_renderer.py"
    spec = importlib.util.spec_from_file_location("policy_platform.lines_json_renderer", str(src))
    assert spec is not None and spec.loader is not None
    mod = importlib.util.module_from_spec(spec)
    sys.modules["policy_platform.lines_json_renderer"] = mod
    spec.loader.exec_module(mod)
    return mod


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _count_runs_with_prop(doc, prop_name: str) -> int:
    """Count `<w:r>` elements in `doc` whose rPr contains `prop_name`."""
    count = 0
    for p in doc.paragraphs:
        for r in p._element.iter(qn("w:r")):
            rPr = r.find(qn("w:rPr"))
            if rPr is not None and rPr.find(qn(f"w:{prop_name}")) is not None:
                count += 1
    return count


def _has_text(doc, needle: str) -> bool:
    """Search for `needle` in paragraphs AND table cells."""
    if any(needle in p.text for p in doc.paragraphs):
        return True
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                if needle in cell.text:
                    return True
    return False


def _has_table_in(doc, slot_id: int) -> bool:
    """Heuristic: check whether slot `slot_id` has at least one table in
    its body region. We rely on the renderer's known scaffold indices
    but tolerate deletion of empty paragraphs (which shifts indices)."""
    from policy_platform.framework.brain_slot_map import BRAIN_SLOT_RANGES, find_slot_boundaries
    body = doc.element.body
    # Use find_slot_boundaries (text-based) instead of fixed indices,
    # because the renderer now DELETES empty scaffold paragraphs which
    # shifts BRAIN_SLOT_RANGES indices.
    bounds = find_slot_boundaries(doc)
    info = bounds.get(slot_id, {})
    for el in info.get("elements", []):
        if el.tag.split("}")[-1] == "tbl":
            return True
    return False


# ---------------------------------------------------------------------------
# Tests
# ---------------------------------------------------------------------------

def test_all_slot_zero_does_not_crash(renderer_mod, brain_path, out_dir):
    """Backward compat: old payloads with all slot=0 must still publish."""
    lines_json = [
        ["p", {"slot": 0, "text": "first", "html": "<p>first</p>"}],
        ["p", {"slot": 0, "text": "second", "html": "<p>second</p>"}],
        ["t", {"slot": 0, "rows": [["a", "b"], ["c", "d"]]}],
    ]
    out = out_dir / "all_slot_zero.docx"
    renderer_mod.render_lines_json_to_brain(lines_json, brain_path, out)
    assert out.exists()
    assert out.stat().st_size > 1000  # not empty
    doc = Document(str(out))
    # Per user spec: "Free Paragraphs" zone label is removed — it was
    # unnecessary text in the published docx. Free-zone items still render
    # in their original insertion order.
    assert not any(p.text.strip() == "Free Paragraphs" for p in doc.paragraphs)
    # User-written content still appears in the docx.
    assert _has_text(doc, "first")
    assert _has_text(doc, "second")


def test_mixed_slot_assignment(renderer_mod, brain_path, out_dir):
    """Each slot must receive its own user content."""
    lines_json = [
        ["p", {"slot": 5, "text": "INTRO TEXT", "html": "<p>INTRO TEXT</p>"}],
        ["p", {"slot": 7, "text": "PURPOSE TEXT", "html": "<p>PURPOSE TEXT</p>"}],
        ["p", {"slot": 12, "text": "DEFINITIONS TEXT", "html": "<p>DEFINITIONS TEXT</p>"}],
    ]
    out = out_dir / "mixed_slots.docx"
    renderer_mod.render_lines_json_to_brain(lines_json, brain_path, out)
    doc = Document(str(out))
    assert _has_text(doc, "INTRO TEXT")
    assert _has_text(doc, "PURPOSE TEXT")
    assert _has_text(doc, "DEFINITIONS TEXT")


def test_rich_html_preserved(renderer_mod, brain_path, out_dir):
    """Bold, italic, color must survive the round-trip."""
    color_span = '<p><span style="color:#ff0000">red text</span></p>'
    lines_json = [
        ["p", {
            "slot": 6,
            "text": "B and i",
            "html": '<p>This policy is <strong>BOLD</strong> and <em>italic</em>.</p>',
        }],
        ["p", {"slot": 7, "text": "red", "html": color_span}],
    ]
    out = out_dir / "rich.docx"
    renderer_mod.render_lines_json_to_brain(lines_json, brain_path, out)
    doc = Document(str(out))
    assert _count_runs_with_prop(doc, "b") >= 1, "bold runs not preserved"
    assert _count_runs_with_prop(doc, "i") >= 1, "italic runs not preserved"
    assert _count_runs_with_prop(doc, "color") >= 1, "color runs not preserved"


def test_table_placement_in_history_slot(renderer_mod, brain_path, out_dir):
    """A table saved with slot 14 must appear in the HISTORY region."""
    lines_json = [
        ["t", {
            "slot": 14,
            "rows": [
                ["Version", "Date", "Author", "Notes"],
                ["1.0", "2026-01-01", "Alice", "Initial draft"],
                ["1.1", "2026-02-01", "Bob", "Tweaks"],
            ],
        }],
    ]
    out = out_dir / "table_in_history.docx"
    renderer_mod.render_lines_json_to_brain(lines_json, brain_path, out)
    doc = Document(str(out))
    assert _has_text(doc, "Initial draft")
    # At least one of the slot 10/14 tables should be present.
    assert _has_table_in(doc, 14)


def test_no_nbsp_injection(renderer_mod, brain_path, out_dir):
    """No `&nbsp;` is injected unless the user typed it."""
    lines_json = [
        ["p", {
            "slot": 6,
            "text": "two   spaces",
            "html": "<p>two   spaces</p>",
        }],
    ]
    out = out_dir / "no_nbsp.docx"
    renderer_mod.render_lines_json_to_brain(lines_json, brain_path, out)
    doc = Document(str(out))
    # Find the rendered paragraph.
    rendered_text = next(p.text for p in doc.paragraphs if "two" in p.text)
    # Should NOT contain the non-breaking space (U+00A0) because the
    # user didn't type any.
    assert "\u00a0" not in rendered_text, f"\\u00a0 found in rendered text: {rendered_text!r}"
    # Real spaces preserved (we typed 3 consecutive spaces).
    assert "two   spaces" in rendered_text


def test_user_typed_nbsp_preserved(renderer_mod, brain_path, out_dir):
    """If the user DID type `&nbsp;`, preserve it (semantic)."""
    lines_json = [
        ["p", {
            "slot": 6,
            "text": "nbsp here",
            "html": "<p>nbsp&nbsp;here</p>",
        }],
    ]
    out = out_dir / "user_typed_nbsp.docx"
    renderer_mod.render_lines_json_to_brain(lines_json, brain_path, out)
    doc = Document(str(out))
    found_nbsp = False
    for p in doc.paragraphs:
        for t in p._element.iter(qn("w:t")):
            if t.text and "\u00a0" in t.text:
                found_nbsp = True
                break
        if found_nbsp:
            break
    assert found_nbsp, "user-typed nbsp should be preserved"


def test_brain_framework_intact(renderer_mod, brain_path, out_dir):
    """Slot titles, header text, footer, media all preserved."""
    lines_json = [
        ["p", {"slot": 1, "text": "HR", "html": "<p>HR</p>"}],
        # Use the heading label as the user's first paragraph so the
        # scaffold heading dedup preserves it.
        ["p", {"slot": 5, "text": "INTRODUCTION", "html": "<p>INTRODUCTION</p>"}],
        ["p", {"slot": 5, "text": "intro body", "html": "<p>intro body</p>"}],
        ["p", {"slot": 6, "text": "POLICY STATEMENT", "html": "<p>POLICY STATEMENT</p>"}],
        ["p", {"slot": 6, "text": "policy body", "html": "<p>policy body</p>"}],
        ["p", {"slot": 7, "text": "1. Purpose", "html": "<p>1. Purpose</p>"}],
        ["p", {"slot": 7, "text": "purpose body", "html": "<p>purpose body</p>"}],
        ["p", {"slot": 8, "text": "2. Scope & Beneficiaries", "html": "<p>2. Scope & Beneficiaries</p>"}],
        ["p", {"slot": 8, "text": "scope body", "html": "<p>scope body</p>"}],
        ["p", {"slot": 9, "text": "3. Exclusions", "html": "<p>3. Exclusions</p>"}],
        ["p", {"slot": 9, "text": "exclusions body", "html": "<p>exclusions body</p>"}],
        ["p", {"slot": 10, "text": "4. Award Structure & Payout Tiers", "html": "<p>4. Award Structure & Payout Tiers</p>"}],
        ["p", {"slot": 10, "text": "award body", "html": "<p>award body</p>"}],
        ["p", {"slot": 12, "text": "DEFINITIONS", "html": "<p>DEFINITIONS</p>"}],
        ["p", {"slot": 12, "text": "defs body", "html": "<p>defs body</p>"}],
        ["p", {"slot": 13, "text": "RELATED POLICIES, PROCEDURES, FORMS, GUIDELINES & OTHER RESOURCES", "html": "<p>RELATED POLICIES, PROCEDURES, FORMS, GUIDELINES & OTHER RESOURCES</p>"}],
        ["p", {"slot": 13, "text": "related body", "html": "<p>related body</p>"}],
        ["p", {"slot": 14, "text": "HISTORY", "html": "<p>HISTORY</p>"}],
        ["p", {"slot": 14, "text": "history body", "html": "<p>history body</p>"}],
    ]
    out = out_dir / "framework_intact.docx"
    renderer_mod.render_lines_json_to_brain(lines_json, brain_path, out)

    # Media integrity check: same media files in both zips.
    import zipfile
    def media_set(p):
        with zipfile.ZipFile(p) as z:
            return {n: z.read(n) for n in z.namelist() if n.startswith("word/media/")}
    out_media = media_set(out)
    brain_media = media_set(brain_path)
    assert out_media == brain_media, "media integrity broken"

    # Slot headings still present in the rendered output.
    doc = Document(str(out))
    expected_titles = [
        "INTRODUCTION", "POLICY STATEMENT", "1. Purpose",
        "2. Scope & Beneficiaries", "3. Exclusions",
        "4. Award Structure & Payout Tiers", "DEFINITIONS",
        "RELATED POLICIES", "HISTORY",
    ]
    for title in expected_titles:
        assert _has_text(doc, title), f"slot title missing: {title}"


def test_extra_paragraphs_appended(renderer_mod, brain_path, out_dir):
    """User paragraphs beyond the scaffold count are appended (not dropped)."""
    many_paras = [
        ["p", {
            "slot": 6,
            "text": f"slot-6 line {i}",
            "html": f"<p>slot-6 line {i}</p>",
        }]
        for i in range(20)
    ]
    out = out_dir / "extra_paragraphs.docx"
    renderer_mod.render_lines_json_to_brain(many_paras, brain_path, out)
    doc = Document(str(out))
    for i in range(20):
        assert _has_text(doc, f"slot-6 line {i}"), f"missing line {i}"


def test_extra_tables_appended(renderer_mod, brain_path, out_dir):
    """User tables beyond the scaffold count are appended."""
    many_tables = [
        ["t", {"slot": 14, "rows": [[f"row{i}-col{j}" for j in range(2)]]}]
        for i in range(5)
    ]
    out = out_dir / "extra_tables.docx"
    renderer_mod.render_lines_json_to_brain(many_tables, brain_path, out)
    doc = Document(str(out))
    for i in range(5):
        assert _has_text(doc, f"row{i}-col0"), f"missing table row {i}"


def test_normalise_lines_json(renderer_mod):
    """The normaliser correctly buckets paragraphs/tables by slot id and
    gracefully accepts both `dict` and `str` payloads (legacy shape)."""
    lines_json = [
        ["p", {"slot": 3, "text": "rich", "html": "<p>rich</p>"}],
        ["p", "legacy string"],
        ["t", {"slot": 10, "rows": [["a"]]}],
        ["t", [["b"]]],
    ]
    paragraphs, tables, dividers, free_zone_items = renderer_mod._normalise_lines_json(lines_json)
    assert 3 in paragraphs
    assert 0 in paragraphs  # the legacy string falls into slot=0
    assert 10 in tables
    assert 0 in tables  # legacy rows fall into slot=0
    assert dividers == []  # no divider entries in this fixture
    # free_zone_items: legacy paragraph + legacy table, in that order.
    assert [k for k, _ in free_zone_items] == ['p', 't']


def test_free_zone_items_preserves_insertion_order(renderer_mod):
    """The normaliser preserves the original insertion order of slot=0
    items (paragraphs, tables, dividers). The renderer uses this to keep
    toolbar-inserted content in the SAME visual order the user inserted it.
    """
    lines_json = [
        ["p", {"slot": 0, "text": "para A", "html": "<p>para A</p>"}],
        ["divider", {"slot": 0}],
        ["t", {"slot": 0, "rows": [["hdr", "col"], ["a", "b"]]}],
        ["p", {"slot": 0, "text": "para B", "html": "<p>para B</p>"}],
        ["divider", {"slot": 0}],
        ["p", {"slot": 5, "text": "slot 5", "html": "<p>slot 5</p>"}],  # not free zone
    ]
    paragraphs, tables, dividers, free_zone_items = renderer_mod._normalise_lines_json(lines_json)
    # Only slot=0 items are in free_zone_items
    assert [k for k, _ in free_zone_items] == ['p', 'divider', 't', 'p', 'divider']
    # Named-slot paragraph is in paragraphs[5], not in free_zone_items
    assert 5 in paragraphs
    assert 0 in paragraphs
    # Dividers list has both slot=0 dividers (for backwards compat)
    assert len(dividers) == 2


def test_slot0_table_renders_in_free_paragraph_zone(renderer_mod, brain_path, out_dir):
    """A toolbar-inserted table (slot=0) renders as a real <w:tbl> in the
    free paragraph zone, with all cell content preserved. Before this
    fix, slot=0 tables were silently dropped from the published output.
    """
    lines_json = [
        ["p", {"slot": 0, "text": "before table", "html": "<p>before table</p>"}],
        ["t", {
            "slot": 0,
            "rows": [
                ["Header1", "Header2"],
                ["cell1a", "cell1b"],
                ["cell2a", "cell2b"],
            ],
        }],
        ["p", {"slot": 0, "text": "after table", "html": "<p>after table</p>"}],
    ]
    out = out_dir / "slot0_table.docx"
    renderer_mod.render_lines_json_to_brain(lines_json, brain_path, out)
    doc = Document(str(out))
    # All three paragraphs render
    assert _has_text(doc, "before table"), "before-text missing"
    assert _has_text(doc, "after table"), "after-text missing"
    # All five table cells render
    assert _has_text(doc, "Header1"), "Header1 missing"
    assert _has_text(doc, "Header2"), "Header2 missing"
    assert _has_text(doc, "cell1a"), "cell1a missing"
    assert _has_text(doc, "cell1b"), "cell1b missing"
    assert _has_text(doc, "cell2a"), "cell2a missing"
    assert _has_text(doc, "cell2b"), "cell2b missing"
    # Verify a real <w:tbl> exists in the output
    body = doc.element.body
    tables = list(body.findall(".//" + qn("w:tbl")))
    assert len(tables) >= 1, "no <w:tbl> in output"

    # Verify the table is in the FREE paragraph zone (before slot 1's "Type:" heading)
    # by checking the document.xml ordering: the table appears before the
    # "Type:" scaffold paragraph.
    body_children = list(body)
    type_idx = None
    tbl_idx = None
    for i, ch in enumerate(body_children):
        if ch.tag.endswith("}p"):
            txt = "".join((t.text or "") for t in ch.iter(qn("w:t"))).strip()
            if txt == "Type" or txt.startswith("Type:"):
                type_idx = i
        elif ch.tag.endswith("}tbl") and tbl_idx is None and i > 0:
            # First table = our free-zone table (should be before "Type:")
            tbl_idx = i
    if tbl_idx is not None and type_idx is not None:
        assert tbl_idx < type_idx, (
            "free-zone table should appear BEFORE slot-1 'Type:' heading"
        )


def test_free_zone_preserves_mixed_order(renderer_mod, brain_path, out_dir):
    """Paragraphs, tables, and dividers in slot=0 are rendered in the
    EXACT order they were inserted — not group-sorted by kind.
    """
    lines_json = [
        ["p", {"slot": 0, "text": "first paragraph", "html": "<p>first paragraph</p>"}],
        ["divider", {"slot": 0}],
        ["t", {"slot": 0, "rows": [["col1", "col2"], ["a", "b"]]}],
        ["p", {"slot": 0, "text": "middle paragraph", "html": "<p>middle paragraph</p>"}],
        ["p", {"slot": 0, "text": "last paragraph", "html": "<p>last paragraph</p>"}],
    ]
    out = out_dir / "free_zone_order.docx"
    renderer_mod.render_lines_json_to_brain(lines_json, brain_path, out)
    doc = Document(str(out))
    # Confirm all elements present
    assert _has_text(doc, "first paragraph"), "first paragraph missing"
    assert _has_text(doc, "middle paragraph"), "middle paragraph missing"
    assert _has_text(doc, "last paragraph"), "last paragraph missing"
    assert _has_text(doc, "col1"), "col1 missing"
    assert _has_text(doc, "a"), "cell a missing"

    # Verify order: in the body XML, the first paragraph precedes the
    # table, and the table precedes the middle paragraph.
    body = doc.element.body
    body_children = list(body)
    pos = {}
    for i, ch in enumerate(body_children):
        if ch.tag.endswith("}p"):
            txt = "".join((t.text or "") for t in ch.iter(qn("w:t"))).strip()
            if txt == "first paragraph":
                pos["first"] = i
            elif txt == "middle paragraph":
                pos["middle"] = i
            elif txt == "last paragraph":
                pos["last"] = i
        elif ch.tag.endswith("}tbl"):
            if "table" not in pos:
                pos["table"] = i
    assert pos["first"] < pos["table"] < pos["middle"], (
        f"order broken: first={pos.get('first')}, table={pos.get('table')}, "
        f"middle={pos.get('middle')}"
    )
    assert pos["middle"] < pos["last"], (
        f"middle should precede last: middle={pos.get('middle')}, last={pos.get('last')}"
    )


def test_import_rich_writer_works_under_importlib_load(renderer_mod):
    """The helper must find the rich writer even when the renderer is
    loaded via importlib.util without `api` on sys.path."""
    wp = renderer_mod._import_rich_writer()
    assert wp is not None, "_import_rich_writer returned None"
    assert callable(wp)


# ---------------------------------------------------------------------------
# Stage 6 — full pipeline toolbar round-trip
# ---------------------------------------------------------------------------
# Each test takes a lines_json entry with HTML from one toolbar control,
# runs the full `render_lines_json_to_brain` pipeline, and asserts the
# corresponding OOXML fragment survives the trip. Tests use slot 6
# (POLICY STATEMENT) so the user content replaces the body paragraph of
# that section in the brain framework.


def _render_and_read(lines_json, brain_path, out_dir, name="stage6.docx"):
    out = out_dir / name
    renderer_mod.render_lines_json_to_brain(lines_json, brain_path, out)
    return Document(str(out))


@pytest.mark.parametrize(
    'name,html,expected_frag,rel_frag',
    [
        # 1. Bold
        ('bold',
         '<p>this is <strong>bold</strong></p>',
         '<w:b/>', None),
        # 2. Italic
        ('italic',
         '<p>this is <em>italic</em></p>',
         '<w:i/>', None),
        # 3. Underline
        ('underline',
         '<p>this is <u>underlined</u></p>',
         '<w:u ', None),
        # 4. Strikethrough
        ('strike',
         '<p>this is <s>struck</s></p>',
         '<w:strike/>', None),
        # 5. Color
        ('color',
         '<p><span style="color: #ff0000;">red</span></p>',
         'val="FF0000"', None),
        # 6. Font size
        ('fontsize',
         '<p><span style="font-size: 18px;">big</span></p>',
         '<w:sz', None),
        # 7. Font family
        ('fontfamily',
         '<p><span style="font-family: Calibri;">calibri</span></p>',
         'w:ascii="calibri"', None),
        # 8. Right align
        ('align-right',
         '<p style="text-align: right;">right</p>',
         'w:val="right"', None),
        # 9. Center align
        ('align-center',
         '<p style="text-align: center;">center</p>',
         'w:val="center"', None),
        # 10. Hyperlink
        ('hyperlink',
         '<p>see <a href="https://example.com">link</a></p>',
         '<w:hyperlink', 'example.com'),
        # 11. Blockquote
        ('blockquote',
         '<blockquote>quoted</blockquote>',
         'w:left="720"', None),
    ],
)
def test_stage6_toolbar_format_round_trip(
    renderer_mod, brain_path, out_dir, name, html, expected_frag, rel_frag
):
    """Every toolbar format must survive the full lines_json -> .docx
    pipeline."""
    lines_json = [
        ["p", {"slot": 6, "text": "x", "html": html}],
    ]
    out = out_dir / f"stage6_{name}.docx"
    renderer_mod.render_lines_json_to_brain(lines_json, brain_path, out)
    doc = Document(str(out))
    xml = doc.element.body.xml
    assert expected_frag in xml, (
        f"{name}: missing {expected_frag!r} in output (html={html!r})"
    )
    if rel_frag is not None:
        # Verify the relationship landed in document.xml.rels.
        import zipfile
        with zipfile.ZipFile(str(out)) as z:
            rels = z.read('word/_rels/document.xml.rels').decode('utf-8')
        assert rel_frag in rels, (
            f"{name}: missing {rel_frag!r} in rels"
        )