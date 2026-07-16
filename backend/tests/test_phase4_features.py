"""Tests for the new behaviors: noise strip (Version|Page X of Y), header
extractor input-driven title, and empty-slot bullet cleanup."""
from __future__ import annotations

from pathlib import Path

from policy_platform.extractors.cleaner import clean_paragraphs
from policy_platform.extractors.header_extractor import extract


# ---------------------------------------------------------------------------
# R3 — Noise strip: Version | Page X of Y pattern
# ---------------------------------------------------------------------------

def test_cleaner_drops_version_page_x_of_y():
    """`10/19 Version | Page 2 of 8` style artifacts must be dropped."""
    paragraphs = [
        "Some real body text.",
        "10/19 Version | Page 2 of 8",
        "More body content here for the policy.",
    ]
    cleaned, dropped, _ = clean_paragraphs(paragraphs)
    kept_texts = [p.strip() for p in cleaned if p.strip()]
    assert "Some real body text." in kept_texts
    assert "More body content here for the policy." in kept_texts
    assert "10/19 Version | Page 2 of 8" not in kept_texts
    reasons = {d["reason"] for d in dropped}
    assert "version_page_noise" in reasons


def test_cleaner_drops_bare_page_of_y():
    """Standalone `Page 2 of 8` lines dropped."""
    paragraphs = ["Policy body.", "Page 2 of 8", "More content."]
    cleaned, dropped, _ = clean_paragraphs(paragraphs)
    kept = [p.strip() for p in cleaned if p.strip()]
    assert "Page 2 of 8" not in kept
    # Either category is acceptable: page_number or version_page_noise
    assert any(d["reason"] in ("page_number", "version_page_noise") for d in dropped)


def test_cleaner_drops_standalone_version_keyword():
    """A line starting with `Version` containing a date is dropped."""
    paragraphs = ["Body.", "Version 1.0 - 2024", "Body 2."]
    cleaned, dropped, _ = clean_paragraphs(paragraphs)
    kept = [p.strip() for p in cleaned if p.strip()]
    assert "Body." in kept
    assert "Body 2." in kept
    assert all(not p.startswith("Version") for p in kept)


def test_cleaner_preserves_real_body_text():
    """Ordinary lines with `version` or `page` words must not be stripped."""
    paragraphs = [
        "All employees must review this version of the policy.",
        "Please turn to page 12 of the appendix.",
    ]
    cleaned, dropped, _ = clean_paragraphs(paragraphs)
    kept = [p.strip() for p in cleaned if p.strip()]
    assert "All employees must review this version of the policy." in kept
    assert "Please turn to page 12 of the appendix." in kept


# ---------------------------------------------------------------------------
# R5 — Header extractor: input-derived title
# ---------------------------------------------------------------------------

def test_header_extractor_returns_title_from_largest_line():
    """The extractor picks the largest / all-caps line in the first chunk."""
    paragraphs = [
        "small line",
        "CITY FAMILY HIGH SCHOOL COMPLETION AWARD POLICY",
        "Some intro text",
    ]
    info = extract(
        input_path=Path("nonexistent.docx"),
        pdf_metadata=None,
        cleaned_paragraphs=paragraphs,
    )
    assert "HIGH SCHOOL" in info["title"] or "FAMILY" in info["title"]
    assert info["source"] in ("first_page_largest", "pdf_metadata", "filename")


def test_header_extractor_finds_version_tag():
    """Extractor surfaces a CL&H_xx/yy style version when present."""
    paragraphs = [
        "CITY FAMILY HIGH SCHOOL COMPLETION AWARD POLICY CL&H_02/24",
        "Body",
    ]
    info = extract(
        input_path=Path("x.docx"),
        pdf_metadata=None,
        cleaned_paragraphs=paragraphs,
    )
    assert info["title"] is not None
    assert "CL&H_02/24" == info["version"]


def test_header_extractor_strips_version_from_title():
    """The title returned has the version tag removed (so it stands alone)."""
    paragraphs = [
        "ANNUAL LEAVE POLICY CL&H_05/24",
        "Body",
    ]
    info = extract(
        input_path=Path("x.docx"),
        pdf_metadata=None,
        cleaned_paragraphs=paragraphs,
    )
    assert "CL&H" not in info["title"]
    assert "ANNUAL LEAVE POLICY" in info["title"]


def test_header_extractor_fallback_to_filename():
    """When no first-page line qualifies, falls back to filename stem."""
    paragraphs = []
    info = extract(
        input_path=Path("Annual_Leave_Policy.docx"),
        pdf_metadata=None,
        cleaned_paragraphs=paragraphs,
    )
    assert info["title"] is not None
    assert "Annual" in info["title"] or "Leave" in info["title"]


def test_header_extractor_skips_bullet_lines():
    """Bullet-only lines are not selected as titles."""
    paragraphs = [
        "•",
        "•",
        "REAL POLICY TITLE HERE",
    ]
    info = extract(
        input_path=Path("x.docx"),
        pdf_metadata=None,
        cleaned_paragraphs=paragraphs,
    )
    assert "REAL POLICY TITLE HERE" in info["title"]


# ---------------------------------------------------------------------------
# R4 — Empty-slot skip: bullets without content are removed
# ---------------------------------------------------------------------------

def test_is_empty_slot_distinguishes_real_content_from_bullets():
    """Bullet-only paragraphs should be detected; real prose should not."""
    from docx import Document
    from docx.oxml.ns import qn
    doc = Document()
    # Add a heading + bullet-only paragraphs.
    body = doc.element.body
    # heading paragraph (will populate via add_paragraph)
    h = doc.add_paragraph("1. Purpose")
    b1 = doc.add_paragraph("")
    b2 = doc.add_paragraph("")
    elems = list(body)
    # Pull out paragraphs only
    ps = [e for e in elems if e.tag.split("}")[-1] == "p"]
    from policy_platform.renderer import _is_empty_slot
    assert _is_empty_slot(ps) is True


def test_is_empty_slot_returns_false_for_real_prose():
    """A slot with one real prose paragraph is not empty."""
    from docx import Document
    from docx.oxml.ns import qn
    doc = Document()
    doc.add_paragraph("1. Purpose")
    doc.add_paragraph("This is a meaningful paragraph with enough prose content to qualify as real data here.")
    body = doc.element.body
    ps = [e for e in body if e.tag.split("}")[-1] == "p"]
    from policy_platform.renderer import _is_empty_slot
    assert _is_empty_slot(ps) is False
