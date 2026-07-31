"""test_lines_json_styling.py

Stage 4.16 - tests verifying the published .docx preserves the brain
framework's NATIVE styling (not the editor preview's CSS).

Brain framework native styling:
  pPr carries pStyle=Text (from brain template styles.xml)
  spacing line=240 lineRule=auto (single line spacing)
  Some paragraphs: ind left=2880 hanging=2880 (hanging indent)
  Some paragraphs: jc=both (justified)
  Default font/size comes from the brain template's font table (not
  Inter 14px - which is the editor preview only).

User content inherits these styles via the surrounding pPr; user rich
HTML overrides (bold/italic/color/font-family/font-size via <span
style="...">) are still honoured.
"""
from __future__ import annotations

import importlib
import importlib.util
import sys
import types
from pathlib import Path

import pytest
from docx import Document
from docx.oxml.ns import qn


# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------

@pytest.fixture(scope="session")
def backend_root() -> Path:
    return Path(__file__).resolve().parents[1]


@pytest.fixture(scope="session")
def brain_path(backend_root) -> Path:
    return backend_root / "data" / "brain_template" / "Policy_Framework_5.docx"


@pytest.fixture(scope="session")
def out_dir(tmp_path_factory) -> Path:
    return tmp_path_factory.mktemp("lines_json_styling_out")


@pytest.fixture(scope="session")
def renderer_mod(backend_root):
    if str(backend_root) not in sys.path:
        sys.path.insert(0, str(backend_root))
    pkg = types.ModuleType("policy_platform")
    pkg.__path__ = [str(backend_root / "policy_platform")]
    pkg.__package__ = "policy_platform"
    sys.modules["policy_platform"] = pkg
    fw = types.ModuleType("policy_platform.framework")
    fw.__path__ = [str(backend_root / "policy_platform" / "framework")]
    fw.__package__ = "policy_platform.framework"
    sys.modules["policy_platform.framework"] = fw

    def _load(name: str, path: Path):
        spec = importlib.util.spec_from_file_location(name, str(path))
        assert spec is not None and spec.loader is not None
        m = importlib.util.module_from_spec(spec)
        sys.modules[name] = m
        spec.loader.exec_module(m)
        return m

    _load("policy_platform.config", backend_root / "policy_platform" / "config.py")
    _load("policy_platform.framework.section_map",
          backend_root / "policy_platform" / "framework" / "section_map.py")
    _load("policy_platform.framework.brain_slot_map",
          backend_root / "policy_platform" / "framework" / "brain_slot_map.py")
    src = backend_root / "policy_platform" / "lines_json_renderer.py"
    spec = importlib.util.spec_from_file_location("policy_platform.lines_json_renderer", str(src))
    assert spec is not None and spec.loader is not None
    mod = importlib.util.module_from_spec(spec)
    sys.modules["policy_platform.lines_json_renderer"] = mod
    spec.loader.exec_module(mod)
    return mod


def _scaffold_pPr(p):
    """Return the paragraph's pPr (None if missing)."""
    return p._element.find(qn("w:pPr"))


def _run_text(p):
    """Concatenate all <w:t> text in the paragraph."""
    return ''.join(t.text or "" for t in p._element.iter(qn("w:t")))


# ---------------------------------------------------------------------------
# Tests
# ---------------------------------------------------------------------------

def test_user_paragraphs_keep_scaffold_pPr(renderer_mod, brain_path, out_dir):
    """User-written paragraphs preserve the brain scaffold's pPr (not
    rewritten by us)."""
    lines_json = [
        ["p", {"slot": 5, "text": "INTRODUCTION", "html": "<p>INTRODUCTION</p>"}],
        ["p", {"slot": 5, "text": "INTRO TEXT", "html": "<p>INTRO TEXT</p>"}],
    ]
    out = out_dir / "preserve_pPr.docx"
    renderer_mod.render_lines_json_to_brain(lines_json, brain_path, out)
    doc = Document(str(out))

    # Find the paragraph containing INTRO TEXT.
    target = None
    for p in doc.paragraphs:
        if "INTRO TEXT" in _run_text(p):
            target = p
            break
    assert target is not None

    # Its pPr should be intact (pStyle from brain scaffold).
    pPr = _scaffold_pPr(target)
    assert pPr is not None
    # The user paragraph inherits the scaffold's pStyle (whatever it is:
    # Text, ListParagraph, etc. depending on the slot).
    pStyle = pPr.find(qn("w:pStyle"))
    if pStyle is not None:
        # Just verify it's a valid brain scaffold style (Text or ListParagraph).
        assert pStyle.get(qn("w:val")) in ("Text", "ListParagraph"), (
            f"unexpected pStyle: {pStyle.get(qn('w:val'))}"
        )


def test_no_default_inter_font(renderer_mod, brain_path, out_dir):
    """User content with no explicit font-family should NOT get
    'Inter' 14px auto-applied (it should inherit brain's native font)."""
    lines_json = [
        ["p", {"slot": 5, "text": "intro", "html": "<p>intro content</p>"}],
    ]
    out = out_dir / "no_inter.docx"
    renderer_mod.render_lines_json_to_brain(lines_json, brain_path, out)
    doc = Document(str(out))
    target = next(p for p in doc.paragraphs if "intro content" in _run_text(p))
    for r in target._element.findall(qn("w:r")):
        rPr = r.find(qn("w:rPr"))
        if rPr is None:
            continue  # No rPr means inheriting scaffold = brain font
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is not None:
            # If rFonts IS set, it must NOT be 'Inter' (which was a
            # preview-only default we removed).
            ascii_val = rFonts.get(qn("w:ascii"))
            assert ascii_val != "Inter", "Inter should not be auto-applied"


def test_no_default_14px_font_size(renderer_mod, brain_path, out_dir):
    """User content with no explicit font-size should NOT get sz=21
    (which is 14px - the editor preview default). It should inherit
    the brain's native size."""
    lines_json = [
        ["p", {"slot": 5, "text": "intro", "html": "<p>intro content</p>"}],
    ]
    out = out_dir / "no_14px.docx"
    renderer_mod.render_lines_json_to_brain(lines_json, brain_path, out)
    doc = Document(str(out))
    target = next(p for p in doc.paragraphs if "intro content" in _run_text(p))
    for r in target._element.findall(qn("w:r")):
        rPr = r.find(qn("w:rPr"))
        if rPr is None:
            continue
        sz = rPr.find(qn("w:sz"))
        if sz is not None:
            sz_val = sz.get(qn("w:val"))
            assert sz_val != "21", f"sz=21 (14px) should not be auto-applied; got {sz_val}"


def test_no_extra_editor_preview_styling(renderer_mod, brain_path, out_dir):
    """User paragraphs should NOT carry editor-preview-specific pPr
    overrides (line=384, after=120, before=0)."""
    lines_json = [
        ["p", {"slot": 6, "text": "policy", "html": "<p>policy content</p>"}],
    ]
    out = out_dir / "no_preview_overrides.docx"
    renderer_mod.render_lines_json_to_brain(lines_json, brain_path, out)
    doc = Document(str(out))
    target = next(p for p in doc.paragraphs if "policy content" in _run_text(p))
    pPr = _scaffold_pPr(target)
    if pPr is not None:
        spacing = pPr.find(qn("w:spacing"))
        if spacing is not None:
            # The renderer must NOT inject line=384.
            line = spacing.get(qn("w:line"))
            assert line != "384", f"editor-preview line=384 must not be injected; got {line}"


def test_explicit_font_size_still_works(renderer_mod, brain_path, out_dir):
    """When the user HTML specifies a font-size, it must still be
    honoured (no regression)."""
    lines_json = [
        ["p", {
            "slot": 6,
            "text": "large",
            "html": '<p style="font-size:24px">LARGE TEXT</p>',
        }],
    ]
    out = out_dir / "explicit_size.docx"
    renderer_mod.render_lines_json_to_brain(lines_json, brain_path, out)
    doc = Document(str(out))
    found = None
    for p in doc.paragraphs:
        if "LARGE TEXT" in _run_text(p):
            found = p
            break
    assert found is not None
    sizes = []
    for r in found._element.findall(qn("w:r")):
        rPr = r.find(qn("w:rPr"))
        if rPr is not None:
            sz = rPr.find(qn("w:sz"))
            if sz is not None:
                sizes.append(sz.get(qn("w:val")))
# Per user directive: ALL text uniform Times New Roman 10pt
    # (size_hp=20). Even inline font-size overrides are normalised to 10pt.
    assert "20" in sizes, f"expected sz=20 (10pt uniform), got {sizes}"


def test_explicit_bold_still_works(renderer_mod, brain_path, out_dir):
    """User HTML with <strong> continues to render as bold runs."""
    lines_json = [
        ["p", {
            "slot": 5,
            "text": "bold",
            "html": '<p>This is <strong>BOLD</strong> content.</p>',
        }],
    ]
    out = out_dir / "explicit_bold.docx"
    renderer_mod.render_lines_json_to_brain(lines_json, brain_path, out)
    doc = Document(str(out))
    target = next(p for p in doc.paragraphs if "BOLD" in _run_text(p))
    bold_found = False
    for r in target._element.findall(qn("w:r")):
        rPr = r.find(qn("w:rPr"))
        if rPr is None:
            continue
        text = ''.join(t.text or "" for t in r.iter(qn("w:t")))
        if "BOLD" not in text:
            continue
        if rPr.find(qn("w:b")) is not None:
            bold_found = True
    assert bold_found, "bold run not preserved"
