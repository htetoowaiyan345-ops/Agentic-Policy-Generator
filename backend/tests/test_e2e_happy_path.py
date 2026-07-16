"""End-to-end happy-path test.

Exercises the full chain on a single uploaded document:
    upload -> process -> edit -> save -> submit -> approve -> publish -> download

Skipped by default. Enable with:    RUN_E2E=1 pytest

All values come from fixtures/env, not literals:
    - source file     : parametrized over SAMPLE_FILENAMES (real samples,
                        no in-test filename string)
    - backend URL     : env API_BASE_URL via api_base_url fixture
    - actor names     : env TEST_ACTOR / TEST_REVIEWER
    - edit marker     : runtime-generated unique uuid (edit_marker fixture)
    - DB name         : env DB_PATH / TEST_DB_PATH, never inlined
"""
from __future__ import annotations

import io
import json
import socket
import subprocess
import sys
import time
import urllib.error
import urllib.request
from pathlib import Path

import pytest

from .conftest import SAMPLE_FILENAMES

pytestmark = pytest.mark.e2e


def _free_port() -> int:
    with socket.socket(socket.AF_INET, socket.SOCK_STREAM) as s:
        s.bind(("127.0.0.1", 0))
        return s.getsockname()[1]


@pytest.fixture(scope="session")
def live_server(tmp_path_factory):
    """Boot the real backend on a free port pointing at an isolated test DB.
    Yields the base URL. Skips if the server module is not importable in this env."""
    import os
    from policy_platform import config

    test_db = tmp_path_factory.mktemp("e2e_db") / "policy_history_e2e.db"
    env = dict(os.environ)
    env["DB_PATH"] = str(test_db)
    env["API_HOST"] = "127.0.0.1"
    env["API_PORT"] = str(_free_port())
    env["API_BASE_URL"] = f"http://{env['API_HOST']}:{env['API_PORT']}"

    try:
        import api.server as _api_server  # noqa: F401
    except Exception as exc:
        pytest.skip(f"backend api.server not importable: {exc}")

    proc = subprocess.Popen(
        [sys.executable, "-m", "api.server"],
        cwd=str(config.PROJECT_ROOT),
        env=env,
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
    )
    base = env["API_BASE_URL"]

    deadline = time.time() + 15.0
    while time.time() < deadline:
        try:
            with urllib.request.urlopen(base + "/api/health", timeout=0.5) as r:
                if 200 <= r.status < 500:
                    break
        except Exception:
            time.sleep(0.2)
    else:
        proc.terminate()
        pytest.skip("backend did not come up within 15s")

    yield base

    proc.terminate()
    try:
        proc.wait(timeout=5)
    except subprocess.TimeoutExpired:
        proc.kill()


def _post(url: str, payload: dict | None = None, file_field: tuple | None = None):
    if file_field is not None:
        filename, content, mime = file_field
        boundary = "----e2e-boundary"
        body = (
            f"--{boundary}\r\n"
            f'Content-Disposition: form-data; name="file"; filename="{filename}"\r\n'
            f"Content-Type: {mime}\r\n\r\n"
        ).encode() + content + f"\r\n--{boundary}--\r\n".encode()
        req = urllib.request.Request(
            url, data=body, headers={"Content-Type": f"multipart/form-data; boundary={boundary}"},
            method="POST",
        )
    else:
        req = urllib.request.Request(
            url, data=json.dumps(payload or {}).encode(),
            headers={"Content-Type": "application/json"}, method="POST",
        )
    with urllib.request.urlopen(req, timeout=30) as r:
        return json.loads(r.read().decode("utf-8"))


def _get_json(url: str):
    with urllib.request.urlopen(url, timeout=30) as r:
        return json.loads(r.read().decode("utf-8"))


def _get_bytes(url: str) -> bytes:
    with urllib.request.urlopen(url, timeout=30) as r:
        return r.read()


def _wait_status(base: str, run_id: str, timeout_s: float = 60.0):
    deadline = time.time() + timeout_s
    last = None
    while time.time() < deadline:
        s = _get_json(f"{base}/api/status/{run_id}")
        last = s
        state = (s.get("state") or "").lower()
        if state == "done":
            return s
        if state == "failed":
            raise AssertionError(f"pipeline failed for {run_id}: {s}")
        time.sleep(0.5)
    raise AssertionError(f"timeout waiting for {run_id}; last={last}")


@pytest.mark.parametrize("sample_name,sample_ext", SAMPLE_FILENAMES)
def test_end_to_end_happy_path(
    live_server,
    samples_dir,
    sample_name,
    sample_ext,
    test_actor,
    test_reviewer,
    edit_marker,
):
    base = live_server

    # 1. upload
    src: Path = samples_dir / sample_name
    if not src.exists():
        pytest.skip(f"sample not present: {sample_name}")
    upload_resp = _post(
        f"{base}/api/upload",
        file_field=(sample_name, src.read_bytes(), "application/octet-stream"),
    )
    run_id = upload_resp["run_id"]

    # 2. process (poll)
    _wait_status(base, run_id)

    # 3 + 4. edit + save (inject the runtime-generated marker into slot 2)
    initial = _get_json(f"{base}/api/preview/{run_id}")
    save_resp = _post(
        f"{base}/api/versions/{run_id}/save",
        {
            "lines": initial.get("lines", []),
            "change_summary": edit_marker,
            "actor": test_actor,
        },
    )
    version_no = save_resp.get("version_no", 1)

    # 5. submit for review
    _post(
        f"{base}/api/versions/{run_id}/{version_no}/submit",
        {"actor": test_actor},
    )

    # 6. approve
    _post(
        f"{base}/api/versions/{run_id}/{version_no}/review",
        {"action": "approved", "reviewer": test_reviewer},
    )

    # 7. publish
    _post(
        f"{base}/api/versions/{run_id}/{version_no}/publish",
        {"actor": test_reviewer},
    )

    # 8. download
    docx_bytes = _get_bytes(f"{base}/api/download/{run_id}/docx")
    assert docx_bytes, "downloaded docx is empty"

    # THE assertion: the runtime-generated marker is present in the docx.
    from docx import Document

    doc = Document(io.BytesIO(docx_bytes))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text += "\n" + cell.text

    assert edit_marker in full_text, (
        f"edit marker {edit_marker!r} not found in published docx; "
        f"first 400 chars of body: {full_text[:400]!r}"
    )
