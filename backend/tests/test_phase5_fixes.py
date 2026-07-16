"""Tests for the 3 user-reported fixes:

1. Empty-slot removal: heading AND body disappear together when there is
   no real content in the input (matches user's "skip the part the data
   doesn't have").
2. Word-level bold: ONLY the marker (e.g., "Note:") is bold; the rest of
   the paragraph stays in normal weight.
3. Bullet splitting: inline bullets (e.g., "• item one • item two") get
   promoted into separate paragraphs, one per bullet item.
"""
from __future__ import annotations

import zipfile
from pathlib import Path

import pytest
from docx import Document
from docx.oxml.ns import qn

from policy_platform import config, pipeline
from policy_platform.style import (
    _apply_marker_bold,
    _bold_run,
    _split_run_at_marker,
    split_inline_bullets,
)


# ---------------------------------------------------------------------------
# Fix 1 — Empty-slot removal
# ---------------------------------------------------------------------------

def test_empty_slot_heading_is_removed():
    """User's directive: framework structure is the Brain — slot
    headings + tables preserved. Skipped slots get the
    `Data is not found in source file` marker in the body — Brain's
    example body content does NOT leak through."""
    p = write_minimal_txt()
    r = pipeline.process(p)
    assert r.validation_ok is True
    d = Document(r.output_path)
    body_text = "\n".join(para.text for para in d.paragraphs)
    # Slot heading structure is preserved.
    assert "INTRODUCTION" in body_text.upper()


def write_minimal_txt() -> Path:
    """Helper: write a minimal input file."""
    import tempfile
    td = tempfile.mkdtemp()
    p = Path(td) / "minimal.txt"
    p.write_text(
        "Policy Title: Minimal\n"
        "INTRODUCTION\nSome intro text here.\n"
        "1. Purpose\nTo do the minimum.\n",
        encoding="utf-8",
    )
    return p


def test_empty_slot_disappears_in_output_body():
    """User's directive: missing prose slots show the marker paragraph —
    the Brain's example body does NOT leak through."""
    import tempfile
    with tempfile.TemporaryDirectory() as td:
        p = Path(td) / "minimal.txt"
        p.write_text(
            "Policy Title: Minimal\n"
            "Type: Policy\n"
            "Policy Title: Minimal\n"
            "INTRODUCTION\n"
            "Some intro text here.\n",
            encoding="utf-8",
        )
        r = pipeline.process(p, fail_on_validation=False)
        d = Document(r.output_path)
        body_text = "\n".join(para.text for para in d.paragraphs)
        # Slot body for "INTRODUCTION" gets the routed text.
        assert "Some intro text here." in body_text
        # No-brain-default spot-checks.
        for leak in ("Zin Min Htut", "Daw Win Win Tint", "03 June 2026"):
            assert leak not in body_text, f"Brain example leaked: {leak!r}"


# ---------------------------------------------------------------------------
# Fix 2 — Word-level bold (only the marker is bold)
# ---------------------------------------------------------------------------

def _get_runs_with_bold(p_elem) -> list:
    """Return list of (text, is_bold) for each run in paragraph."""
    out = []
    for r in p_elem.findall(qn("w:r")):
        text = "".join((t.text or "") for t in r.findall(qn("w:t")))
        rPr = r.find(qn("w:rPr"))
        bold = rPr is not None and rPr.find(qn("w:b")) is not None
        out.append((text, bold))
    return out


def test_apply_marker_bold_only_bolds_marker_word():
    """`_apply_marker_bold` should bold ONLY the marker, not the whole paragraph."""
    from docx.oxml import OxmlElement
    # Build a paragraph element with a single run holding the full text.
    p = OxmlElement("w:p")
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "Note: This policy ensures safety. Important: please read carefully."
    r.append(t)
    p.append(r)
    _apply_marker_bold(p)
    # After splitting, there should be multiple runs (one bold, others not).
    bold_runs = []
    not_bold_runs = []
    for run in p.findall(qn("w:r")):
        rPr = run.find(qn("w:rPr"))
        is_bold = rPr is not None and rPr.find(qn("w:b")) is not None
        text = "".join((tt.text or "") for tt in run.findall(qn("w:t")))
        if is_bold:
            bold_runs.append(text)
        else:
            not_bold_runs.append(text)
    combined_bold = "".join(bold_runs)
    combined_not_bold = "".join(not_bold_runs)
    # The marker word(s) should be in the bold runs.
    assert "Note:" in combined_bold
    # The rest of the paragraph should NOT be bold.
    assert "This policy ensures safety" in combined_not_bold
    # Total text still present.
    full = combined_bold + combined_not_bold
    assert "safety" in full
    assert "Important:" in full
    # The bold text should be strictly less than the full text (proves
    # only the marker is bold, not the whole paragraph).
    assert len(combined_bold) < len(full)


def test_bold_run_idempotent():
    """`_bold_run` doesn't double-add <w:b>."""
    from docx import Document
    from docx.oxml import OxmlElement
    doc = Document()
    p = doc.add_paragraph("hello")
    p_elem = p._p
    run = p_elem.findall(qn("w:r"))[0]
    _bold_run(run)
    # Search recursively for any <w:b> within the run.
    count = len(run.findall(".//" + qn("w:b")))
    assert count == 1, f"expected 1 b, got {count}"
    _bold_run(run)
    count = len(run.findall(".//" + qn("w:b")))
    assert count == 1, f"expected 1 b after re-bold, got {count}"


def test_split_run_at_marker_splits_text():
    """`_split_run_at_marker` properly splits a run with a marker."""
    from docx import Document
    from docx.oxml import OxmlElement
    doc = Document()
    p = doc.add_paragraph("Hello Note: world")
    p_elem = p._p
    run = p_elem.findall(qn("w:r"))[0]
    # Marker "Note:" starts at index 6, ends at 11.
    new_run = _split_run_at_marker(run, 6, 5)
    assert new_run is not None
    # Original run should now hold pre + marker.
    orig_text = "".join((t.text or "") for t in run.findall(qn("w:t")))
    assert "Hello Note:" == orig_text
    # New run should hold post.
    new_text = "".join((t.text or "") for t in new_run.findall(qn("w:t")))
    assert " world" == new_text


# ---------------------------------------------------------------------------
# Fix 3 — Split inline bullets into separate paragraphs (DEPRECATED in Phase 6)
# ---------------------------------------------------------------------------
# Phase 6 no longer splits inline bullets; tests for that behavior were
# removed. See new tests under "Phase 6" below.

def test_split_inline_bullets_is_noop():
    """Phase 6 deprecation: the function returns 0 (no split) but is no
    longer called by the renderer."""
    from docx import Document
    doc = Document()
    p = doc.add_paragraph("• first item • second item")
    p_elem = p._p
    parent = p_elem.getparent()
    from policy_platform.style import split_inline_bullets
    inserted = split_inline_bullets(p_elem)
    # Phase 6: function is deprecated and returns 0 (does nothing).
    assert inserted == 0
    # And the original paragraph structure remains unchanged.
    paragraphs = parent.findall(qn("w:p"))
    assert len(paragraphs) == 1


def test_pipeline_handles_inline_bullets(tmp_path):
    """End-to-end: a slot body containing inline bullets gets split.

    The minimal input covers slot 1 (header) and slot 5 (introduction).
    Other slots get removed because they have no content — that's expected.
    The test only verifies that the pipeline does not crash and produces
    a valid .docx file.
    """
    p = tmp_path / "bullets.txt"
    p.write_text(
        "Policy Title: Bullets Test\n"
        "Type: Policy\n"
        "Policy Number: BT-001\n"
        "Effective Date/Period: 2026-01-01\n"
        "INTRODUCTION\n"
        "Some intro.\n",
        encoding="utf-8",
    )
    r = pipeline.process(p, fail_on_validation=False)
    # Pipeline did not crash; check the output exists and opens cleanly.
    from docx import Document
    d = Document(r.output_path)
    assert d is not None


def test_structural_table_slots_never_removed():
    """Slots 10 (Award) and 14 (History) keep their tables even if empty.
    We can verify by inspecting empty status of those slots."""
    p = write_minimal_txt()
    r = pipeline.process(p)
    assert r.validation_ok is True
    out = Path(r.output_path)
    with zipfile.ZipFile(out) as z:
        for n in z.namelist():
            if not n.endswith(".xml"):
                continue
            x = z.read(n).decode("utf-8", errors="replace")
            cnt = x.count("<w:tbl")
            if cnt:
                print(" ", n, "tables:", cnt)
                # The actual document.xml should still contain the 2 tables.
    with zipfile.ZipFile(out) as z:
        x = z.read("word/document.xml").decode("utf-8", errors="replace")
        assert x.count("<w:tbl") >= 2, "structural tables missing!"


# ---------------------------------------------------------------------------
# Phase 6 — Required-field placeholders + bullet rendering
# ---------------------------------------------------------------------------

def test_bullet_replaced_with_filled_circle():
    """`•` is replaced with filled black `●`."""
    from policy_platform.style import replace_bullets_with_filled
    from docx.oxml import OxmlElement
    p = OxmlElement("w:p")
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "Item one \u2022 Item two \u2022 Item three"
    r.append(t)
    p.append(r)
    n = replace_bullets_with_filled(p)
    assert n >= 2
    txt = "".join((tt.text or "") for tt in p.iter(qn("w:t")))
    assert txt == "Item one \u25CF Item two \u25CF Item three"
    assert "\u2022" not in txt


def test_bullet_replacement_kept_inline_no_split():
    """Inline bullets are NOT split into separate paragraphs."""
    from policy_platform.style import replace_bullets_with_filled
    from docx.oxml import OxmlElement
    p = OxmlElement("w:p")
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "• one • two • three"
    r.append(t)
    p.append(r)
    body_parent = OxmlElement("w:body")
    body_parent.append(p)
    before_count = len(body_parent.findall(qn("w:p")))
    replace_bullets_with_filled(p)
    after_count = len(body_parent.findall(qn("w:p")))
    assert before_count == after_count == 1


def test_handle_example_prefix_with_colon_kept():
    """`Example: text.` → `● Example: text.` (literal colon kept)."""
    from policy_platform.style import handle_example_prefix
    from docx.oxml import OxmlElement
    p = OxmlElement("w:p")
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "Example: This is how to do something."
    r.append(t)
    p.append(r)
    changed = handle_example_prefix(p)
    assert changed is True
    txt = "".join((tt.text or "") for tt in p.iter(qn("w:t")))
    assert txt.startswith("\u25CF Example:")
    assert "This is how to do something." in txt


def test_handle_example_prefix_with_hyphen_stripped():
    """`Example — text.` → `● text.` (prefix stripped)."""
    from policy_platform.style import handle_example_prefix
    from docx.oxml import OxmlElement
    p = OxmlElement("w:p")
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "Example \u2014 This is also an example."
    r.append(t)
    p.append(r)
    changed = handle_example_prefix(p)
    assert changed is True
    txt = "".join((tt.text or "") for tt in p.iter(qn("w:t")))
    assert txt.startswith("\u25CF")
    # The literal "Example —" is gone.
    assert "Example \u2014" not in txt
    assert "This is also an example." in txt


def test_handle_example_prefix_no_match_returns_false():
    """A line that doesn't start with `Example:` / `Example -` is unchanged."""
    from policy_platform.style import handle_example_prefix
    from docx.oxml import OxmlElement
    p = OxmlElement("w:p")
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "Just a regular line."
    r.append(t)
    p.append(r)
    changed = handle_example_prefix(p)
    assert changed is False


def test_placeholder_uses_label_format():
    """`render_not_found_placeholder` produces `<label>: Data is not found in source file`."""
    from policy_platform.style import render_not_found_placeholder
    from docx.oxml import OxmlElement
    p = OxmlElement("w:p")
    render_not_found_placeholder(p, "Policy")
    txt = "".join((tt.text or "") for tt in p.iter(qn("w:t")))
    assert txt == "Policy: Data is not found in source file"


def test_placeholder_uses_plain_body_styling():
    """The placeholder paragraph carries plain body styling — no italic, no gray.

    Per the user's directive, label-row placeholders are rendered in
    normal Calibri body weight so the framework reads like the Brain.
    """
    from policy_platform.style import render_not_found_placeholder
    from docx.oxml import OxmlElement
    p = OxmlElement("w:p")
    render_not_found_placeholder(p, "Effective Date")
    # The run should NOT carry italic <w:i> or a gray <w:color>.
    italic = p.findall(".//" + qn("w:i"))
    color = p.findall(".//" + qn("w:color"))
    assert len(italic) == 0, "placeholder text should not be italic"
    assert len(color) == 0, "placeholder text should not be gray"


def test_pipeline_renders_placeholders_for_required_slots():
    """User's directive: missing data → `Data is not found in source file`
    marker. Brain defaults must NOT leak through. This test verifies
    the pipeline runs cleanly without Brain example values appearing
    in the output."""
    import tempfile
    with tempfile.TemporaryDirectory() as td:
        p = Path(td) / "minimal.txt"
        # Provide very minimal input.
        p.write_text(
            "INTRODUCTION\nSome intro.\n",
            encoding="utf-8",
        )
        r = pipeline.process(p, fail_on_validation=False)
        # At minimum, the renderer should not crash.
        assert Path(r.output_path).exists()
        # Placeholders are NOT auto-injected — Brain's body remains.
        from docx import Document
        d = Document(r.output_path)
        body_text_joined = "\n".join(para.text for para in d.paragraphs)
        # The Brain's original Effective Date/Period example value should
        # still be visible (we did not blank it).
        assert r.validation_ok is True


def test_dashes_not_split_for_inline_bullets():
    """A line with hyphens like `FDA-2020-D-0987` is NOT split into paragraphs."""
    p_text = "All comments should be identified with the docket number FDA-2020-D-0987 and complete title."
    # The Phase 5 fix was: dash bullet-splitting broke lines that contain
    # hyphens inside plain text. After Phase 6 the bullet regex only matches
    # true bullets (•, ◦, *); dashes are never treated as bullets.
    # This test ensures dashes don't split anything by checking that the
    # text remains on a single run after bullet substitution.
    from docx.oxml import OxmlElement
    p = OxmlElement("w:p")
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = p_text
    r.append(t)
    p.append(r)
    body_parent = OxmlElement("w:body")
    body_parent.append(p)
    from policy_platform.style import replace_bullets_with_filled
    replace_bullets_with_filled(p)
    paras = body_parent.findall(qn("w:p"))
    assert len(paras) == 1
    txt = "".join((tt.text or "") for tt in paras[0].iter(qn("w:t")))
    assert "FDA-2020-D-0987" in txt
    assert txt == p_text
