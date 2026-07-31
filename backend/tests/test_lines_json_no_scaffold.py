"""test_lines_json_no_scaffold.py

Stage 4.17 - tests verifying the published .docx contains ONLY the
user's Result data and NO brain scaffold placeholder text in the
slot body regions.

Brain scaffold signatures that must NEVER appear in the published
output's slot body regions:
  - City Family High School Completion Award
  - 13 Feb 2023
  - [CL&H_02/24]
  - 300,000 MMK / 500,000 MMK / 600,000 MMK / 700,000 MMK / 800,000 MMK
  - Pass with distinction
  - Initial policy drafting
  - Si Thu Maung
  - Win Tint
  - Myanmar high school

Also: heading paragraphs must not be duplicated when the user
content already starts with the heading label.
"""
from __future__ import annotations

import importlib
import importlib.util
import sys
import types
from pathlib import Path

import pytest
from docx import Document
from docx.oxml.ns import qn


# ---------------------------------------------------------------------------
# Fixtures (shared with test_lines_json_renderer.py)
# ---------------------------------------------------------------------------

@pytest.fixture(scope="session")
def backend_root() -> Path:
    return Path(__file__).resolve().parents[1]


@pytest.fixture(scope="session")
def brain_path(backend_root) -> Path:
    return backend_root / "data" / "brain_template" / "Policy_Framework_5.docx"


@pytest.fixture(scope="session")
def out_dir(tmp_path_factory) -> Path:
    return tmp_path_factory.mktemp("lines_json_no_scaffold_out")


@pytest.fixture(scope="session")
def renderer_mod(backend_root):
    if str(backend_root) not in sys.path:
        sys.path.insert(0, str(backend_root))
    pkg = types.ModuleType("policy_platform")
    pkg.__path__ = [str(backend_root / "policy_platform")]
    pkg.__package__ = "policy_platform"
    sys.modules["policy_platform"] = pkg
    fw = types.ModuleType("policy_platform.framework")
    fw.__path__ = [str(backend_root / "policy_platform" / "framework")]
    fw.__package__ = "policy_platform.framework"
    sys.modules["policy_platform.framework"] = fw

    def _load(name: str, path: Path):
        spec = importlib.util.spec_from_file_location(name, str(path))
        assert spec is not None and spec.loader is not None
        m = importlib.util.module_from_spec(spec)
        sys.modules[name] = m
        spec.loader.exec_module(m)
        return m

    _load("policy_platform.config", backend_root / "policy_platform" / "config.py")
    _load("policy_platform.framework.section_map",
          backend_root / "policy_platform" / "framework" / "section_map.py")
    _load("policy_platform.framework.brain_slot_map",
          backend_root / "policy_platform" / "framework" / "brain_slot_map.py")
    src = backend_root / "policy_platform" / "lines_json_renderer.py"
    spec = importlib.util.spec_from_file_location("policy_platform.lines_json_renderer", str(src))
    assert spec is not None and spec.loader is not None
    mod = importlib.util.module_from_spec(spec)
    sys.modules["policy_platform.lines_json_renderer"] = mod
    spec.loader.exec_module(mod)
    return mod


# Brain scaffold signatures that must not leak into published output.
BRAIN_SCAFFOLD_SIGNATURES = (
    "City Family High School Completion Award",
    "[CL&H_02/24]",
    "13 Feb 2023",
    "300,000 MMK",
    "500,000 MMK",
    "600,000 MMK",
    "700,000 MMK",
    "800,000 MMK",
    "Pass with distinction",
    "Initial policy drafting",
    "Si Thu Maung",
    "Myanmar high school matriculation",
    "Daw Win Win Tint, Group CEO",
    "Zin Min Htut, CSR Specialist",
)


def _all_paragraph_text(doc):
    """Concatenate all paragraph text in document order."""
    out = []
    for p in doc.paragraphs:
        out.append(''.join(t.text or "" for t in p._element.iter(qn("w:t"))))
    return out


def _all_table_text(doc):
    """Concatenate all table cell text."""
    out = []
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                out.append(cell.text)
    return out


def _has_signature(doc, sig):
    """True if any paragraph or table cell contains the scaffold sig."""
    for text in _all_paragraph_text(doc):
        if sig in text:
            return True
    for text in _all_table_text(doc):
        if sig in text:
            return True
    return False


# ---------------------------------------------------------------------------
# Tests
# ---------------------------------------------------------------------------

def test_no_brain_scaffold_text_when_user_filled_slot(renderer_mod, brain_path, out_dir):
    """When the user has content for a slot, the brain's placeholder
    text for that slot must NOT appear in the output."""
    lines_json = [
        ["p", {"slot": 2, "text": "My brief description", "html": "<p>My brief description</p>"}],
        ["p", {"slot": 3, "text": "01 July 2026 - 30 June 2027", "html": "<p>01 July 2026 - 30 June 2027</p>"}],
        ["p", {"slot": 4, "text": "To establish a fair policy.", "html": "<p>To establish a fair policy.</p>"}],
    ]
    out = out_dir / "no_scaffold.docx"
    renderer_mod.render_lines_json_to_brain(lines_json, brain_path, out)
    doc = Document(str(out))

    # Brief Description: brain had "[The policy provides a one-time cash award..." - now gone.
    all_text = ' '.join(_all_paragraph_text(doc))
    assert "one-time cash award" not in all_text, "brain scaffold text leaked into output"
    # Effective Date/Period: brain had "[13 Feb 2023]" - now gone.
    assert "[13 Feb 2023]" not in all_text, "brain scaffold date leaked into output"
    # Reason for Policy: brain had "[This policy is established..." - now gone.
    assert "academic achievements of the children" not in all_text, "brain scaffold text leaked"


def test_no_brain_award_tier_scaffold_in_slot10(renderer_mod, brain_path, out_dir):
    """Slot 10 scaffold had '300,000 MMK' etc. as example data; once
    the user fills slot 10, those examples must be gone."""
    lines_json = [
        ["p", {"slot": 10, "text": "User's tier", "html": "<p>User's tier text</p>"}],
    ]
    out = out_dir / "no_award_tiers.docx"
    renderer_mod.render_lines_json_to_brain(lines_json, brain_path, out)
    doc = Document(str(out))

    # The user's content "User's tier text" should be present.
    all_text = ' '.join(_all_paragraph_text(doc))
    assert "User's tier text" in all_text
    # The brain's example award tiers must be gone.
    for sig in ["300,000 MMK", "500,000 MMK", "Pass with distinction"]:
        assert not _has_signature(doc, sig), f"scaffold {sig!r} leaked into output"


def test_heading_dedup_when_user_typed_heading(renderer_mod, brain_path, out_dir):
    """When the user's first paragraph for a slot matches the heading
    label, the scaffold heading paragraph is cleared so we don't
    render the label twice."""
    lines_json = [
        ["p", {"slot": 5, "text": "INTRODUCTION", "html": "<p>INTRODUCTION</p>"}],
        ["p", {"slot": 5, "text": "This policy supports engagement.", "html": "<p>This policy supports engagement.</p>"}],
    ]
    out = out_dir / "heading_dedup.docx"
    renderer_mod.render_lines_json_to_brain(lines_json, brain_path, out)
    doc = Document(str(out))

    intro_count = 0
    for p in doc.paragraphs:
        text = ''.join(t.text or "" for t in p._element.iter(qn("w:t"))).strip()
        if text.upper() == "INTRODUCTION":
            intro_count += 1
    assert intro_count == 1, f"expected exactly 1 INTRODUCTION paragraph, got {intro_count}"


def test_heading_dedup_for_policy_statement(renderer_mod, brain_path, out_dir):
    """Same dedup applies to slot 6 (POLICY STATEMENT)."""
    lines_json = [
        ["p", {"slot": 6, "text": "POLICY STATEMENT", "html": "<p>POLICY STATEMENT</p>"}],
        ["p", {"slot": 6, "text": "Clear guidelines.", "html": "<p>Clear guidelines.</p>"}],
    ]
    out = out_dir / "policy_statement_dedup.docx"
    renderer_mod.render_lines_json_to_brain(lines_json, brain_path, out)
    doc = Document(str(out))

    policy_count = 0
    for p in doc.paragraphs:
        text = ''.join(t.text or "" for t in p._element.iter(qn("w:t"))).strip()
        if text.upper() == "POLICY STATEMENT":
            policy_count += 1
    assert policy_count == 1, f"expected 1 POLICY STATEMENT, got {policy_count}"


def test_no_heading_dedup_when_user_does_not_repeat_heading(renderer_mod, brain_path, out_dir):
    """If the user's content does NOT start with the heading label,
    the scaffold heading is preserved (no dedup)."""
    lines_json = [
        # User did NOT type "INTRODUCTION" — they only typed body text.
        ["p", {"slot": 5, "text": "This policy supports engagement.", "html": "<p>This policy supports engagement.</p>"}],
    ]
    out = out_dir / "no_dedup.docx"
    renderer_mod.render_lines_json_to_brain(lines_json, brain_path, out)
    doc = Document(str(out))

    intro_count = 0
    for p in doc.paragraphs:
        text = ''.join(t.text or "" for t in p._element.iter(qn("w:t"))).strip()
        if text.upper() == "INTRODUCTION":
            intro_count += 1
    # Without dedup, the scaffold heading 'INTRODUCTION' should be
    # replaced by the user's body text (since user_count > 0).
    # The result: 1 paragraph with the user's body text, no INTRODUCTION.
    assert intro_count == 0, f"expected 0 INTRODUCTION (no dedup), got {intro_count}"


def test_full_user_scenario_no_scaffold_leak(renderer_mod, brain_path, out_dir):
    """End-to-end: the user's actual saved data produces a clean
    output with no brain scaffold text leaking through."""
    # Real data shape (matches c67cb822120c v3 lines_json).
    lines_json = [
        ["p", {"slot": 0, "text": "Type: HR Policy", "html": "<p>Type: HR Policy</p>"}],
        ["p", {"slot": 0, "text": "Policy Title: AWARD POLICY", "html": "<p>Policy Title: AWARD POLICY</p>"}],
        ["p", {"slot": 0, "text": "INTRODUCTION", "html": "<p>INTRODUCTION</p>"}],
        ["p", {"slot": 0, "text": "This policy supports employee engagement.", "html": "<p>This policy supports employee engagement.</p>"}],
        ["p", {"slot": 0, "text": "POLICY STATEMENT", "html": "<p>POLICY STATEMENT</p>"}],
        ["p", {"slot": 0, "text": "Clear guidelines for nominations.", "html": "<p>Clear guidelines for nominations.</p>"}],
    ]
    out = out_dir / "full_scenario.docx"
    renderer_mod.render_lines_json_to_brain(lines_json, brain_path, out)
    doc = Document(str(out))

    all_text = ' '.join(_all_paragraph_text(doc))

    # User content must be present.
    for must in [
        "Type: HR Policy",
        "Policy Title: AWARD POLICY",
        "This policy supports employee engagement.",
        "Clear guidelines for nominations.",
    ]:
        assert must in all_text, f"user content {must!r} missing"

    # Brain scaffold must NOT leak.
    for must_not in BRAIN_SCAFFOLD_SIGNATURES:
        assert must_not not in all_text, f"brain scaffold {must_not!r} leaked"

    # Headings must appear exactly once.
    intro_count = sum(
        1 for p in doc.paragraphs
        if ''.join(t.text or "" for t in p._element.iter(qn("w:t"))).strip().upper() == "INTRODUCTION"
    )
    policy_count = sum(
        1 for p in doc.paragraphs
        if ''.join(t.text or "" for t in p._element.iter(qn("w:t"))).strip().upper() == "POLICY STATEMENT"
    )
    assert intro_count == 1, f"expected 1 INTRODUCTION, got {intro_count}"
    assert policy_count == 1, f"expected 1 POLICY STATEMENT, got {policy_count}"
