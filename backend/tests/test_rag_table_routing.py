"""Tests for table passthrough routing."""
from __future__ import annotations

from policy_platform.rag.table_routing import (
    TABLE_SLOTS,
    TABLE_SLOT_SIGNALS,
    find_table_for_slot,
)


def test_table_slots_are_9_10_14():
    assert TABLE_SLOTS == {9, 10, 14}


def test_label_row_slots_not_in_table_set():
    for sid in (1, 2, 3, 4, 5, 6, 7, 8, 11, 12, 13, 15):
        assert sid not in TABLE_SLOTS


def test_table_with_award_signal_routes_to_slot_10():
    table = [
        ["Tier", "Damage Level", "Payout (MMK)"],
        ["1", "Minor", "100,000"],
        ["2", "Severe", "1,000,000"],
    ]
    result = find_table_for_slot(10, [table])
    assert result is not None
    assert result == table  # whole table passed through


def test_table_with_history_signal_routes_to_slot_14():
    table = [
        ["Version", "Date", "Description"],
        ["v1.0", "2024-01-15", "Initial release"],
        ["v1.1", "2024-06-15", "Updated"],
    ]
    result = find_table_for_slot(14, [table])
    assert result is not None
    assert result == table


def test_table_with_exclusion_signal_routes_to_slot_9():
    table = [
        ["Exclusion Type", "Description"],
        ["Type A", "Not covered"],
        ["Type B", "Exception case"],
    ]
    result = find_table_for_slot(9, [table])
    assert result is not None
    assert result == table


def test_irrelevant_table_not_matched():
    table = [
        ["Header"],
        ["Some random content"],
    ]
    # No keywords from any slot signal
    assert find_table_for_slot(10, [table]) is None
    assert find_table_for_slot(14, [table]) is None
    assert find_table_for_slot(9, [table]) is None


def test_empty_table_list_returns_none():
    assert find_table_for_slot(10, []) is None


def test_label_row_slots_return_none():
    table = [["anything"]]
    for sid in (1, 2, 3, 4, 5, 6, 7, 8, 11, 12, 13, 15):
        assert find_table_for_slot(sid, [table]) is None


def test_table_preserves_all_rows_and_columns():
    table = [
        ["Tier", "Damage Level", "Award Amount"],
        ["1", "Minor", "100,000"],
        ["2", "Severe", "1,000,000"],
        ["3", "Complete Loss", "2,000,000"],
    ]
    result = find_table_for_slot(10, [table])
    assert result is not None
    assert len(result) == 4
    assert len(result[0]) == 3
    assert result[0][0] == "Tier"
    assert result[3][2] == "2,000,000"


def test_first_matching_table_returned():
    # Multiple tables - the first match wins.
    table_unrelated = [["header"], ["nothing"]]
    table_award = [
        ["Tier", "Award Amount"],
        ["Gold", "$1000"],
    ]
    result = find_table_for_slot(10, [table_unrelated, table_award])
    assert result is not None
    assert result[0][0] == "Tier"


# -- Context-aware table routing (issue 1-3) --

def test_table_routed_by_section_context_exclusion():
    """Regression: table under Exclusions should NOT be routed to slot 10.

    A facility/level table that appears in the Exclusions section
    must be routed to slot 9 (Exclusions), not slot 10 (Award Structure).
    """
    from policy_platform.rag.table_routing import find_table_for_slot_with_context

    paragraphs = [
        "Introduction: This policy supports safety.",
        "Policy Statement: Company shall provide assistance.",
        "Exclusions:",
        "Some exclusion text here.",
        "Level", "Facility", "Maintenance", "Priority",
        "Level 1", "Critical", "Monthly", "High",
    ]
    section_index = {0: 5, 1: 6, 2: 9, 3: 9}
    # Table_paragraph_indices: the table appears after paragraph 3.
    table_paragraph_indices = [4]
    facility_table = [
        ["Level", "Facility", "Maintenance", "Priority"],
        ["Level 1", "Critical Care", "Monthly", "High"],
    ]
    # Slot 10 should NOT get this table (it's in slot 9's section).
    result10 = find_table_for_slot_with_context(
        10, [facility_table], section_index, paragraphs,
        table_paragraph_indices=table_paragraph_indices,
    )
    assert result10 is None
    # Slot 9 SHOULD get this table.
    result9 = find_table_for_slot_with_context(
        9, [facility_table], section_index, paragraphs,
        table_paragraph_indices=table_paragraph_indices,
    )
    assert result9 is not None
    assert result9 == facility_table


def test_table_routed_to_award_structure_when_in_award_section():
    """A table under Award Structure should be routed to slot 10."""
    from policy_platform.rag.table_routing import find_table_for_slot_with_context

    paragraphs = [
        "Award Structure:",
        "Some intro text.",
        "Tier", "Amount",
        "Tier 1", "100",
    ]
    section_index = {0: 10, 1: 10}
    table_paragraph_indices = [2]
    award_table = [
        ["Tier", "Amount"],
        ["Tier 1", "100"],
        ["Tier 2", "200"],
    ]
    result = find_table_for_slot_with_context(
        10, [award_table], section_index, paragraphs,
        table_paragraph_indices=table_paragraph_indices,
    )
    assert result is not None
    assert result == award_table


def test_table_without_section_context_uses_content_signals():
    """Table with no section context falls back to content-signal matching."""
    from policy_platform.rag.table_routing import find_table_for_slot_with_context

    paragraphs = ["Some intro."]
    section_index = {}  # No section context
    award_table = [
        ["Tier", "Amount"],
        ["Tier 1", "100"],
    ]
    result = find_table_for_slot_with_context(
        10, [award_table], section_index, paragraphs,
    )
    assert result is not None
    assert result == award_table


def test_degenerate_table_rejected():
    # Empty table or all-empty cells
    table = [["", ""], ["", ""]]
    result = find_table_for_slot(10, [table])
    # All cells are blank - flatten returns empty, so no signal hits.
    # But the _table_shape_ok check needs at least one non-empty cell.
    assert result is None


def test_multi_table_routing_returns_all_in_section():
    """All tables in a slot's section are returned, not just the first.

    Phase K.6 update: with target_pi-first lookup, tpi=2 (Award
    Structure heading paragraph) directly maps to slot 10, so all
    tables with tpi=2 land in slot 10. The exclusions table (tpi=1)
    lands in slot 9.
    """
    from policy_platform.rag.table_routing import find_all_tables_for_slot_with_context
    section_index = {
        0: 5,  # Introduction
        1: 9,  # Exclusions
        2: 10,  # Award Structure
        3: 12,  # Definitions
    }
    paragraphs = ["Intro", "Exclusions", "Award Structure", "Definitions"]
    tables = [
        [["Category", "Excluded"], ["Contractors", "Yes"]],  # exclusions
        [["Tier", "Amount"], ["1", "100"], ["2", "200"]],  # award tier
        [["Level", "Priority"], ["Critical", "High"]],  # level
    ]
    # table_paragraph_indices: paragraph index of the section the table
    # belongs to. tpi=N means the table sits AFTER paragraph N, in N's
    # section.
    # - exclusions table: tpi=1 (after Exclusions heading paragraph)
    # - tier table: tpi=2 (after Award Structure heading paragraph)
    # - level table: tpi=2 (still in Award Structure section)
    table_paragraph_indices = [1, 2, 2]
    result_9 = find_all_tables_for_slot_with_context(
        9, tables, section_index, paragraphs, table_paragraph_indices
    )
    result_10 = find_all_tables_for_slot_with_context(
        10, tables, section_index, paragraphs, table_paragraph_indices
    )
    # Slot 9 should get 1 table (exclusions).
    assert len(result_9) == 1
    # Slot 10 should get 2 tables (tier + level).
    assert len(result_10) == 2


def test_table_paragraph_indices_off_by_one_fix():
    """Regression: table_paragraph_indices should point to LAST
    paragraph BEFORE the table, not after.

    Previously the index was `len(paragraphs)` (one past the last
    paragraph), which caused tables to be attributed to the next
    section's slot instead of their actual section.
    """
    from policy_platform.extractors import dispatch
    from docx import Document
    import tempfile
    d = Document()
    d.add_paragraph("Policy Title: Test")
    d.add_paragraph("Type: HR")
    d.add_paragraph("Effective Date: 01 July 2026.")
    d.add_paragraph("Applies to: All")
    d.add_paragraph("Reason for Policy: Test reason")
    d.add_paragraph("Introduction: Intro body.")
    d.add_paragraph("Policy Statement: Policy body.")
    d.add_paragraph("Purpose: Purpose body.")
    d.add_paragraph("Scope and Beneficiaries: All employees.")
    d.add_paragraph("Exclusions")
    t1 = d.add_table(rows=2, cols=2)
    t1.cell(0, 0).text = "Category"
    t1.cell(0, 1).text = "Excluded"
    t1.cell(1, 0).text = "Contractors"
    t1.cell(1, 1).text = "Yes"
    d.add_paragraph("Award Structure and Payout Tiers")
    t2 = d.add_table(rows=2, cols=2)
    t2.cell(0, 0).text = "Tier"
    t2.cell(0, 1).text = "Amount"
    t2.cell(1, 0).text = "1"
    t2.cell(1, 1).text = "100"
    d.add_paragraph("Definitions: Term definitions.")
    d.add_paragraph("RELATED POLICIES, PROCEDURES, FORMS, GUIDELINES & OTHER RESOURCES")
    d.add_paragraph("Reference.")
    d.add_paragraph("HISTORY")
    d.add_paragraph("Initial release.")

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        d.save(f.name)
        from pathlib import Path
        ext = dispatch(Path(f.name))
    # The exclusions table should be at the index of the Exclusions
    # heading paragraph, NOT one past it.
    assert ext.table_paragraph_indices[0] == 9  # Exclusions heading
    assert ext.table_paragraph_indices[1] == 10  # Award Structure heading
    # The exclusions table should be classifiable as slot 9.
    from policy_platform.rag.table_routing import _find_table_section_slot
    from policy_platform.rag.heading_anchors import build_section_index
    section_index = build_section_index(ext.paragraphs)
    slot = _find_table_section_slot(
        0, ext.tables, section_index, ext.table_paragraph_indices, ext.paragraphs
    )
    assert slot == 9, f"Expected slot 9 for exclusions table, got {slot}"
    slot = _find_table_section_slot(
        1, ext.tables, section_index, ext.table_paragraph_indices, ext.paragraphs
    )
    assert slot == 10, f"Expected slot 10 for award tier table, got {slot}"
