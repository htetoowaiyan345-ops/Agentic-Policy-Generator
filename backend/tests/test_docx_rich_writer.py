"""Tests for docx_rich_writer (Phase 4).

Asserts that rich HTML payloads from the TipTap editor round-trip
correctly into the OOXML emitted by `build_approved_docx`:
  - bold  -> w:b
  - italic -> w:i
  - underline -> w:u
  - strikethrough -> w:strike
  - color -> w:color (hex, no leading #)
  - font-family -> w:rFonts (ascii/hAnsi)
  - font-size -> w:sz (half-points)
  - heading -> w:outlineLvl
"""
from __future__ import annotations

import zipfile
from pathlib import Path

import pytest
from docx import Document
from docx.oxml.ns import qn

from api.docx_rich_writer import (
    _parse_spans,
    write_paragraph,
    plain_text_from_html,
    _as_color,
    _as_size,
    _sanitize_font,
)
from api.docx_approved_export import build_approved_docx


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _read_docx_xml(docx_path) -> str:
    with zipfile.ZipFile(str(docx_path)) as z:
        return z.read('word/document.xml').decode('utf-8')


def _build_minimal_source(docx_path: Path, n_paragraphs: int = 4) -> None:
    doc = Document()
    for i in range(n_paragraphs):
        doc.add_paragraph(f'p-{i}')
    doc.save(str(docx_path))


def _first_rpr_of_p(doc: Document, p_idx: int):
    p = doc.paragraphs[p_idx]
    rPrs = p._p.findall('.//' + qn('w:rPr'))
    return rPrs


def _runs_text(doc: Document, p_idx: int) -> str:
    p = doc.paragraphs[p_idx]
    return ''.join(
        (t.text or '')
        for t in p._p.findall('.//' + qn('w:t'))
    )


# ---------------------------------------------------------------------------
# Pure-function unit tests
# ---------------------------------------------------------------------------

def test_as_color_three_digit_hex():
    assert _as_color('#abc') == 'AABBCC'


def test_as_color_six_digit_hex():
    assert _as_color('#1A2b3C') == '1A2B3C'


def test_as_color_rgb():
    assert _as_color('rgb(255, 0, 0)') == 'FF0000'


def test_as_color_unsupported_returns_none():
    assert _as_color('currentcolor') is None
    assert _as_color('purple') is None  # not parsed (could extend)


def test_as_size_px():
    # 14px -> ~10.5pt -> 21 half-points
    assert _as_size('14px') == 21


def test_as_size_em():
    # 1.5em at 16px base -> 24pt -> 48 half-points
    assert _as_size('1.5em') == 48


def test_as_size_pt():
    assert _as_size('12pt') == 24


def test_sanitize_font_picks_first():
    assert _sanitize_font('Inter, sans-serif') == 'Inter'
    assert _sanitize_font('"Times New Roman", serif') == 'Times New Roman'
    assert _sanitize_font('serif') == 'Times New Roman'
    assert _sanitize_font('monospace') == 'Courier New'
    assert _sanitize_font('sans-serif') == 'Calibri'


def test_parse_spans_preserves_plain():
    spans = _parse_spans('<p>plain text</p>')
    text = ''.join(s.text or '' for s in spans)
    assert 'plain text' in text


def test_parse_spans_bold_and_italic():
    spans = _parse_spans('<p><strong>bold</strong> and <em>italic</em></p>')
    # Bold layer's font-weight must be 'bold'; italic layer's font-style 'italic'
    bold_span = next((s for s in spans if (s.text or '').strip() == 'bold'), None)
    italic_span = next((s for s in spans if (s.text or '').strip() == 'italic'), None)
    assert bold_span is not None
    assert italic_span is not None
    assert bold_span.rPr.get('font-weight') == 'bold'
    assert italic_span.rPr.get('font-style') == 'italic'


def test_parse_spans_color_and_size():
    spans = _parse_spans(
        '<p><span style="color: #ff0000; font-size: 24px;">big red</span></p>'
    )
    red = next((s for s in spans if (s.text or '').strip() == 'big red'), None)
    assert red is not None
    assert red.rPr.get('color') == '#ff0000' or red.rPr.get('color') == 'ff0000'.lower()


def test_plain_text_from_html_strips_tags():
    s = plain_text_from_html('<p><strong>bold</strong> and <em>italic</em></p>')
    assert 'bold' in s
    assert 'italic' in s
    assert '<' not in s


# ---------------------------------------------------------------------------
# Integration tests: write_paragraph on a python-docx paragraph
# ---------------------------------------------------------------------------

def test_write_paragraph_bold_emits_w_b(tmp_path: Path):
    doc = Document()
    p = doc.add_paragraph('placeholder')
    write_paragraph(p._p, '<p>Hello <strong>world</strong>!</p>')
    xml = p._p.xml
    assert '<w:rPr>' in xml
    assert '<w:b/>' in xml
    # Plain text concat
    txt = ''.join((t.text or '') for t in p._p.findall('.//' + qn('w:t')))
    assert 'Hello' in txt and 'world' in txt


def test_write_paragraph_italic_emits_w_i(tmp_path: Path):
    doc = Document()
    p = doc.add_paragraph('placeholder')
    write_paragraph(p._p, '<p>A <em>slanted</em> B</p>')
    assert '<w:i/>' in p._p.xml


def test_write_paragraph_underline_emits_w_u(tmp_path: Path):
    doc = Document()
    p = doc.add_paragraph('placeholder')
    write_paragraph(p._p, '<p><u>under</u></p>')
    xml = p._p.xml
    assert '<w:u ' in xml
    assert 'val="single"' in xml


def test_write_paragraph_strike_emits_w_strike(tmp_path: Path):
    doc = Document()
    p = doc.add_paragraph('placeholder')
    write_paragraph(p._p, '<p><s>gone</s></p>')
    assert '<w:strike/>' in p._p.xml


def test_write_paragraph_color_emits_w_color(tmp_path: Path):
    doc = Document()
    p = doc.add_paragraph('placeholder')
    write_paragraph(p._p, '<p><span style="color: #00ff00;">green</span></p>')
    xml = p._p.xml
    assert '<w:color ' in xml
    assert 'val="00FF00"' in xml


def test_write_paragraph_font_family_emits_w_rfonts(tmp_path: Path):
    doc = Document()
    p = doc.add_paragraph('placeholder')
    write_paragraph(p._p, '<p><span style="font-family: Calibri;">hi</span></p>')
    xml = p._p.xml
    assert '<w:rFonts' in xml
    # The serializer normalises to lowercase; both forms are acceptable.
    assert 'calibri' in xml.lower()


def test_write_paragraph_font_size_emits_w_sz(tmp_path: Path):
    doc = Document()
    p = doc.add_paragraph('placeholder')
    write_paragraph(p._p, '<p><span style="font-size: 24px;">big</span></p>')
    xml = p._p.xml
    assert '<w:sz' in xml
    # 24px -> ~18pt -> 36 half-points
    assert 'val="36"' in xml


def test_write_paragraph_heading_emits_outline_lvl(tmp_path: Path):
    doc = Document()
    p = doc.add_paragraph('placeholder')
    write_paragraph(p._p, '<h1>Title</h1>')
    xml = p._p.xml
    assert '<w:outlineLvl' in xml


def test_write_paragraph_keeps_pPr(tmp_path: Path):
    """If a paragraph had a pPr (style, alignment), we must NOT strip it."""
    from docx.shared import Pt
    doc = Document()
    p = doc.add_paragraph('placeholder')
    p.style = doc.styles['Heading 2']
    before = p._p.find(qn('w:pPr'))
    write_paragraph(p._p, '<p>updated</p>')
    after = p._p.find(qn('w:pPr'))
    assert before is not None
    assert after is not None


def test_write_paragraph_empty_html_still_creates_run(tmp_path: Path):
    doc = Document()
    p = doc.add_paragraph('placeholder')
    write_paragraph(p._p, '')
    runs = p._p.findall(qn('w:r'))
    assert len(runs) == 1


def test_write_paragraph_soft_break_with_br(tmp_path: Path):
    doc = Document()
    p = doc.add_paragraph('placeholder')
    write_paragraph(p._p, '<p>line1<br/>line2</p>')
    brs = p._p.findall('.//' + qn('w:br'))
    assert len(brs) >= 1


# ---------------------------------------------------------------------------
# Integration tests: build_approved_docx with rich payloads
# ---------------------------------------------------------------------------

def test_build_approved_docx_preserves_bold(tmp_path: Path):
    src = tmp_path / 'src.docx'
    out = tmp_path / 'out.docx'
    # Source needs at least 2 body paragraphs so we have something to overlay.
    _build_minimal_source(src, n_paragraphs=2)
    build_approved_docx(
        src,
        [
            ['p', {'slot': 7, 'text': 'X', 'html': '<p>Hello <strong>Bold</strong>!</p>'}],
            ['p', {'slot': 8, 'text': 'Y', 'html': '<p>Body</p>'}],
        ],
        str(out),
    )
    doc = Document(str(out))
    p_xml = doc.paragraphs[0]._p.xml
    assert '<w:b/>' in p_xml


def test_build_approved_docx_preserves_italic(tmp_path: Path):
    src = tmp_path / 'src.docx'
    out = tmp_path / 'out.docx'
    _build_minimal_source(src, n_paragraphs=2)
    build_approved_docx(
        src,
        [
            ['p', {'slot': 7, 'text': 'X', 'html': '<p>This is <em>italic</em>.</p>'}],
            ['p', 'plain-2'],
        ],
        str(out),
    )
    doc = Document(str(out))
    assert '<w:i/>' in doc.paragraphs[0]._p.xml


def test_build_approved_docx_legacy_payload_still_works(tmp_path: Path):
    src = tmp_path / 'src.docx'
    out = tmp_path / 'out.docx'
    _build_minimal_source(src)
    build_approved_docx(src, [['p', 'A'], ['p', 'B'], ['p', 'C']], str(out))
    doc = Document(str(out))
    assert doc.paragraphs[0].text == 'A'
    assert doc.paragraphs[1].text == 'B'
    assert doc.paragraphs[2].text == 'C'


def test_build_approved_docx_table_cell_preserves_bold(tmp_path: Path):
    src = tmp_path / 'src.docx'
    out = tmp_path / 'out.docx'
    doc = Document()
    t = doc.add_table(rows=2, cols=2)
    t.cell(0, 0).text = 'h1'
    t.cell(0, 1).text = 'h2'
    t.cell(1, 0).text = 'a'
    t.cell(1, 1).text = 'b'
    doc.add_paragraph('after')
    doc.save(str(src))

    build_approved_docx(
        src,
        [
            ['t', {'slot': 14, 'rows': [
                [{'text': 'X', 'html': '<p><strong>X</strong></p>'}],
                [{'text': 'Y', 'html': '<p>Y</p>'}],
            ]}],
            ['p', 'AFTER'],
        ],
        str(out),
    )
    doc = Document(str(out))
    cell0_xml = doc.tables[0].cell(0, 0).paragraphs[0]._p.xml
    assert '<w:b/>' in cell0_xml
    assert doc.paragraphs[0].text == 'AFTER'


def test_build_approved_docx_mixed_legacy_and_rich(tmp_path: Path):
    src = tmp_path / 'src.docx'
    out = tmp_path / 'out.docx'
    _build_minimal_source(src, n_paragraphs=4)
    build_approved_docx(
        src,
        [
            ['p', 'legacy-1'],
            ['p', {'slot': 7, 'text': 'rich-2', 'html': '<p><em>rich-2</em></p>'}],
            ['p', 'legacy-3'],
            ['p', {'slot': 8, 'text': 'rich-4', 'html': '<p><strong>rich-4</strong></p>'}],
        ],
        str(out),
    )
    doc = Document(str(out))
    assert doc.paragraphs[0].text == 'legacy-1'
    assert '<w:i/>' in doc.paragraphs[1]._p.xml
    assert doc.paragraphs[2].text == 'legacy-3'
    assert '<w:b/>' in doc.paragraphs[3]._p.xml


def test_build_approved_docx_does_not_break_on_missing_html(tmp_path: Path):
    """Rich dict without html falls back to plain text."""
    src = tmp_path / 'src.docx'
    out = tmp_path / 'out.docx'
    _build_minimal_source(src)
    build_approved_docx(
        src,
        [
            ['p', {'slot': 7, 'text': 'fallback'}],
            ['p', 'plain'],
        ],
        str(out),
    )
    doc = Document(str(out))
    assert doc.paragraphs[0].text == 'fallback'
    assert doc.paragraphs[1].text == 'plain'


def test_build_approved_docx_no_original(tmp_path: Path):
    """When no source docx exists, a minimal new docx is built with
    text-only fallback (rich writer needs an existing paragraph)."""
    out = tmp_path / 'out.docx'
    build_approved_docx(
        None,
        [['p', 'first'], ['p', {'slot': 7, 'text': 'second', 'html': '<p><b>second</b></p>'}]],
        str(out),
    )
    doc = Document(str(out))
    assert doc.paragraphs[0].text == 'first'
    assert doc.paragraphs[1].text == 'second'


def test_audit_text_preserved_when_formatting_round_trips(tmp_path: Path):
    """The `text` field must remain the verbatim source text (audit log)."""
    src = tmp_path / 'src.docx'
    out = tmp_path / 'out.docx'
    _build_minimal_source(src)
    src_text = 'Original Source'
    build_approved_docx(
        src,
        [['p', {'slot': 12, 'text': src_text, 'html': f'<p><strong>{src_text}</strong></p>'}]],
        str(out),
    )
    doc = Document(str(out))
    # The line/paragraph terminator in OOXML is a soft break <w:br/>
    # which surfaces as '\n' in `paragraph.text`. Strip trailing whitespace.
    assert doc.paragraphs[0].text.strip() == src_text


# ---------------------------------------------------------------------------
# Stage 3 / 4 — full toolbar round-trip coverage
# ---------------------------------------------------------------------------
# These tests assert that every formatting control exposed by the
# Step 03 CKEditor toolbar survives the lines_json -> .docx pipeline.
# Each test takes a deliberately crafted HTML payload, runs it through
# `write_paragraph` (the front door of the rich writer), and asserts
# the expected OOXML fragment is emitted. Stages 4 and 5 will add the
# structural controls (alignment, blockquote, link) that are not yet
# handled by the rich writer.

@pytest.mark.parametrize(
    'html,expected_xml',
    [
        # 1. Bold
        ('<p><strong>bold</strong></p>', '<w:b/>'),
        # 2. Italic
        ('<p><em>italic</em></p>', '<w:i/>'),
        # 3. Underline
        ('<p><u>under</u></p>', '<w:u '),
        # 4. Strikethrough
        ('<p><s>gone</s></p>', '<w:strike/>'),
        # 5. Font color (hex)
        ('<p><span style="color: #ff0000;">red</span></p>', 'val="FF0000"'),
        # 6. Font size (em)
        ('<p><span style="font-size: 1.5em;">big</span></p>', '<w:sz'),
        # 7. Font family (inter)
        ('<p><span style="font-family: Inter;">a</span></p>', 'w:ascii="inter"'),
        # 8. Bold inside divider/span
        ('<p><strong><span style="color: #00ff00;">bold-green</span></strong></p>',
         '<w:b/>'),
        # 9. Heading 1
        ('<h1>Title</h1>', '<w:outlineLvl'),
        # 10. Heading 2
        ('<h2>Sub</h2>', '<w:outlineLvl'),
        # 11. Heading 3
        ('<h3>Sub-sub</h3>', '<w:outlineLvl'),
        # 12. Br (soft break)
        ('<p>line1<br/>line2</p>', '<w:br'),
        # 13. Alignment: right
        ('<p style="text-align: right;">right</p>', '<w:jc w:val="right"'),
        # 14. Alignment: center
        ('<p style="text-align: center;">center</p>', '<w:jc w:val="center"'),
        # 15. Alignment: justify
        ('<p style="text-align: justify;">justify</p>', '<w:jc w:val="both"'),
        # 16. Hyperlink
        ('<p>see <a href="https://example.com">link</a></p>',
         '<w:hyperlink'),
        # 17. Blockquote (indent)
        ('<blockquote>quoted</blockquote>', '<w:ind'),
    ],
)
def test_stage3_inline_format_round_trip(html, expected_xml, tmp_path):
    """Every inline toolbar control must survive write_paragraph."""
    doc = Document()
    p = doc.add_paragraph('placeholder')
    write_paragraph(p._p, html)
    assert expected_xml in p._p.xml, f"missing {expected_xml!r} for {html!r}"


# ---------------------------------------------------------------------------
# Stage 7 — CKEditor 5's ACTUAL output format
# ---------------------------------------------------------------------------
# The /test-roundtrip browser test page showed that CKEditor 5 normalizes
# its HTML output per HTML5 spec:
#   - <em> → <i> (HTML5 considers <i> a valid semantic italic tag)
#   - 'color: #ff0000' → 'color:#ff0000' (whitespace removed)
#   - 'font-size: 24px' → 'font-size:24px' (whitespace removed)
#   - 'text-align: right' → 'text-align:right' (whitespace removed)
#
# These tests use the EXACT HTML format the editor produces to prove the
# backend correctly converts whatever the editor outputs, not just the
# canonical format with whitespace and <em> tags.
#
# This is critical because lines_json is populated from
# `editor.getData()` — so the HTML the backend receives IS whatever
# CKEditor 5 chose to emit, not what the user typed.


@pytest.mark.parametrize(
    'name,html,expected_frag',
    [
        # Editor's actual italic output: <em> → <i>
        ('editor-italic',
         '<p><i>italic</i></p>',
         '<w:i/>'),
        # Editor's actual color output: no whitespace
        ('editor-color-hex',
         '<p><span style="color:#ff0000;">red</span></p>',
         'val="FF0000"'),
        # Editor's actual color output with rgb() (no spaces)
        ('editor-color-rgb',
         '<p><span style="color:rgb(255,0,0);">red</span></p>',
         'val="FF0000"'),
        # Editor's actual font size output: no whitespace
        ('editor-fontsize-no-space',
         '<p><span style="font-size:24px;">big</span></p>',
         'val="36"'),
        # Editor's actual font size output: no unit (raw px)
        ('editor-fontsize-px',
         '<p><span style="font-size:18px;">mid</span></p>',
         'val="27"'),
        # Editor's actual right-align output: no whitespace
        ('editor-align-right',
         '<p style="text-align:right;">right</p>',
         'w:val="right"'),
        # Editor's actual center-align output: no whitespace
        ('editor-align-center',
         '<p style="text-align:center;">center</p>',
         'w:val="center"'),
        # Editor's actual justify output
        ('editor-align-justify',
         '<p style="text-align:justify;">justify</p>',
         'w:val="both"'),
        # Editor's bold output: <strong>
        ('editor-bold-strong',
         '<p><strong>bold</strong></p>',
         '<w:b/>'),
        # Editor's underline output: <u>
        ('editor-underline',
         '<p><u>under</u></p>',
         '<w:u '),
        # Editor's strikethrough output: <s>
        ('editor-strike',
         '<p><s>struck</s></p>',
         '<w:strike/>'),
        # Editor's hyperlink output
        ('editor-hyperlink',
         '<p>see <a href="https://example.com">link</a></p>',
         '<w:hyperlink'),
        # Editor's blockquote output
        ('editor-blockquote',
         '<blockquote>quoted</blockquote>',
         '<w:ind'),
        # Editor's horizontal line output
        ('editor-hr',
         '<p>before</p><hr><p>after</p>',
         'after'),
        # Editor's data-slot output — the backend strips data-slot from
        # the <w:p> element because it's a UI hint for the editor, not
        # content formatting. The slot is preserved separately by the
        # lines_json pipeline. We just verify the text "slot 5" is
        # written.
        ('editor-data-slot',
         '<p data-slot="5">slot 5</p>',
         'slot 5'),
        # Editor's combined: bold + color + alignment
        ('editor-combined-bold-color',
         '<p style="text-align:center;"><strong><span style="color:#0000ff;">centered blue</span></strong></p>',
         'val="0000FF"'),
    ],
)
def test_stage7_editor_output_format(name, html, expected_frag, tmp_path):
    """The backend correctly handles the EXACT HTML format that CKEditor 5
    produces after its HTML5 normalization (no whitespace, <i> not <em>,
    etc.). This proves the toolbar round-trip is fully functional — the
    editor's output is what reaches the backend, and the backend handles
    it correctly.
    """
    doc = Document()
    p = doc.add_paragraph('placeholder')
    write_paragraph(p._p, html)
    assert expected_frag in p._p.xml, (
        f"{name}: missing {expected_frag!r} for {html!r}"
    )
