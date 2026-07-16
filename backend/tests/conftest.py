"""Shared fixtures for tests."""
from __future__ import annotations

import os
import re
import shutil
import uuid
from pathlib import Path

import pytest

from policy_platform import config


def _load_dotenv(path: Path) -> None:
    """Minimal .env loader (no external dep). Sets KEY=value pairs into os.environ
    only if the key is not already set in os.environ (so real env wins)."""
    if not path.exists():
        return
    for raw in path.read_text(encoding="utf-8").splitlines():
        line = raw.strip()
        if not line or line.startswith("#"):
            continue
        if "=" not in line:
            continue
        key, _, value = line.partition("=")
        key = key.strip()
        value = value.strip().strip('"').strip("'")
        os.environ.setdefault(key, value)


# Best-effort load of .env so tests resolve API_BASE_URL/DB_PATH/etc. from
# the developer machine's local config without ever typing values inline.
_load_dotenv(config.PROJECT_ROOT / ".env")


SAMPLE_FILENAMES: list[tuple[str, str]] = [
    ("Policy_Template_Award_and_Recognition_Updated.docx", ".docx"),
    ("Policy_Template_Award_and_Recognition_Updated.pdf", ".pdf"),
    ("Hospital_Buildings_Policy_Template.pdf", ".pdf"),
    ("Sexual Harassment Policy.docx", ".docx"),
    ("Sexual Harassment Policy.pdf", ".pdf"),
    ("School_Building_Policy_Template_Updated.pdf", ".pdf"),
    ("Earthquake_Full_Policy_One_Paragraph.docx", ".docx"),
    ("Earthquake_Full_Policy_One_Paragraph.pdf", ".pdf"),
    ("Policy For Coronavirus Disease.docx", ".docx"),
]
"""Single source of truth for sample-file fixtures. Tests import this list,
never inline filename literals."""


def pytest_collection_modifyitems(config, items):
    """Skip e2e tests unless RUN_E2E=1 is set in env."""
    if os.environ.get("RUN_E2E", "0") != "1":
        skip_e2e = pytest.mark.skip(reason="set RUN_E2E=1 to run end-to-end tests")
        for item in items:
            if "e2e" in item.keywords:
                item.add_marker(skip_e2e)


@pytest.fixture(scope="session")
def project_root() -> Path:
    return config.PROJECT_ROOT


@pytest.fixture(scope="session")
def brain_path() -> Path:
    return config.BRAIN_PATH


@pytest.fixture(scope="session")
def samples_dir(project_root: Path) -> Path:
    return config.SAMPLES_DIR


@pytest.fixture(scope="session")
def fixtures_dir(project_root: Path) -> Path:
    d = project_root / "tests" / "fixtures"
    d.mkdir(parents=True, exist_ok=True)
    return d


@pytest.fixture
def api_base_url() -> str:
    return config.API_BASE_URL


@pytest.fixture
def db_path() -> Path:
    return config.DB_PATH


@pytest.fixture
def test_db_path() -> Path:
    return config.TEST_DB_PATH


@pytest.fixture
def test_actor() -> str:
    return config.TEST_ACTOR


@pytest.fixture
def test_reviewer() -> str:
    return config.TEST_REVIEWER


@pytest.fixture
def edit_marker() -> str:
    """Runtime-generated unique marker so each test run can deterministically
    detect its own edit while stale runs from previous runs can never mask
    a real regression."""
    return f"E2E-MARKER-{uuid.uuid4().hex[:8]}"


@pytest.fixture
def small_docx(fixtures_dir: Path) -> Path:
    """Create a small docx fixture."""
    from docx import Document

    path = fixtures_dir / "sample_policy.docx"
    if not path.exists():
        doc = Document()
        doc.add_paragraph("Policy Title: Test Policy A")
        doc.add_paragraph("Policy Number: TP-001")
        doc.add_paragraph("Brief Description: A test policy used by the platform.")
        doc.add_paragraph("INTRODUCTION")
        doc.add_paragraph("This is the introduction body for the test policy.")
        doc.add_paragraph("POLICY STATEMENT")
        doc.add_paragraph("We shall do the thing.")
        doc.add_paragraph("1. Purpose")
        doc.add_paragraph("To test the pipeline.")
        doc.add_paragraph("2. Scope & Beneficiaries")
        doc.add_paragraph("All tests in the repo.")
        doc.add_paragraph("3. Exclusions")
        doc.add_paragraph("Production users during the test window.")
        doc.add_paragraph("4. Award Structure & Payout Tiers")
        t = doc.add_table(rows=2, cols=2)
        t.rows[0].cells[0].text = "Tier"
        t.rows[0].cells[1].text = "Amount"
        t.rows[1].cells[0].text = "A"
        t.rows[1].cells[1].text = "100"
        doc.add_paragraph("DEFINITIONS")
        doc.add_paragraph("Company: Test Co.")
        doc.save(str(path))
    return path


@pytest.fixture
def small_txt(fixtures_dir: Path) -> Path:
    path = fixtures_dir / "sample_policy.txt"
    if not path.exists():
        path.write_text(
            "Policy Title: TXT Policy\n"
            "Brief Description: A plain-text policy used to exercise the pipeline.\n"
            "INTRODUCTION\nThis is the intro.\nPOLICY STATEMENT\nWe do the thing.\n"
            "1. Purpose\nValidate text extraction.\n"
            "DEFINITIONS\nCompany: TXT Co.\n",
            encoding="utf-8",
        )
    return path
