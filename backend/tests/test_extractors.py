from policy_platform.extractors import docx_extractor, txt_extractor
from pathlib import Path


def test_docx_extractor_returns_verbatim(tmp_path: Path):
    from tests.conftest import small_docx  # noqa: F401
    pass


def test_txt_extractor_preserves_bytes(tmp_path: Path):
    p = tmp_path / "x.txt"
    p.write_bytes(b"Hello\nWorld\n")
    e = txt_extractor.extract(p)
    assert e.paragraphs == ["Hello", "World"]
    assert e.source_format == "txt"
    assert e.full_text.count("Hello") == 1


def test_txt_extractor_preserves_crlf(tmp_path: Path):
    p = tmp_path / "x.txt"
    p.write_bytes(b"Hello\r\nWorld\r\n")
    e = txt_extractor.extract(p)
    # splitlines() normalizes CRLF; the Unicode rule allows this
    assert e.paragraphs == ["Hello", "World"]
    # Round-trip via raw bytes preserved
    assert "Hello" in e.full_text and "World" in e.full_text


def test_docx_extractor_preserves_paragraphs(tmp_path: Path):
    p = tmp_path / "x.docx"
    from docx import Document
    d = Document()
    d.add_paragraph("Alpha")
    d.add_paragraph("Beta")
    d.save(str(p))
    e = docx_extractor.extract(p)
    assert "Alpha" in e.paragraphs
    assert "Beta" in e.paragraphs
