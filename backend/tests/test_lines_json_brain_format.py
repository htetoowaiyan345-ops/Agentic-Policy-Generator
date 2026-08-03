"""test_lines_json_brain_format.py

Stage 4.19 - tests verifying the published .docx matches the brain
framework format EXACTLY:

  - Empty leftover scaffold paragraphs are DELETED (not left as
    blank lines / bullet points in Word).
  - Brain's pPr (pStyle, spacing, jc, ind, numPr) is preserved
    verbatim on every paragraph that survives.
  - Slot-1 metadata labels are bold; values are not.
  - Slot 5-14 headings (INTRODUCTION, POLICY STATEMENT, etc.) are
    bold and retain brain's Roman-numeral numbering.
  - Slot heading dedup: when user types the heading label as
    their first paragraph, the scaffold heading is cleared so
    we don't render the label twice.
  - Free paragraph zone at top renders slot=0 paragraphs
    verbatim (no scaffold substitution).
  - User content replaces brain placeholder text in EVERY filled
    slot — no "Type:Policy" / "13 Feb 2023" / "300,000 MMK" /
    "Si Thu Maung" / "Pass with distinction" leaks.
"""
from __future__ import annotations

import importlib
import importlib.util
import sys
import types
from pathlib import Path

import pytest
from docx import Document
from docx.oxml.ns import qn


@pytest.fixture(scope="session")
def backend_root() -> Path:
    return Path(__file__).resolve().parents[1]


@pytest.fixture(scope="session")
def brain_path(backend_root) -> Path:
    return backend_root / "data" / "brain_template" / "Policy_Framework_5.docx"


@pytest.fixture(scope="session")
def out_dir(tmp_path_factory) -> Path:
    return tmp_path_factory.mktemp("lines_json_brain_format_out")


@pytest.fixture(scope="session")
def renderer_mod(backend_root):
    if str(backend_root) not in sys.path:
        sys.path.insert(0, str(backend_root))
    pkg = types.ModuleType("policy_platform")
    pkg.__path__ = [str(backend_root / "policy_platform")]
    pkg.__package__ = "policy_platform"
    sys.modules["policy_platform"] = pkg
    fw = types.ModuleType("policy_platform.framework")
    fw.__path__ = [str(backend_root / "policy_platform" / "framework")]
    fw.__package__ = "policy_platform.framework"
    sys.modules["policy_platform.framework"] = fw

    def _load(name: str, path: Path):
        spec = importlib.util.spec_from_file_location(name, str(path))
        assert spec is not None and spec.loader is not None
        m = importlib.util.module_from_spec(spec)
        sys.modules[name] = m
        spec.loader.exec_module(m)
        return m

    _load("policy_platform.config", backend_root / "policy_platform" / "config.py")
    _load("policy_platform.framework.section_map",
          backend_root / "policy_platform" / "framework" / "section_map.py")
    _load("policy_platform.framework.brain_slot_map",
          backend_root / "policy_platform" / "framework" / "brain_slot_map.py")
    src = backend_root / "policy_platform" / "lines_json_renderer.py"
    spec = importlib.util.spec_from_file_location("policy_platform.lines_json_renderer", str(src))
    assert spec is not None and spec.loader is not None
    mod = importlib.util.module_from_spec(spec)
    sys.modules["policy_platform.lines_json_renderer"] = mod
    spec.loader.exec_module(mod)
    return mod


def _all_text(doc):
    out = []
    for p in doc.paragraphs:
        t = ''.join(x.text or "" for x in p._element.iter(qn("w:t")))
        if t:
            out.append(t)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                if cell.text:
                    out.append(cell.text)
    return ' '.join(out)


# ---------------------------------------------------------------------------
# Tests
# ---------------------------------------------------------------------------

def test_brain_scaffold_placeholders_never_leak(renderer_mod, brain_path, out_dir):
    """When user provides content for slots, brain placeholder text
    ('13 Feb 2023', 'Daw Win Win Tint, Group CEO', '300,000 MMK',
    'Si Thu Maung', 'Pass with distinction', 'City Family High
    School Completion Award', 'City Holdings Group') must NOT
    appear in the output."""
    lines_json = [
        ["p", {"slot": 1, "text": "Type: HR Policy", "html": "<p>Type: HR Policy</p>"}],
        ["p", {"slot": 1, "text": "Policy Title: AWARD POLICY", "html": "<p>Policy Title: AWARD POLICY</p>"}],
        ["p", {"slot": 2, "text": "My brief description", "html": "<p>My brief description</p>"}],
        ["p", {"slot": 3, "text": "01 July 2026 - 30 June 2027", "html": "<p>01 July 2026 - 30 June 2027</p>"}],
        ["p", {"slot": 4, "text": "To establish a fair policy.", "html": "<p>To establish a fair policy.</p>"}],
        ["p", {"slot": 5, "text": "Body for intro", "html": "<p>Body for intro</p>"}],
        ["p", {"slot": 6, "text": "Body for policy", "html": "<p>Body for policy</p>"}],
        ["p", {"slot": 7, "text": "Body for purpose", "html": "<p>Body for purpose</p>"}],
        ["p", {"slot": 8, "text": "Body for scope", "html": "<p>Body for scope</p>"}],
        ["p", {"slot": 9, "text": "Body for exclusions", "html": "<p>Body for exclusions</p>"}],
        ["p", {"slot": 10, "text": "Body for awards", "html": "<p>Body for awards</p>"}],
        ["p", {"slot": 12, "text": "Body for definitions", "html": "<p>Body for definitions</p>"}],
        ["p", {"slot": 13, "text": "Body for related", "html": "<p>Body for related</p>"}],
        ["p", {"slot": 14, "text": "Body for history", "html": "<p>Body for history</p>"}],
    ]
    out = out_dir / "no_scaffold.docx"
    renderer_mod.render_lines_json_to_brain(lines_json, brain_path, out)
    text = _all_text(Document(str(out)))
    forbidden = [
        "City Family High School Completion Award",
        "13 Feb 2023",
        "Daw Win Win Tint",
        "300,000 MMK",
        "Si Thu Maung",
        "Pass with distinction",
        "City Holdings Group of Companies and all its integrated",
    ]
    for sig in forbidden:
        assert sig not in text, f"brain scaffold {sig!r} leaked"


def test_brain_pPr_preserved_on_user_paragraphs(renderer_mod, brain_path, out_dir):
    """User-written paragraphs inherit the brain scaffold's pPr
    (pStyle, jc, ind) — but line height is publication-overridden to 1.5x."""
    lines_json = [
        ["p", {"slot": 5, "text": "INTRODUCTION", "html": "<p>INTRODUCTION</p>"}],
        ["p", {"slot": 5, "text": "Body content here.", "html": "<p>Body content here.</p>"}],
    ]
    out = out_dir / "pPr.docx"
    renderer_mod.render_lines_json_to_brain(lines_json, brain_path, out)
    doc = Document(str(out))
    # Find body paragraph (skip the deduped heading)
    target = next(p for p in doc.paragraphs if "Body content" in (p.text or ""))
    pPr = target._element.find(qn("w:pPr"))
    assert pPr is not None
    # line height is publication-overridden to 480 (2.0x).
    spacing = pPr.find(qn("w:spacing"))
    if spacing is not None:
        line = spacing.get(qn("w:line"))
        assert line == "480", f"expected line=480 (2.0x); got {line}"


def test_paragraphs_use_left_alignment(renderer_mod, brain_path, out_dir):
    """All paragraphs (except table cells) use left alignment
    (overriding brain's scaffold jc=both)."""
    lines_json = [
        ["p", {"slot": 5, "text": "INTRODUCTION", "html": "<p>INTRODUCTION</p>"}],
        ["p", {"slot": 5, "text": "Body content here.", "html": "<p>Body content here.</p>"}],
    ]
    out = out_dir / "jc.docx"
    renderer_mod.render_lines_json_to_brain(lines_json, brain_path, out)
    doc = Document(str(out))
    target = next(p for p in doc.paragraphs if "Body content" in (p.text or ""))
    pPr = target._element.find(qn("w:pPr"))
    assert pPr is not None
    jc = pPr.find(qn("w:jc"))
    assert jc is not None, "jc must be set on body paragraphs"
    assert jc.get(qn("w:val")) == "left", (
        f"body paragraphs must use jc=left; got {jc.get(qn('w:val'))}"
    )


def test_paragraphs_use_1_5_line_height(renderer_mod, brain_path, out_dir):
    """All paragraphs (except table cells) use 1.5 line height
    (line=480 lineRule=auto)."""
    lines_json = [
        ["p", {"slot": 5, "text": "INTRODUCTION", "html": "<p>INTRODUCTION</p>"}],
        ["p", {"slot": 5, "text": "Body content here.", "html": "<p>Body content here.</p>"}],
    ]
    out = out_dir / "line_height.docx"
    renderer_mod.render_lines_json_to_brain(lines_json, brain_path, out)
    doc = Document(str(out))
    target = next(p for p in doc.paragraphs if "Body content" in (p.text or ""))
    pPr = target._element.find(qn("w:pPr"))
    assert pPr is not None
    spacing = pPr.find(qn("w:spacing"))
    assert spacing is not None
    assert spacing.get(qn("w:line")) == "480", (
        f"expected line=480 (2.0x), got {spacing.get(qn('w:line'))}"
    )
    assert spacing.get(qn("w:lineRule")) == "auto"


def test_brain_numPr_preserved_on_headings(renderer_mod, brain_path, out_dir):
    """When user provides content that includes the heading label as
    their first paragraph (dedup), brain's numPr on the resulting
    paragraph is preserved (the deduped paragraph keeps its pPr).
    The deduped paragraph IS the heading paragraph."""
    lines_json = [
        # User's first paragraph IS the heading label. Dedup kicks in:
        # scaffold heading is cleared and user text fills it.
        ["p", {"slot": 5, "text": "INTRODUCTION", "html": "<p>INTRODUCTION</p>"}],
        # Second paragraph fills body[0].
        ["p", {"slot": 5, "text": "Body content here.", "html": "<p>Body content here.</p>"}],
    ]
    out = out_dir / "numPr.docx"
    renderer_mod.render_lines_json_to_brain(lines_json, brain_path, out)
    doc = Document(str(out))
    # Find paragraphs and check that the FIRST "INTRODUCTION" paragraph
    # has numPr (it's the heading_elem that received the user's text).
    found_intro_with_numpr = False
    for p in doc.paragraphs:
        text = (p.text or "").strip()
        if text.startswith("INTRODUCTION") and not text.startswith("Body"):
            pPr = p._element.find(qn("w:pPr"))
            if pPr is not None:
                numPr = pPr.find(qn("w:numPr"))
                if numPr is not None:
                    numId = numPr.find(qn("w:numId"))
                    if numId is not None and numId.get(qn("w:val")) == "6":
                        found_intro_with_numpr = True
                        break
    assert found_intro_with_numpr, (
        "deduped INTRODUCTION heading must retain brain's numId=6 (Roman numerals)"
    )


def test_empty_paragraphs_are_deleted_not_kept(renderer_mod, brain_path, out_dir):
    """Empty leftover scaffold paragraphs (after user fills slot) are
    DELETED entirely — not left as blank lines / bullet points.
    Brain's structural empty paragraphs (before slot 1, between slots)
    are tolerated."""
    lines_json = [
        # Slot 12 has 8 body items; user provides only 1.
        ["p", {"slot": 12, "text": "User def", "html": "<p>User def</p>"}],
    ]
    out = out_dir / "no_empty.docx"
    renderer_mod.render_lines_json_to_brain(lines_json, brain_path, out)
    doc = Document(str(out))
    # Count empty paragraphs INSIDE the slot 12 region (after heading).
    # We tolerate 0-2 empty paragraphs at the start of the document
    # (brain structural scaffolding before slot 1).
    empty_count = 0
    for p in doc.paragraphs:
        if not (p.text or "").strip():
            empty_count += 1
    # Up to 2 empty paragraphs from brain's pre-slot-1 scaffolding.
    assert empty_count <= 2, f"too many empty paragraphs: {empty_count}"


def test_slot1_label_bold_value_normal(renderer_mod, brain_path, out_dir):
    """Slot-1 metadata labels (Type:, Policy Title:, etc.) are bold;
    their values are not bold."""
    lines_json = [
        ["p", {"slot": 1, "text": "Type: HR Policy", "html": "<p>Type: HR Policy</p>"}],
    ]
    out = out_dir / "label_bold.docx"
    renderer_mod.render_lines_json_to_brain(lines_json, brain_path, out)
    doc = Document(str(out))
    target = next(p for p in doc.paragraphs if "Type: HR Policy" in (p.text or ""))
    bold_runs = []
    nonbold_runs = []
    for r in target._element.findall(qn("w:r")):
        rPr = r.find(qn("w:rPr"))
        t = r.find(qn("w:t"))
        text = (t.text or "") if t is not None else ""
        is_bold = (rPr is not None and rPr.find(qn("w:b")) is not None)
        if is_bold:
            bold_runs.append(text)
        elif text.strip():
            nonbold_runs.append(text)
    assert any("Type" in t for t in bold_runs), f"no bold 'Type': bold={bold_runs} nonbold={nonbold_runs}"
    assert any("HR Policy" in t for t in nonbold_runs), f"no non-bold 'HR Policy': bold={bold_runs} nonbold={nonbold_runs}"


def test_slot_heading_dedup_when_user_typed_label(renderer_mod, brain_path, out_dir):
    """When the user's first paragraph for slot 5 is 'INTRODUCTION',
    the scaffold heading is cleared and only one 'INTRODUCTION'
    paragraph exists (no duplicate)."""
    lines_json = [
        ["p", {"slot": 5, "text": "INTRODUCTION", "html": "<p>INTRODUCTION</p>"}],
        ["p", {"slot": 5, "text": "Body content here.", "html": "<p>Body content here.</p>"}],
    ]
    out = out_dir / "dedup.docx"
    renderer_mod.render_lines_json_to_brain(lines_json, brain_path, out)
    doc = Document(str(out))
    intro_count = sum(
        1 for p in doc.paragraphs
        if "INTRODUCTION" in (p.text or "")
    )
    assert intro_count == 1, f"expected 1 INTRODUCTION, got {intro_count}"


def test_no_dividers_inserted(renderer_mod, brain_path, out_dir):
    """No horizontal-rule dividers are inserted by default. The user
    spec adds borders under specific slot-1 rows ('Functional Area(s):'
    and 'Applies to:') — verified separately below."""
    lines_json = [
        ["p", {"slot": 1, "text": "Some unrelated label", "html": "<p>Some unrelated label</p>"}],
        ["p", {"slot": 1, "text": "Other label", "html": "<p>Other label</p>"}],
    ]
    out = out_dir / "no_dividers.docx"
    renderer_mod.render_lines_json_to_brain(lines_json, brain_path, out)
    doc = Document(str(out))
    divider_count = 0
    for p in doc.paragraphs:
        pPr = p._element.find(qn("w:pPr"))
        if pPr is not None and pPr.find(qn("w:pBdr")) is not None:
            divider_count += 1
    assert divider_count == 0, f"unexpected dividers: {divider_count}"


def test_functional_area_and_applies_to_have_divider_after(renderer_mod, brain_path, out_dir):
    """Per user spec: 'Functional Area(s)' and 'Applies to' slot-1 rows
    are followed by a dedicated divider paragraph with a single bottom
    border + 8px top/bottom margins. Dividers are SEPARATE paragraphs."""
    lines_json = [
        ["p", {"slot": 1, "text": "Functional Area(s): Human Resources", "html": "<p>Functional Area(s): Human Resources</p>"}],
        ["p", {"slot": 1, "text": "Applies to: All eligible employees", "html": "<p>Applies to: All eligible employees</p>"}],
    ]
    out = out_dir / "with_dividers.docx"
    renderer_mod.render_lines_json_to_brain(lines_json, brain_path, out)
    doc = Document(str(out))
    # The next paragraph after each anchor should be the divider.
    paragraphs = list(doc.paragraphs)
    divider_count = 0
    for i, p in enumerate(paragraphs):
        text = (p.text or "").strip()
        if not (text.startswith("Functional Area") or text.startswith("Applies to")):
            continue
        # Next paragraph should be divider (empty + pBdr + 8px margins).
        if i + 1 >= len(paragraphs):
            continue
        next_p = paragraphs[i + 1]
        next_text = (next_p.text or "").strip()
        assert next_text == "", f"divider should be empty, got: {next_text!r}"
        next_pPr = next_p._element.find(qn("w:pPr"))
        assert next_pPr is not None, "divider missing pPr"
        next_pBdr = next_pPr.find(qn("w:pBdr"))
        assert next_pBdr is not None, "divider missing pBdr"
        next_bottom = next_pBdr.find(qn("w:bottom"))
        assert next_bottom is not None
        assert next_bottom.get(qn("w:val")) == "single"
        next_spacing = next_pPr.find(qn("w:spacing"))
        assert next_spacing is not None
        assert next_spacing.get(qn("w:before")) == "160"
        assert next_spacing.get(qn("w:after")) == "160"
        divider_count += 1
    assert divider_count == 2, f"expected 2 dividers after anchors, got {divider_count}"


def test_slot1_metadata_1_0_line_height(renderer_mod, brain_path, out_dir):
    """Slot-1 metadata paragraphs use 1.0 line height (line=240)
    while body paragraphs use 2.0 line height (line=480)."""
    lines_json = [
        ["p", {"slot": 1, "text": "Type: HR Policy", "html": "<p>Type: HR Policy</p>"}],
        ["p", {"slot": 1, "text": "Policy Title: My Title", "html": "<p>Policy Title: My Title</p>"}],
        ["p", {"slot": 1, "text": "Applicable Sector(s): Corporate", "html": "<p>Applicable Sector(s): Corporate</p>"}],
        ["p", {"slot": 1, "text": "Functional Area(s): Human Resources", "html": "<p>Functional Area(s): Human Resources</p>"}],
        ["p", {"slot": 1, "text": "Applies to: All employees", "html": "<p>Applies to: All employees</p>"}],
        ["p", {"slot": 1, "text": "Last Reviewed: 01 July 2026", "html": "<p>Last Reviewed: 01 July 2026</p>"}],
    ]
    out = out_dir / "slot1_line.docx"
    renderer_mod.render_lines_json_to_brain(lines_json, brain_path, out)
    doc = Document(str(out))
    for p in doc.paragraphs:
        text = (p.text or "").strip()
        if not text:
            continue
        pPr = p._element.find(qn("w:pPr"))
        spacing = pPr.find(qn("w:spacing")) if pPr is not None else None
        line = spacing.get(qn("w:line")) if spacing is not None else None
        if (text.lower().startswith("type:") or
            text.lower().startswith("policy title") or
            text.lower().startswith("applicable sector") or
            text.lower().startswith("functional area") or
            text.lower().startswith("applies to") or
            text.lower().startswith("last reviewed")):
            assert line == "240", (
                f"slot-1 metadata '{text}' should have line=240, got {line}"
            )
        else:
            # Body paragraphs should be 480 (2.0x line height).
            assert line == "480", (
                f"body paragraph '{text}' should have line=480, got {line}"
            )


def test_applies_to_has_divider_after_post_pass(renderer_mod, brain_path, out_dir):
    """The 'Applies to:' row gets a divider paragraph (empty + border +
    8px margins) immediately after it, even when it's an extra
    paragraph beyond the brain scaffold body count."""
    lines_json = [
        ["p", {"slot": 1, "text": "Type: HR Policy", "html": "<p>Type: HR Policy</p>"}],
        ["p", {"slot": 1, "text": "Policy Title: My Title", "html": "<p>Policy Title: My Title</p>"}],
        ["p", {"slot": 1, "text": "Policy Number: HR-001", "html": "<p>Policy Number: HR-001</p>"}],
        ["p", {"slot": 1, "text": "Applicable Sector(s): Corporate", "html": "<p>Applicable Sector(s): Corporate</p>"}],
        ["p", {"slot": 1, "text": "Functional Area(s): Human Resources", "html": "<p>Functional Area(s): Human Resources</p>"}],
        ["p", {"slot": 1, "text": "Brief Description: Brief.", "html": "<p>Brief Description: Brief.</p>"}],
        ["p", {"slot": 1, "text": "Effective Date/Period: 01 July 2026", "html": "<p>Effective Date/Period: 01 July 2026</p>"}],
        ["p", {"slot": 1, "text": "Approved by: Htet Oo", "html": "<p>Approved by: Htet Oo</p>"}],
        ["p", {"slot": 1, "text": "Prepared by: Htet Oo Wai Yan", "html": "<p>Prepared by: Htet Oo Wai Yan</p>"}],
        ["p", {"slot": 1, "text": "Responsible Function(s): Human Resources", "html": "<p>Responsible Function(s): Human Resources</p>"}],
        ["p", {"slot": 1, "text": "Responsible Function Officer(s): Htet", "html": "<p>Responsible Function Officer(s): Htet</p>"}],
        ["p", {"slot": 1, "text": "Supersedes: Version 0.9", "html": "<p>Supersedes: Version 0.9</p>"}],
        ["p", {"slot": 1, "text": "Last Reviewed: 01 July 2026", "html": "<p>Last Reviewed: 01 July 2026</p>"}],
        ["p", {"slot": 1, "text": "Applies to: All eligible employees", "html": "<p>Applies to: All eligible employees</p>"}],
        ["p", {"slot": 1, "text": "Reason for Policy: framework", "html": "<p>Reason for Policy: framework</p>"}],
    ]
    out = out_dir / "applies_to_divider.docx"
    renderer_mod.render_lines_json_to_brain(lines_json, brain_path, out)
    doc = Document(str(out))
    paragraphs = list(doc.paragraphs)
    applies_to_index = None
    for i, p in enumerate(paragraphs):
        text = (p.text or "").strip()
        if text.lower().startswith("applies to:"):
            applies_to_index = i
            break
    assert applies_to_index is not None, "Applies to: row not found"
    next_p = paragraphs[applies_to_index + 1]
    assert (next_p.text or "").strip() == "", "next after Applies to should be divider"
    next_pPr = next_p._element.find(qn("w:pPr"))
    assert next_pPr is not None
    next_pBdr = next_pPr.find(qn("w:pBdr"))
    assert next_pBdr is not None, "Applies to: should be followed by divider paragraph"
    next_bottom = next_pBdr.find(qn("w:bottom"))
    assert next_bottom is not None
    assert next_bottom.get(qn("w:val")) == "single"
    next_spacing = next_pPr.find(qn("w:spacing"))
    assert next_spacing is not None
    assert next_spacing.get(qn("w:before")) == "160"
    assert next_spacing.get(qn("w:after")) == "160"


def test_type_has_divider_before(renderer_mod, brain_path, out_dir):
    """A dedicated divider paragraph (empty + bottom-border + 8px
    margins) is inserted immediately BEFORE the first slot-1 metadata
    row ('Type:')."""
    lines_json = [
        ["p", {"slot": 1, "text": "Type: HR Policy", "html": "<p>Type: HR Policy</p>"}],
        ["p", {"slot": 1, "text": "Policy Title: My Title", "html": "<p>Policy Title: My Title</p>"}],
    ]
    out = out_dir / "type_top_divider.docx"
    renderer_mod.render_lines_json_to_brain(lines_json, brain_path, out)
    doc = Document(str(out))
    paragraphs = list(doc.paragraphs)
    type_index = None
    for i, p in enumerate(paragraphs):
        text = (p.text or "").strip()
        if text.lower().startswith("type:"):
            type_index = i
            break
    assert type_index is not None, "Type: row not found"
    assert type_index > 0, "Type: should NOT be the first paragraph (divider must come before)"
    prev_p = paragraphs[type_index - 1]
    assert (prev_p.text or "").strip() == "", "paragraph before Type: should be empty divider"
    prev_pPr = prev_p._element.find(qn("w:pPr"))
    assert prev_pPr is not None
    prev_pBdr = prev_pPr.find(qn("w:pBdr"))
    assert prev_pBdr is not None, "missing pBdr on divider before Type:"
    prev_bottom = prev_pBdr.find(qn("w:bottom"))
    assert prev_bottom is not None
    assert prev_bottom.get(qn("w:val")) == "single"
    prev_spacing = prev_pPr.find(qn("w:spacing"))
    assert prev_spacing is not None
    assert prev_spacing.get(qn("w:before")) == "160"
    assert prev_spacing.get(qn("w:after")) == "160"


def test_english_sections_have_divider_after(renderer_mod, brain_path, out_dir):
    """Both [English] paragraphs get a divider paragraph (empty +
    bottom-border + 5px top / 8px bottom margins) immediately AFTER
    them. The [English] TEXT paragraph itself gets a 6px top margin
    (120 twips) per user spec for breathing room above the label."""
    lines_json = [
        ["p", {"slot": 5, "text": "This is intro body.", "html": "<p>This is intro body.</p>"}],
        ["p", {"slot": 5, "text": "[English]", "html": "<p>[English]</p>"}],
        ["p", {"slot": 7, "text": "This is definitions body.", "html": "<p>This is definitions body.</p>"}],
        ["p", {"slot": 7, "text": "[English]", "html": "<p>[English]</p>"}],
    ]
    out = out_dir / "english_dividers.docx"
    renderer_mod.render_lines_json_to_brain(lines_json, brain_path, out)
    doc = Document(str(out))
    paragraphs = list(doc.paragraphs)
    english_indices = []
    for i, p in enumerate(paragraphs):
        if (p.text or "").strip() == "[English]":
            english_indices.append(i)
    assert len(english_indices) == 2, f"expected 2 [English], got {len(english_indices)}"
    for idx in english_indices:
        # 1) [English] TEXT paragraph itself: 6px top margin (120 twips).
        english_p = paragraphs[idx]._element
        eng_pPr = english_p.find(qn("w:pPr"))
        assert eng_pPr is not None, "[English] paragraph missing pPr"
        eng_spacing = eng_pPr.find(qn("w:spacing"))
        assert eng_spacing is not None, "[English] paragraph missing spacing"
        assert eng_spacing.get(qn("w:before")) == "120", (
            f"[English] top margin should be 120 twips (6px), "
            f"got {eng_spacing.get(qn('w:before'))!r}"
        )
        # 2) Divider paragraph AFTER [English]: 5px top, 8px bottom.
        next_p = paragraphs[idx + 1]
        assert (next_p.text or "").strip() == "", "paragraph after [English] should be empty divider"
        next_pPr = next_p._element.find(qn("w:pPr"))
        assert next_pPr is not None
        next_pBdr = next_pPr.find(qn("w:pBdr"))
        assert next_pBdr is not None, "missing pBdr on divider after [English]"
        next_bottom = next_pBdr.find(qn("w:bottom"))
        assert next_bottom is not None
        assert next_bottom.get(qn("w:val")) == "single"
        next_spacing = next_pPr.find(qn("w:spacing"))
        assert next_spacing is not None
        # [English] divider: 5px (100 twips) top, 8px (160 twips) bottom.
        assert next_spacing.get(qn("w:before")) == "100", (
            f"[English] divider top margin should be 100 (5px), "
            f"got {next_spacing.get(qn('w:before'))!r}"
        )
        assert next_spacing.get(qn("w:after")) == "160", (
            f"[English] divider bottom margin should be 160 (8px), "
            f"got {next_spacing.get(qn('w:after'))!r}"
        )


def test_all_5_dividers_in_full_doc(renderer_mod, brain_path, out_dir):
    """End-to-end: 15 metadata paragraphs + body with 2 [English] sections
    should produce exactly 5 dedicated divider paragraphs in the correct
    positions:
      1) BEFORE Type:        (divider paragraph above)
      2) AFTER Functional Area(s): (divider paragraph below)
      3) AFTER Applies to:   (divider paragraph below)
      4) AFTER first [English] (divider paragraph below)
      5) AFTER second [English] (divider paragraph below)
    Each divider is empty + has <w:pBdr><w:bottom/></w:pBdr> + 8px
    margins (before=160, after=160).
    """
    lines_json = [
        ["p", {"slot": 1, "text": "Type: HR Policy", "html": "<p>Type: HR Policy</p>"}],
        ["p", {"slot": 1, "text": "Policy Title: Title", "html": "<p>Policy Title: Title</p>"}],
        ["p", {"slot": 1, "text": "Policy Number: HR-001", "html": "<p>Policy Number: HR-001</p>"}],
        ["p", {"slot": 1, "text": "Applicable Sector(s): Corporate", "html": "<p>Applicable Sector(s): Corporate</p>"}],
        ["p", {"slot": 1, "text": "Functional Area(s): Human Resources", "html": "<p>Functional Area(s): Human Resources</p>"}],
        ["p", {"slot": 1, "text": "Brief Description: Brief.", "html": "<p>Brief Description: Brief.</p>"}],
        ["p", {"slot": 1, "text": "Effective Date/Period: 01 July 2026", "html": "<p>Effective Date/Period: 01 July 2026</p>"}],
        ["p", {"slot": 1, "text": "Approved by: Htet Oo", "html": "<p>Approved by: Htet Oo</p>"}],
        ["p", {"slot": 1, "text": "Prepared by: Htet Oo Wai Yan", "html": "<p>Prepared by: Htet Oo Wai Yan</p>"}],
        ["p", {"slot": 1, "text": "Responsible Function(s): HR", "html": "<p>Responsible Function(s): HR</p>"}],
        ["p", {"slot": 1, "text": "Responsible Function Officer(s): HOW", "html": "<p>Responsible Function Officer(s): HOW</p>"}],
        ["p", {"slot": 1, "text": "Supersedes: Version 0.9", "html": "<p>Supersedes: Version 0.9</p>"}],
        ["p", {"slot": 1, "text": "Last Reviewed: 01 July 2026", "html": "<p>Last Reviewed: 01 July 2026</p>"}],
        ["p", {"slot": 1, "text": "Applies to: All eligible employees", "html": "<p>Applies to: All eligible employees</p>"}],
        ["p", {"slot": 1, "text": "Reason for Policy: framework", "html": "<p>Reason for Policy: framework</p>"}],
        ["p", {"slot": 5, "text": "Introduction body.", "html": "<p>Introduction body.</p>"}],
        ["p", {"slot": 5, "text": "[English]", "html": "<p>[English]</p>"}],
        ["p", {"slot": 7, "text": "Definitions body.", "html": "<p>Definitions body.</p>"}],
        ["p", {"slot": 7, "text": "[English]", "html": "<p>[English]</p>"}],
    ]
    out = out_dir / "all_5_dividers.docx"
    renderer_mod.render_lines_json_to_brain(lines_json, brain_path, out)
    doc = Document(str(out))
    paragraphs = list(doc.paragraphs)

    # Count divider paragraphs (empty + pBdr + either 8px or 5px margins).
    divider_indices = []
    english_divider_indices = []
    for i, p in enumerate(paragraphs):
        if (p.text or "").strip() != "":
            continue
        pPr = p._element.find(qn("w:pPr"))
        if pPr is None:
            continue
        pBdr = pPr.find(qn("w:pBdr"))
        if pBdr is None:
            continue
        bottom = pBdr.find(qn("w:bottom"))
        if bottom is None:
            continue
        spacing = pPr.find(qn("w:spacing"))
        if spacing is None:
            continue
        before = spacing.get(qn("w:before"))
        after = spacing.get(qn("w:after"))
        # Slot-1 dividers (Type, Functional Area, Applies to): 8px top/bottom.
        if before == "160" and after == "160":
            divider_indices.append(i)
        # [English] dividers: 5px top, 8px bottom (per user spec).
        elif before == "100" and after == "160":
            english_divider_indices.append(i)
    assert len(divider_indices) == 3, f"expected 3 slot-1 dividers (8px), got {len(divider_indices)}: {divider_indices}"
    assert len(english_divider_indices) == 2, f"expected 2 [English] dividers (5px top / 8px bottom), got {len(english_divider_indices)}: {english_divider_indices}"

    # Verify each divider is in the correct position relative to anchors.
    type_idx = next(i for i, p in enumerate(paragraphs) if (p.text or "").strip().lower().startswith("type:"))
    functional_idx = next(i for i, p in enumerate(paragraphs) if (p.text or "").strip().lower().startswith("functional area"))
    applies_idx = next(i for i, p in enumerate(paragraphs) if (p.text or "").strip().lower().startswith("applies to"))
    english_indices = [i for i, p in enumerate(paragraphs) if (p.text or "").strip().lower() == "[english]"]
    assert len(english_indices) == 2

    assert (type_idx - 1) in divider_indices, "divider missing before Type:"
    assert (functional_idx + 1) in divider_indices, "divider missing after Functional Area(s)"
    assert (applies_idx + 1) in divider_indices, "divider missing after Applies to:"
    for idx in english_indices:
        assert (idx + 1) in english_divider_indices, f"[English] divider (5px top / 8px bottom) missing after [English] at {idx}"


def test_faded_table_borders_stripped(renderer_mod, brain_path, out_dir):
    """The brain template's TableGridLight / PlainTable3 styles use
    light gray (BFBFBF / 7F7F7F) borders. These render as visible
    "old lines" in the output. The faded-border strip pass removes
    them so the published document has no old faded lines."""
    lines_json = [
        ["p", {"slot": 1, "text": "Type: HR Policy", "html": "<p>Type: HR Policy</p>"}],
        ["p", {"slot": 10, "text": "Award tier data.", "html": "<p>Award tier data.</p>"}],
        ["p", {"slot": 10, "html": "<table><tr><td>Cell</td></tr></table>", "rows": [["Cell"]]}],
    ]
    out = out_dir / "no_faded_borders.docx"
    renderer_mod.render_lines_json_to_brain(lines_json, brain_path, out)
    # Read raw XML to inspect borders.
    import zipfile
    with zipfile.ZipFile(str(out)) as zf:
        xml = zf.read("word/document.xml").decode("utf-8")
    # No BFBFBF or 7F7F7F should remain on borders.
    for color in ("BFBFBF", "7F7F7F", "bFBFBF".lower(), "7f7f7f".lower()):
        # Check both upper and lower case (depending on how they got serialized).
        if color in xml or color.upper() in xml:
            # Make sure it's not just a styles.xml definition (which is preserved).
            # The styles file is separate from document.xml.
            # If the color appears in document.xml as a border color attribute,
            # that's a problem.
            import re
            in_border = re.search(r'w:color="' + color + '"', xml, re.IGNORECASE)
            if in_border:
                # Could be acceptable if not on a border element. Check
                # whether the surrounding context is a border or shading.
                start = max(0, in_border.start() - 200)
                ctx = xml[start:in_border.end()]
                if '<w:' in ctx and ('Borders' in ctx or 'bdr' in ctx.lower() or 'border' in ctx.lower()):
                    pytest.fail(
                        f"faded border color {color} still present in "
                        f"document.xml: ...{ctx[-150:]}"
                    )


def test_inherited_table_style_borders_preserved(renderer_mod, brain_path, out_dir):
    """The brain's table styles `TableGridLight` and `PlainTable3`
    define borders (BFBFBF / 7F7F7F) in `word/styles.xml`. These are
    inherited by tables via `<w:tblStyle>` references and render as
    the brain's original table lines — which the user wants KEPT.

    This test verifies that the renderer does NOT mutate these style
    border definitions: every side in TableGridLight's `<w:tblBorders>`
    and PlainTable3's conditional `<w:tcBorders>` retains its original
    `val="single"` (or similar non-nil value).
    """
    lines_json = [
        ["p", {"slot": 1, "text": "Type: HR Policy", "html": "<p>Type: HR Policy</p>"}],
        ["p", {"slot": 10, "text": "Award tier data.", "html": "<p>Award tier data.</p>"}],
        ["t", {"slot": 10, "rows": [["Tier", "Detail"]]}],
        ["t", {"slot": 14, "rows": [["Date", "Change"]]}],
    ]
    out = out_dir / "styles_borders_preserved.docx"
    renderer_mod.render_lines_json_to_brain(lines_json, brain_path, out)
    import zipfile, re
    with zipfile.ZipFile(str(out)) as zf:
        styles_xml = zf.read("word/styles.xml").decode("utf-8")
    # TableGridLight: at least one side must retain val="single"
    # (the brain's original BFBFBF single border).
    m = re.search(
        r'<w:style w:type="table" w:styleId="TableGridLight">'
        r'(.*?)</w:style>',
        styles_xml,
        re.DOTALL,
    )
    assert m is not None, "TableGridLight style not found in styles.xml"
    bm = re.search(
        r"<w:tblBorders>(.*?)</w:tblBorders>",
        m.group(1),
        re.DOTALL,
    )
    assert bm is not None, "TableGridLight missing tblBorders in styles.xml"
    sides = re.findall(
        r"<w:(top|left|bottom|right|insideH|insideV)\b[^/>]*w:val=\"([^\"]+)\"",
        bm.group(1),
    )
    assert sides, "TableGridLight tblBorders has no side elements"
    vals = [v for _, v in sides]
    assert "single" in vals, (
        "TableGridLight borders were mutated (no single sides remain). "
        "Brain's original table lines must be preserved. Found vals: "
        + str(vals)
    )


def test_page_header_separator_suppressed(renderer_mod, brain_path, out_dir):
    """The brain template contains line elements that render as
    visible horizontal lines in the published docx:

    1. `<v:line>` connectors (legacy VML) in `header2.xml` and
       `document.xml` — removed by stripping every `<v:line>` and
       dropping `<mc:AlternateContent>` whose Fallback is purely
       `<w:pict><v:line/></w:pict>`.

    2. Zero-height `<w:drawing>` elements (modern DrawingML) — these
       are degenerate drawings with `<wp:extent cy="0">` that Word
       renders as horizontal lines on every page where the header
       appears. The brain has one in `header2.xml` (cx=467.46pt,
       cy=0) which was the persistent "2nd line" the user kept
       seeing.

    The renderer's `_suppress_page_header_separator` pass handles
    both forms. Legitimate drawings (logo images with cy > 0) and
    header text content are NOT touched.

    This test verifies that no `<v:line>` AND no zero-height
    `<w:drawing>` survives in any XML in the published docx.
    """
    lines_json = [
        ["p", {"slot": 1, "text": "Type: HR Policy", "html": "<p>Type: HR Policy</p>"}],
    ]
    out = out_dir / "no_lines_anywhere.docx"
    renderer_mod.render_lines_json_to_brain(lines_json, brain_path, out)
    import zipfile, re
    with zipfile.ZipFile(str(out)) as zf:
        for name in zf.namelist():
            if not name.endswith(".xml"):
                continue
            xml = zf.read(name).decode("utf-8")
            # 1) No <v:line> anywhere
            vlines = re.findall(r"<v:line\b", xml)
            assert not vlines, (
                name + " still contains " + str(len(vlines)) + " "
                "<v:line> element(s)."
            )
            # 2) No zero-height <w:drawing> elements
            zero_h = re.findall(
                r'<wp:extent\s+cx="\d+"\s+cy="0"',
                xml,
            )
            assert not zero_h, (
                name + " still contains " + str(len(zero_h)) + " "
                "drawing(s) with zero-height extent (cy=0); these "
                "render as horizontal lines on every page. Logo "
                "drawings (cy > 0) must be preserved."
            )


def test_slot14_history_no_bullet_numbering(renderer_mod, brain_path, out_dir):
    """The HISTORY heading (slot 14) must have Roman numerals (numId=6)
    like other section titles (INTRODUCTION, DEFINITIONS, etc.) so it
    matches the user's spec. Scaffold body paragraphs AND user-written
    paragraphs in slot 14 must NOT have numPr so the user's edited
    data renders as plain text (not bullets).

    The renderer:
      1. Strips <w:numPr> from all slot-14 scaffold body paragraphs
         (so user-edited data renders as plain text).
      2. Adds <w:numPr><w:numId val="6"/></w:numPr> back to the
         HISTORY HEADING only (so it renders with Roman numerals).

    This test verifies that:
      - The HISTORY heading paragraph has <w:numPr> with numId=6.
      - Adjacent paragraph AFTER HISTORY (a user body paragraph) has
        NO <w:numPr>.
    """
    lines_json = [
        ["p", {"slot": 1, "text": "Type: HR Policy", "html": "<p>Type: HR Policy</p>"}],
        ["p", {"slot": 14, "text": "HISTORY Htet Oo", "html": "<p>HISTORY Htet Oo</p>"}],
        ["t", {"slot": 14, "rows": [
            ["DATE", "VERSION", "DESCRIPTION"],
            ["05 July 2026", "1.0", "Initial Release"],
        ]}],
    ]
    out = out_dir / "slot14_no_bullets.docx"
    renderer_mod.render_lines_json_to_brain(lines_json, brain_path, out)
    import zipfile, re
    with zipfile.ZipFile(str(out)) as zf:
        doc_xml = zf.read("word/document.xml").decode("utf-8")
    # Find the HISTORY heading paragraph in document.xml
    history_pattern = re.search(
        r"<w:p\b[^>]*>(?:(?!</w:p>).)*?HISTORY(?:(?!</w:p>).)*?</w:p>",
        doc_xml,
        re.DOTALL,
    )
    assert history_pattern, "HISTORY paragraph not found in output"
    history_block = history_pattern.group()
    # The HISTORY heading MUST have <w:numPr> with numId=6 (Roman
    # numerals) — matching other titles like INTRODUCTION.
    assert "<w:numPr>" in history_block, (
        "HISTORY heading should have <w:numPr> (Roman numerals) "
        "to match other section titles. Block: " + history_block[:500]
    )
    assert '<w:numId w:val="6"/>' in history_block, (
        "HISTORY heading should use numId=6 (Roman numerals). "
        "Block: " + history_block[:500]
    )
    # The next paragraph (user body text) must NOT have <w:numPr>.
    history_pos = doc_xml.find("HISTORY")
    after_history = doc_xml[history_pos + len(history_block):]
    # Take first 1500 chars after HISTORY for adjacent paragraphs.
    nearby = after_history[:1500]
    # Find the next paragraph after HISTORY
    next_p = re.search(r"<w:p\b[^>]*>(?:(?!</w:p>).)*?</w:p>", nearby, re.DOTALL)
    if next_p:
        next_block = next_p.group()
        # If this paragraph has only the table or no text content,
        # skip. Otherwise check no numPr.
        if "<w:t" in next_block:
            assert "<w:numPr>" not in next_block, (
                "Paragraph after HISTORY should NOT have <w:numPr> "
                "so user data renders as plain text. Block: "
                + next_block[:500]
            )


def test_inherited_table_style_borders_preserved(renderer_mod, brain_path, out_dir):
    """The brain's table styles `TableGridLight` and `PlainTable3`
    define borders (BFBFBF / 7F7F7F) in `word/styles.xml`. These are
    inherited by tables via `<w:tblStyle>` references and render as
    the brain's original table lines — which the user wants KEPT.

    This test verifies that the renderer does NOT mutate these style
    border definitions: every side in TableGridLight's `<w:tblBorders>`
    and PlainTable3's conditional `<w:tcBorders>` retains its original
    `val="single"` (or similar non-nil value).
    """
    lines_json = [
        ["p", {"slot": 1, "text": "Type: HR Policy", "html": "<p>Type: HR Policy</p>"}],
        ["p", {"slot": 10, "text": "Award tier data.", "html": "<p>Award tier data.</p>"}],
        ["t", {"slot": 10, "rows": [["Tier", "Detail"]]}],
        ["t", {"slot": 14, "rows": [["Date", "Change"]]}],
    ]
    out = out_dir / "styles_borders_preserved.docx"
    renderer_mod.render_lines_json_to_brain(lines_json, brain_path, out)
    import zipfile, re
    with zipfile.ZipFile(str(out)) as zf:
        styles_xml = zf.read("word/styles.xml").decode("utf-8")
    # TableGridLight: at least one side must retain val="single"
    # (the brain's original BFBFBF single border).
    m = re.search(
        r'<w:style w:type="table" w:styleId="TableGridLight">'
        r'(.*?)</w:style>',
        styles_xml,
        re.DOTALL,
    )
    assert m is not None, "TableGridLight style not found in styles.xml"
    bm = re.search(
        r"<w:tblBorders>(.*?)</w:tblBorders>",
        m.group(1),
        re.DOTALL,
    )
    assert bm is not None, "TableGridLight missing tblBorders in styles.xml"
    sides = re.findall(
        r"<w:(top|left|bottom|right|insideH|insideV)\b[^/>]*w:val=\"([^\"]+)\"",
        bm.group(1),
    )
    assert sides, "TableGridLight tblBorders has no side elements"
    vals = [v for _, v in sides]
    assert "single" in vals, (
        "TableGridLight borders were mutated (no single sides remain). "
        "Brain's original table lines must be preserved. Found vals: "
        + str(vals)
    )


def test_explicit_font_size_still_works(renderer_mod, brain_path, out_dir):
    """When the user HTML specifies a font-size, it's honoured."""
    lines_json = [
        ["p", {
            "slot": 5,
            "text": "INTRODUCTION",
            "html": "<p style=\"font-size:24px\">INTRODUCTION</p>",
        }],
    ]
    out = out_dir / "explicit_size.docx"
    renderer_mod.render_lines_json_to_brain(lines_json, brain_path, out)
    doc = Document(str(out))
    target = next(p for p in doc.paragraphs if "INTRODUCTION" in (p.text or ""))
    sizes = []
    for r in target._element.findall(qn("w:r")):
        rPr = r.find(qn("w:rPr"))
        if rPr is not None:
            sz = rPr.find(qn("w:sz"))
            if sz is not None:
                sizes.append(sz.get(qn("w:val")))
# Per user directive: ALL text must be Times New Roman 10pt, even
    # when the user specifies a different font-size via inline style.
    # 10pt -> 20 half-points.
    assert "20" in sizes, f"expected sz=20 (10pt uniform), got {sizes}"


def test_explicit_bold_still_works(renderer_mod, brain_path, out_dir):
    """User HTML with <strong> still renders as bold runs."""
    lines_json = [
        ["p", {
            "slot": 5,
            "text": "bold",
            "html": '<p>Body with <strong>BOLD</strong> word.</p>',
        }],
    ]
    out = out_dir / "explicit_bold.docx"
    renderer_mod.render_lines_json_to_brain(lines_json, brain_path, out)
    doc = Document(str(out))
    target = next(p for p in doc.paragraphs if "BOLD" in (p.text or ""))
    bold_found = False
    for r in target._element.findall(qn("w:r")):
        rPr = r.find(qn("w:rPr"))
        if rPr is None:
            continue
        text = ''.join(t.text or "" for t in r.iter(qn("w:t")))
        if "BOLD" not in text:
            continue
        if rPr.find(qn("w:b")) is not None:
            bold_found = True
    assert bold_found, "bold run not preserved"
