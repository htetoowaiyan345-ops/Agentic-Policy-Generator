"""Tests for Pyidaungsu font embedding in the DOCX.

Verifies that when ``data/fonts/Pyidaungsu.ttf`` is bundled:
  1. The renderer embeds the font inside the DOCX zip.
  2. ``word/fontTable.xml`` registers the font with the right name.
  3. ``[Content_Types].xml`` has the font content type override.
  4. The rFonts element on Burmese runs sets all four attributes
     (ascii, hAnsi, eastAsia, cs) to "Pyidaungsu".

Tests skip gracefully when the Pyidaungsu.ttf file is not bundled
(this is the case for fresh checkouts before the user provides
the font). This matches the rest of the i18n test pattern.
"""
from __future__ import annotations

import io
import sys
import tempfile
import unittest
import zipfile
from pathlib import Path

from policy_platform.style import (
    BURMESE_FONT_NAME,
    get_burmese_font_path,
    apply_myanmar_font_if_needed,
)
from policy_platform.renderer import _embed_burmese_font_in_docx


def _make_minimal_docx(path: Path) -> None:
    """Create a tiny DOCX with Burmese text."""
    from docx import Document
    doc = Document()
    p = doc.add_paragraph("Hello world: မိတ်ဆက် ဤမူဝါဒသည်။")
    doc.save(str(path))


class TestPyidaungsuBundled(unittest.TestCase):
    """Verify the Pyidaungsu TTF is present and loadable."""

    def test_pyidaungsu_ttf_bundled(self):
        path = get_burmese_font_path()
        if path is None:
            self.skipTest("Pyidaungsu.ttf not bundled (drop at data/fonts/)")
        self.assertTrue(Path(path).exists())
        self.assertGreater(Path(path).stat().st_size, 100_000)

    def test_font_name_is_pyidaungsu(self):
        self.assertEqual(BURMESE_FONT_NAME, "Pyidaungsu")

    def test_bold_font_bundled(self):
        """Pyidaungsu-Bold.ttf should also be bundled for bold runs."""
        regular = get_burmese_font_path()
        if regular is None:
            self.skipTest("Pyidaungsu.ttf not bundled")
        bold = Path(regular).with_name("Pyidaungsu-Bold.ttf")
        if not bold.exists():
            self.skipTest("Pyidaungsu-Bold.ttf not bundled")
        self.assertGreater(bold.stat().st_size, 100_000)


class TestFontEmbedding(unittest.TestCase):
    """Verify the embedding helper writes a valid embedded-font DOCX."""

    def setUp(self):
        self.regular = get_burmese_font_path()
        if self.regular is None:
            self.skipTest("Pyidaungsu.ttf not bundled")
        self.bold = Path(self.regular).with_name("Pyidaungsu-Bold.ttf")
        self.tmpdir = tempfile.mkdtemp()
        self.docx_path = Path(self.tmpdir) / "test.docx"

    def tearDown(self):
        import shutil
        shutil.rmtree(self.tmpdir, ignore_errors=True)

    def _build_and_embed(self):
        _make_minimal_docx(self.docx_path)
        ok = _embed_burmese_font_in_docx(
            self.docx_path, self.regular,
            str(self.bold) if self.bold.exists() else None,
        )
        return ok

    def test_embedding_returns_true(self):
        self.assertTrue(self._build_and_embed())

    def test_docx_zip_contains_font(self):
        self._build_and_embed()
        with zipfile.ZipFile(str(self.docx_path)) as z:
            names = z.namelist()
            self.assertIn("word/fonts/pyidaungsu-regular.ttf", names)

    def test_docx_zip_contains_bold_when_present(self):
        if not self.bold.exists():
            self.skipTest("Pyidaungsu-Bold.ttf not bundled")
        self._build_and_embed()
        with zipfile.ZipFile(str(self.docx_path)) as z:
            self.assertIn("word/fonts/pyidaungsu-bold.ttf", z.namelist())

    def test_font_table_xml_registers_pyidaungsu(self):
        self._build_and_embed()
        with zipfile.ZipFile(str(self.docx_path)) as z:
            ft = z.read("word/fontTable.xml").decode("utf-8")
        self.assertIn('w:name="Pyidaungsu"', ft)
        self.assertIn("w:embedRegular", ft)
        if self.bold.exists():
            self.assertIn("w:embedBold", ft)

    def test_relationships_reference_font(self):
        self._build_and_embed()
        with zipfile.ZipFile(str(self.docx_path)) as z:
            rels = z.read("word/_rels/document.xml.rels").decode("utf-8")
        # The relationship Type for font embedding.
        self.assertIn(
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/font",
            rels,
        )
        # The relationship target points to the embedded font.
        self.assertIn("fonts/pyidaungsu-regular.ttf", rels)

    def test_content_types_include_font(self):
        self._build_and_embed()
        with zipfile.ZipFile(str(self.docx_path)) as z:
            ct = z.read("[Content_Types].xml").decode("utf-8")
        self.assertIn(
            "PartName=\"/word/fonts/pyidaungsu-regular.ttf\"", ct
        )
        self.assertIn("ContentType=\"application/x-font-ttf\"", ct)


class TestMyanmarFontRPr(unittest.TestCase):
    """Verify the rFonts element on Burmese runs sets all four attributes."""

    def test_apply_myanmar_font_sets_all_four_attributes(self):
        from docx import Document
        from docx.oxml.ns import qn
        doc = Document()
        p = doc.add_paragraph("Burmese: မိတ်ဆက်")
        apply_myanmar_font_if_needed(p._p)
        # The paragraph should now have at least one run with rFonts.
        rPr = p._p.find(qn("w:r")).find(qn("w:rPr"))
        rFonts = rPr.find(qn("w:rFonts"))
        self.assertIsNotNone(rFonts)
        self.assertEqual(rFonts.get(qn("w:ascii")), "Pyidaungsu")
        self.assertEqual(rFonts.get(qn("w:hAnsi")), "Pyidaungsu")
        self.assertEqual(rFonts.get(qn("w:eastAsia")), "Pyidaungsu")
        self.assertEqual(rFonts.get(qn("w:cs")), "Pyidaungsu")


if __name__ == "__main__":
    unittest.main()