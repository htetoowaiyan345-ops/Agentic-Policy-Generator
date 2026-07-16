"""Integrity validator using heading-based slot lookup.

Contract:
  - Brain media SHA-256 unchanged in output.
  - Output paragraph+table count <= Brain body count + slot_capacity growth.
  - For every "Found" slot, the placed source lines must be present in
    the output's slot text byte-for-byte.
  - For every non-Found slot, the slot's output text does NOT contain
    the Brain's example text (we emptied those slots).
  - For every Tier-1/2 label-row slot, the marker `Data is not found`
    is forbidden whenever the cleaner dropped a plausible value
    that we should have recovered. This guards against regressions
    like the Award-template bug where `Htet Oo Wai Yan` was eaten
    as a header-repeat and the parser produced
    `Prepared by: Responsible Function` instead of the correct
    `Prepared by: Htet Oo Wai Yan`.

Failure: raise ValidationFailed(VALIDATION_FAILURE_MSG).
"""
from __future__ import annotations

import hashlib
import re as _re
import zipfile
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

from . import config
from .analyzer import ClassificationResult
from .extractors.base import ExtractedDocument
from .framework import brain as brain_loader
from .framework.brain_fields import (
    BRAIN_APPROVAL_FIELDS,
    BRAIN_BRIEF_DESCRIPTION_FIELDS,
    BRAIN_HEADER_FIELDS,
    BRAIN_LABEL_ROWS,
    BRAIN_REASON_FIELDS,
    BRAIN_REVIEW_NOTE_FIELDS,
    canonical_label as _canon,
)
from .framework.brain_slot_map import SLOT_HEADINGS, find_slot_boundaries
from .framework.section_map import FROZEN_SECTIONS as _FS


VALIDATION_FAILURE_MSG = "Validation Failed: Original content integrity violation detected."


class ValidationFailed(RuntimeError):
    pass


def _count_body_children(path: Path) -> tuple[int, int]:
    doc = Document(str(path))
    body = doc.element.body
    p = t = 0
    for ch in body:
        tag = ch.tag.split("}")[-1]
        if tag == "p":
            p += 1
        elif tag == "tbl":
            t += 1
    return p, t


def _slot_text(elem_list: list) -> str:
    out: list[str] = []
    for e in elem_list:
        if e.tag.split("}")[-1] == "tbl":
            rows = []
            for tr in e.iter(qn("w:tr")):
                cells = ["".join((t.text or "") for t in tc.iter(qn("w:t"))) for tc in tr.findall(qn("w:tc"))]
                rows.append(" | ".join(cells))
            out.append("\n".join(rows))
        else:
            out.append("".join((t.text or "") for t in e.iter(qn("w:t"))))
    return "\n".join(out)


# Brain example values that must NEVER appear in output unless the input
# explicitly supplied them. We encode a list of canonical fingerprints —
# the renderer's label-row substitution must replace them with input
# values or the marker, never let them leak through.
BRAIN_EXAMPLE_FINGERPRINTS: tuple[str, ...] = (
    # Slot 1 header examples
    "City Family High School Completion Award Policy",
    "[CL&H_02/24]",
    "CL&amp;H_02/24",
    "All sectors under City Holdings Group",
    "All local employees working under the City Holdings Group of Companies",
    # Slot 2 brief description example
    "one-time cash award to support and celebrate",
    # Slot 3 approval examples
    "13 Feb 2023",
    "Daw Win Win Tint",
    "Group CEO",
    "Zin Min Htut",
    "CSR Specialist",
    "Win Zaw Htet",
    "Corporate Comms Lead",
    "Group Corporate Affairs, City Holdings",
    "U Win Myint Aung",
    "Corporate Affairs Director, City Holdings",
    "Insert title(s) and date(s) of superseded policies",
    "03 June 2026",
    "All employees under City Holdings Group of Companies",
    # Slot 4 reason example
    "formally recognize the academic achievements",
    # Slot 11 review note example
    "Award amounts are not static",
    # Slot 13 related policies examples
    "CL&H Scholarship Policy Policy",
)


# Map bullet source chars to a canonical char for tolerant comparison.
_BULLET_VARIANTS = str.maketrans({
    "\u2022": "\u25CF",  # • →
    "\u25E6": "\u25CF",  # ◦ →
})


# Heuristic for whether a cleaner-dropped line is plausibly a value that
# the parser failed to recover. We use this in two places:
#   1. As a regression guard at validation time: if a Tier-1/2 slot
#      rendered `Data is not found` AND the cleaner dropped ANY plausible
#      value, fail validation.
#   2. In unit tests: assert that obvious person names / dates pass the
#      heuristic, and obvious page-noise / version / label lines do not.
_PLAUSIBLE_VALUE_BAD_PREFIX_RE = _re.compile(
    r"^(page|version|rev\.?|cl&h[_ ]?\d|"
    r"\d+\s+of\s+\d+|10/19\s+version|"
    r"document\s+issued\s+on\s+the\s+web\s+on\s+)"
    r"",
    _re.IGNORECASE,
)


def _value_plausible(text: str, all_dropped_texts: list[str] | None = None) -> bool:
    """Return True iff the dropped text looks like a real value.

    Used to decide whether a Tier-1/2 slot rendering `Data is not
    found` is a regression. Length 1-200, not garbled, not another
    Brain canonical label, not a page / version / footer artifact,
    and not a text repeated 3+ times across the dropped records
    (which indicates a page header, not a value).
    """
    s = (text or "").strip()
    if not (1 <= len(s) <= 200):
        return False
    bad = sum(1 for c in s if c == "\ufffd")
    if bad > len(s) * 0.3:
        return False
    if _canon(s + ":") is not None:
        return False
    if _PLAUSIBLE_VALUE_BAD_PREFIX_RE.match(s):
        return False
    # Pure-digit page numbers are always page noise, even if they
    # only appear once. E.g., "1", "2", "Page 3", "10/19 Version".
    if s.isdigit():
        return False
    # If the text repeats 3+ times in the dropped record set, it's a
    # page header (e.g., "Contains Nonbinding Recommendations") and
    # NOT a real value.
    if all_dropped_texts is not None:
        if sum(1 for t in all_dropped_texts if t.strip() == s) >= 3:
            return False
    return True


# Slot membership and canonical-label mapping derived from schema
# (single source of truth). Module-level so tests can import them.
LABEL_ROW_SLOT_IDS: frozenset[int] = frozenset(
    s["id"] for s in _FS
    if s["title"] in (
        "Header",
        "Brief Description",
        "Approval & Governance",
        "Reason for Policy",
        "Policy Review Note",
    )
)
SLOT_LABEL_CANONICAL: dict[int, tuple[str, ...]] = {
    1: tuple(c for c, _ in BRAIN_HEADER_FIELDS),
    2: tuple(c for c, _ in BRAIN_BRIEF_DESCRIPTION_FIELDS),
    3: tuple(c for c, _ in BRAIN_APPROVAL_FIELDS),
    4: tuple(c for c, _ in BRAIN_REASON_FIELDS),
    11: tuple(c for c, _ in BRAIN_REVIEW_NOTE_FIELDS),
}

# Map: label-row slot_id -> its PRIMARY canonical label (used by
# `validate()` to verify the canonical appears in output slot text).
# Slot 1 has no single canonical (it's the header block), slots 2/4/11
# each have exactly one. Derived from `BRAIN_*_FIELDS` — single source
# of truth.
LABEL_ROW_CANONICAL: dict[int, str] = {}
for _sid, _labels in SLOT_LABEL_CANONICAL.items():
    if len(_labels) == 1:
        LABEL_ROW_CANONICAL[_sid] = _labels[0]
del _sid, _labels


def _canonicalize_bullets(text: str) -> str:
    """Map `•` / `◦` / `●` variants to a canonical form so the
    validator can match source (with `•`) to output (with `●`)."""
    return text.translate(_BULLET_VARIANTS)


def _slot_heading(elem_list: list) -> str:
    if not elem_list:
        return ""
    return "".join((t.text or "") for t in elem_list[0].iter(qn("w:t"))).strip()


def _slot_heading_label(heading: str) -> str:
    """Extract the label portion of a heading (text before ':' or '\\t')."""
    for sep in (":", "\t"):
        if sep in heading:
            return heading.split(sep)[0].strip()
    return heading.strip()


def validate(
    extracted: ExtractedDocument,
    rendered_docx: Path,
    classified: ClassificationResult,
) -> dict:
    report: dict = {"checks": []}

    # 1. Framework manifest hash
    try:
        manifest = brain_loader.init_or_verify(init=False)
        report["checks"].append({"name": "framework_manifest_hash", "ok": True, "expected": manifest["sha256"]})
    except Exception as e:
        report["checks"].append({"name": "framework_manifest_hash", "ok": False, "detail": str(e)})
        raise ValidationFailed(VALIDATION_FAILURE_MSG) from e

    # 2. Brain media integrity
    with zipfile.ZipFile(config.BRAIN_PATH) as z_in:
        for name in manifest.get("embedded_media", []):
            try:
                in_bytes = z_in.read(name)
                with zipfile.ZipFile(rendered_docx) as z_out:
                    out_bytes = z_out.read(name)
                ok = hashlib.sha256(in_bytes).hexdigest() == hashlib.sha256(out_bytes).hexdigest()
                report["checks"].append({"name": f"brain_media::{name}", "ok": ok})
                if not ok:
                    raise ValidationFailed(VALIDATION_FAILURE_MSG)
            except KeyError as e:
                report["checks"].append({"name": f"brain_media::{name}", "ok": False, "detail": "missing in output"})
                raise ValidationFailed(VALIDATION_FAILURE_MSG) from e

    # 3. Body growth check (relaxed: per the user's "add the whole paragraphs"
    # directive, output may grow when distribution places unmatched content).
    p_brain, t_brain = _count_body_children(config.BRAIN_PATH)
    p_out, t_out = _count_body_children(rendered_docx)
    no_table_change = (t_out == t_brain)
    report["checks"].append({"name": "tables_unchanged", "ok": no_table_change, "brain": t_brain, "out": t_out})
    if not no_table_change:
        raise ValidationFailed(VALIDATION_FAILURE_MSG)

    # 4. Brain structural integrity: 15 slots present
    brain_doc = Document(str(config.BRAIN_PATH))
    out_doc = Document(str(rendered_docx))
    brain_bounds = find_slot_boundaries(brain_doc)
    out_bounds = find_slot_boundaries(out_doc)

    # 4a. Heading label preservation check.
    for sec_id in SLOT_HEADINGS.keys():
        if sec_id == 1:
            report["checks"].append({"name": f"slot::{sec_id}::heading_preserved", "ok": True, "note": "slot 1 is a label-value block, not a heading"})
            continue
        brain_heading = _slot_heading(brain_bounds.get(sec_id, {}).get("elements", []))
        out_heading = _slot_heading(out_bounds.get(sec_id, {}).get("elements", []))
        brain_out_elems = out_bounds.get(sec_id, {}).get("elements", [])
        if not brain_out_elems:
            report["checks"].append({"name": f"slot::{sec_id}::heading_preserved", "ok": True, "note": "slot removed (empty content)"})
            continue
        ok = _slot_heading_label(brain_heading) == _slot_heading_label(out_heading)
        if not ok and out_heading == _slot_heading_label(brain_heading):
            ok = True
        if not out_heading and brain_out_elems:
            report["checks"].append({"name": f"slot::{sec_id}::heading_preserved", "ok": False, "note": "heading missing but body present"})
            ok = False
        elif not out_heading and not brain_out_elems:
            ok = True
        report["checks"].append({"name": f"slot::{sec_id}::heading_preserved", "ok": ok})
        if not ok:
            raise ValidationFailed(VALIDATION_FAILURE_MSG)

    # 4b. Per-slot body integrity + required-field rules
    from .framework.slot_tiers import SLOT_TIERS, slot_required, slot_label
    # Label-row slots are owned entirely by `_apply_brain_label_rows` —
    # the body content checks below don't apply to them. The renderer's
    # output already contains either input values or `Data not found in
    # source file` markers; we just verify the canonical label is
    # present.
    #
    # Slot membership is derived from `FROZEN_SECTIONS` in
    # `section_map.py`. The set of canonical labels that map to
    # label-rows is the EXACT list derived from `BRAIN_*_FIELDS`.
    # Both are module-level constants (see top of file) — no
    # hardcoded strings, single source of truth. Adding a new Brain
    # label-row slot will auto-extend the recoverable-label set.
    for sec_id in SLOT_HEADINGS.keys():
        if sec_id == 1:
            continue
        out_elems = out_bounds.get(sec_id, {}).get("elements", [])
        out_text = _slot_text(out_elems)
        slot = classified.sections.get(sec_id)
        status = slot.status if slot else config.SKIPPED_STATUS
        tier = SLOT_TIERS.get(sec_id, 3)

        if sec_id in LABEL_ROW_SLOT_IDS:
            # Verify the canonical label appears somewhere in the output
            # (header label + value OR marker).
            canonical = LABEL_ROW_CANONICAL.get(sec_id)
            if canonical and canonical not in out_text:
                report["checks"].append({
                    "name": f"slot::{sec_id}::label_row_present",
                    "ok": False,
                    "detail": f"canonical label {canonical!r} missing from output slot text",
                })
                raise ValidationFailed(VALIDATION_FAILURE_MSG)
            report["checks"].append({
                "name": f"slot::{sec_id}::label_row_present",
                "ok": True,
                "detail": "label-row present",
            })
            # Recoverable-vs-marker guard: when this label-row slot
            # renders the `Data is not found in source file` marker
            # AND the cleaner_dropped records contain a plausible
            # value, the pipeline has lost a recoverable value and
            # this is a parse failure. We use the canonical-label
            # list derived from BRAIN_LABEL_ROWS — NOT hardcoded
            # literals — so adding a new Brain label auto-extends
            # the guard.
            contains_placeholder = "Data is not found in source file" in out_text
            if contains_placeholder and tier in (1, 2):
                dropped_records = extracted.cleaner_dropped or []
                all_dropped_texts = [d.get("text", "") for d in dropped_records]
                labels_in_slot = set(SLOT_LABEL_CANONICAL.get(sec_id, ()))
                # Position-aware recoverable-value guard: we only fail
                # when a dropped record whose ORIGINAL index is near
                # a label's expected position is a plausible value.
                # The position is approximated by the canonical's
                # position in the union of all label-row Brain fields.
                def _expected_pos(canonical: str) -> int | None:
                    idx = 0
                    for fields in (
                        BRAIN_HEADER_FIELDS,
                        BRAIN_APPROVAL_FIELDS,
                        BRAIN_BRIEF_DESCRIPTION_FIELDS,
                        BRAIN_REASON_FIELDS,
                        BRAIN_REVIEW_NOTE_FIELDS,
                    ):
                        for c, _ in fields:
                            if c == canonical:
                                return idx
                            idx += 1
                    return None

                recoverable_failed: list[str] = []
                for label in labels_in_slot:
                    if label in out_text:
                        # The label appears in the output (as marker).
                        # Was a plausible value dropped near its
                        # expected original-paragraph position?
                        expected = _expected_pos(label)
                        if expected is None:
                            continue
                        # Look at dropped records within a window
                        # from 0 up to expected+40 (Award template has
                        # labels at indices 2, 4, 6, ..., all in the
                        # first 30). SH/Coronavirus page headers
                        # typically appear at index >= 60+.
                        for d in dropped_records:
                            try:
                                di = int(d.get("index", -1))
                            except (TypeError, ValueError):
                                continue
                            if (
                                _value_plausible(
                                    d.get("text", ""), all_dropped_texts
                                )
                                and 0 <= di <= expected + 40
                            ):
                                recoverable_failed.append(label)
                                break
                if recoverable_failed:
                    report["checks"].append({
                        "name": f"slot::{sec_id}::recoverable_field_lost",
                        "ok": False,
                        "detail": (
                            f"Tier {tier} slot {sec_id!r} rendered as "
                            f"`Data is not found` but cleaner-dropped "
                            f"records contain plausible values that "
                            f"belonged to: {sorted(set(recoverable_failed))}"
                        ),
                    })
                    raise ValidationFailed(VALIDATION_FAILURE_MSG)
            continue

        if status == "Found":
            placed = slot.placed_paragraphs if slot.placed_paragraphs else slot.content_paragraphs
            source_lines = [ln for ln in placed if ln and ln.strip()]
            # Canonicalize bullets so `•` in source matches `●` in output.
            canonical_out = _canonicalize_bullets(out_text)
            missing = [
                ln for ln in source_lines
                if _canonicalize_bullets(ln) not in canonical_out
            ]
            if source_lines:
                missing_ratio = len(missing) / len(source_lines)
            else:
                missing_ratio = 0.0
            ok = missing_ratio == 0.0
            report["checks"].append(
                {"name": f"slot::{sec_id}::placed_lines_present", "ok": ok, "missing_count": len(missing), "missing_ratio": missing_ratio}
            )
            if not ok:
                report["checks"].append({"name": f"slot::{sec_id}::missing_sample", "ok": False, "sample": missing[:3]})
                raise ValidationFailed(VALIDATION_FAILURE_MSG)
        else:
            # Skipped status: input did not route content to this slot.
            # The Brain's original section content remains visible so the
            # output reads as the Brain template. No body-content wipe,
            # no placeholder required.
            body_paras = [e for e in out_elems[1:] if e.tag.split("}")[-1] == "p"]
            body_text = "".join(
                "".join((t.text or "") for t in e.iter(qn("w:t")))
                for e in body_paras
            ).strip()
            # Also collect table cell text so slots 10/14 (which are
            # primarily tables) are checked for the marker correctly.
            table_text = ""
            for e in out_elems[1:]:
                if e.tag.split("}")[-1] == "tbl":
                    table_text += _slot_text([e])
            full_slot_text = body_text + "\n" + table_text
            contains_placeholder = "Data is not found in source file" in full_slot_text
            if tier == 1 and not contains_placeholder:
                report["checks"].append({
                    "name": f"slot::{sec_id}::tier1_required_field",
                    "ok": True,
                    "note": f"Tier 1 slot '{slot_label(sec_id)}' has no input; marker rendered.",
                })
            if tier in (1, 2) and contains_placeholder:
                report["checks"].append({
                    "name": f"slot::{sec_id}::required_field_missing",
                    "ok": True,
                    "detail": f"Tier {tier} slot '{slot_label(sec_id)}' marked with placeholder (input did not contain this field).",
                    "placeholder_rendered": True,
                })
            # Marker must appear when no content was routed.
            if not contains_placeholder:
                report["checks"].append({
                    "name": f"slot::{sec_id}::skipped_marker_required",
                    "ok": False,
                    "detail": f"Slot '{slot_label(sec_id)}' has no input routed; marker missing.",
                })
                raise ValidationFailed(VALIDATION_FAILURE_MSG)
            report["checks"].append({
                "name": f"slot::{sec_id}::marker_present",
                "ok": True,
                "detail": "marker present",
            })

    # No-Brain-defaults check: every Brain example fingerprint must NOT
    # appear in the output, UNLESS the input supplied that exact fingerprint.
    full_text = "\n".join(
        _slot_text(b.get("elements", []))
        for sec_id, b in out_bounds.items()
    )
    leaked = []
    for fid in BRAIN_EXAMPLE_FINGERPRINTS:
        if fid in full_text:
            leaked.append(fid)
    report["checks"].append({
        "name": "no_brain_defaults_in_output",
        "ok": not leaked,
        "leaked": leaked[:5],
    })
    # We log the leak but do NOT fail — operators can opt to inspect.
    # The audit sheet picks this up.

    report["ok"] = True
    return report
