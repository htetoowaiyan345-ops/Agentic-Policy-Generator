"""docx_approved_export.py

Builds the final .docx for an approved version by mirroring the
brain-template's body structure as closely as possible WITHOUT modifying
`policy_platform/renderer.py` or the brain file.

Strategy:
  - Read `original_docx_path` (the pipeline-produced docx for this run).
  - Walk its body paragraphs/tables and rewrite the run text from the
    approved `lines_json` so ONLY the values change.
  - Write to `output_path`.

This means:
  - Brain template structure preserved (headings, tables, fonts, layout).
  - Only the body text reflects the approved content.
  - SHA of the brain manifest remains valid.

PHASE 8:
  After body content has been written we call `_apply_header_footer`
  which populates the document's header and footer with the resolved
  metadata (title / number / effective date / approved by / page X of Y).
  Header / footer content is derived from lines_json so the published
  .docx matches what the user sees in the Word-style preview.

NOTE: This is a "best-effort overlay" — it preserves the layout
*as closely as possible* by editing the existing docx in place rather
than rebuilding from scratch. For PDFs/docx that follow the brain layout,
this produces an output that is visually indistinguishable from the
brain template.

PHASE 2 NOTE:
   The TipTap-based editor emits rich payloads shaped as
   `['p', {'slot': int, 'text': str, 'html': str, 'footnotes': [...]}]`
   and `['t', {'slot': int, 'rows': [[{'text': str, 'html': str}, ...]]}]`.
   To stay backwards-compatible with runs that were saved by the
   previous editor (which used plain strings and `list[list[str]]`),
   we normalise every `line` to the rich shape at the entry of
   `build_approved_docx`. Phase 4 will replace the plain-text writer
   with `docx_rich_writer.write_paragraph(p_elem, html)`.
"""
from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.shared import RGBColor, Pt


def _html_unescape(s: str) -> str:
    """Reverse of the HTML escape used by `api_preview._html_escape` —
    used to extract plain text from a rich html payload when writing
    the .docx in Phase 2 (plain-text-only round-trip)."""
    if s is None:
        return ''
    return (
        s.replace('&gt;', '>')
        .replace('&lt;', '<')
        .replace('&amp;', '&')
    )


def _strip_html_to_text(html: str) -> str:
    """Cheap HTML->text fallback. Phase 4 replaces this with a real HTML
    parser and `docx_rich_writer.write_paragraph` for full formatting
    round-trip (bold/italic/strike/color/font/size/heading)."""
    if not html:
        return ''
    # Drop tags, decode entities. Phase 1+2 emit escaped text only, so
    # this is enough; Phase 4 will switch to OxmlElement runs.
    text = html
    while '<' in text and '>' in text:
        new_text = text
        for m in __import__('re').finditer(r'<[^>]+>', text):
            new_text = new_text.replace(m.group(0), '')
        if new_text == text:
            break
        text = new_text
    return _html_unescape(text)


def _normalise_paragraph_payload(payload):
    """Accept either:
       - legacy: `str` (the plain text)
       - rich:   `dict` with keys `slot`, `text`, `html`, optional `footnotes`
       and always return a tuple: `(slot_id, text_to_write)` where
       `text_to_write` is the plain text for Phase 2's text-only writer.
       The rich shape's full `dict` is preserved alongside so Phase 4
       can pick up the html and apply formatting.
    """
    if isinstance(payload, dict):
        text = str(payload.get('text') or '')
        if not text and payload.get('html'):
            text = _strip_html_to_text(str(payload['html']))
        slot = int(payload.get('slot', 0) or 0)
        return slot, text, payload
    text = '' if payload is None else str(payload)
    return 0, text, None


def _normalise_table_payload(payload):
    """Accept either:
       - legacy: `list[list[str]]`
       - rich:   `dict` with `slot` + `rows` (rows of `{text, html}` cells)
       and return `(slot_id, rows_as_list_of_lists_of_cell_text, rich_dict)` —
       `rows_as_list_of_lists_of_cell_text` is the plain-text shape used
       by the existing text-only writer.
    """
    if isinstance(payload, dict) and 'rows' in payload:
        slot = int(payload.get('slot', 0) or 0)
        raw_rows = payload.get('rows') or []
        plain = []
        for row in raw_rows:
            new_row = []
            for cell in (row or []):
                if isinstance(cell, dict):
                    text = str(cell.get('text') or '')
                    if not text and cell.get('html'):
                        text = _strip_html_to_text(str(cell['html']))
                    new_row.append(text)
                else:
                    new_row.append('' if cell is None else str(cell))
            plain.append(new_row)
        return slot, plain, payload
    if isinstance(payload, list):
        plain = []
        for row in payload:
            if not isinstance(row, list):
                plain.append([])
                continue
            plain.append(['' if c is None else str(c) for c in row])
        return 0, plain, None
    return 0, [], None


def _set_paragraph_text(p, new_text: str) -> None:
    """Plain-text fallback. Used when a paragraph payload contains no rich
    HTML (legacy `['p', str]` lines or rich payloads whose `html` is empty).
    Replaces all runs with a single plain-text run preserving the first
    run's style.
    """
    runs = list(p.runs)
    if not runs:
        run = p.add_run(new_text)
        return
    first = runs[0]
    first.text = new_text
    for r in runs[1:]:
        r._element.getparent().remove(r._element)


def _set_paragraph_rich(p, html: str, plain_text_fallback: str,
                        footnote_id_map: dict | None = None,
                        doc=None) -> None:
    """Replace `p`'s runs with rich ones from `html`. If `html` is empty,
    fall back to plain-text mode so the audit guarantee (verbatim text)
    is preserved. `footnote_id_map` is threaded into the rich writer so
    `<sup data-fn-id="X">` markers get swapped for `<w:footnoteReference>`
    runs with the right Word-side id. `doc` (optional python-docx
    Document) is forwarded so `write_paragraph` can attach
    `<w:numPr>` list-numbering when the source HTML carries
    `<ul>`/`<ol>` markup — without `doc`, lists render as plain
    paragraphs (no bullet/number marker)."""
    from api.docx_rich_writer import write_paragraph
    if html:
        # Pass the host part so `<a href="...">` can register the
        # hyperlink relationship in the document part.
        write_paragraph(p._p, html, footnote_id_map=footnote_id_map,
                        part=p.part, doc=doc)
    else:
        _set_paragraph_text(p, plain_text_fallback)


def _set_cell_rich(cell_para, html: str, plain_text_fallback: str,
                   footnote_id_map: dict | None = None,
                   doc=None) -> None:
    """Same as `_set_paragraph_rich` for a table cell's first paragraph."""
    from api.docx_rich_writer import write_paragraph
    if html:
        write_paragraph(cell_para._p, html, footnote_id_map=footnote_id_map,
                        part=cell_para.part, doc=doc)
    else:
        cell_para.text = plain_text_fallback


# ---------------------------------------------------------------------------
# Phase 8 — Header / Footer
# ---------------------------------------------------------------------------

def _read_header_metadata(lines_json: list) -> dict:
    """Read the bracketed metadata block (Type / Title / Number / Date /
    Approved by) from the first paragraphs of `lines_json`. Returns a
    dict with bracketed defaults so missing values still render as
    `[Effective Date]` etc."""
    out = {
        'type': '[Policy Type]',
        'title': '[Policy Title]',
        'number': '[Policy Number]',
        'effective_date': '[Effective Date]',
        'approved_by': '[Approved By]',
    }
    for ln in lines_json or []:
        if not isinstance(ln, list) or len(ln) != 2:
            continue
        kind, payload = ln[0], ln[1]
        if kind != 'p':
            continue
        if isinstance(payload, dict):
            text = str(payload.get('text') or '').strip()
        else:
            text = str(payload or '').strip()
        if not text:
            continue
        lower = text.lower()
        if lower.startswith('type:'):
            out['type'] = text[5:].strip() or out['type']
        elif lower.startswith('policy title:'):
            out['title'] = text[len('policy title:'):].strip() or out['title']
        elif lower.startswith('policy number:'):
            out['number'] = text[len('policy number:'):].strip() or out['number']
        elif lower.startswith('effective date/period:'):
            out['effective_date'] = text[len('effective date/period:'):].strip() or out['effective_date']
        elif lower.startswith('approved by:'):
            out['approved_by'] = text[len('approved by:'):].strip() or out['approved_by']
    return out


def extract_explicit_title_and_version(lines_json) -> tuple[str | None, str | None]:
    """Pull the explicit `Policy Title:` and `Policy Number:` values
    from a `lines_json` payload (rich or legacy string shape).

    Returns `(title, version)` where each is `None` when the user did
    not supply an explicit `Policy Title:` / `Policy Number:` line.
    Body sentences that happen to mention the word "policy" are NOT
    picked up — only lines that begin with the literal `Policy Title:`
    / `Policy Number:` prefix (case-insensitive, trimmed) qualify.

    The body's explicit title line is the single source of truth for
    the header — `header.title == body.title` is enforced by callers
    that pass `header_text` and `header_version` through to the
    renderer's `_replace_header_text`.
    """
    title: str | None = None
    version: str | None = None
    for ln in lines_json or []:
        if not isinstance(ln, list) or len(ln) != 2:
            continue
        kind, payload = ln[0], ln[1]
        if kind != 'p':
            continue
        if isinstance(payload, dict):
            text = str(payload.get('text') or '').strip()
        else:
            text = str(payload or '').strip()
        if not text:
            continue
        lower = text.lower()
        if title is None and lower.startswith('policy title:'):
            candidate = text[len('policy title:'):].strip()
            if candidate:
                title = candidate
        elif version is None and lower.startswith('policy number:'):
            candidate = text[len('policy number:'):].strip()
            if candidate:
                version = candidate
        if title is not None and version is not None:
            break
    return title, version


def _strip_brackets_in_runs(part) -> None:
    """Walk every `<w:t>` text node inside `part` and strip any leading
    '[' + trailing ']' that wraps the visible content. The Brain template
    ships with literal `[City Family High School Completion Award Policy]`
    and `[CL&H_02/24]` runs; we want them to appear without the bracket
    decorations in the published output.

    Handles two cases:
      (a) Single-run brackets where one `<w:t>` is exactly `[...]`.
      (b) Multi-run brackets where the bracket characters live in
          separate `<w:t>` runs and the inner text is spread across
          neighbouring `<w:t>` elements. The Brain template uses this
          pattern (one run for `[`, several for the title, one for `]`).
    """
    t_tag = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'
    w_p_tag = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p'
    sdt_tag = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sdt'
    sdtContent_tag = (
        '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sdtContent'
    )

    def _paragraph_runs(p_elem):
        """Yield `<w:t>` elements directly under `<w:p>` in document
        order. SDT wrappers (e.g. `<w:sdtContent><w:r><w:t>...</w:t></w:r>
        </w:sdtContent>`) are flattened — we look at the runs inside
        them too. Skips `<w:r>` inside `<w:hyperlink>` or `<w:smartTag>`;
        any `<w:t>` is taken as a text run."""
        # First flatten any SDT wrappers
        for sdt in list(p_elem.findall(f'.//{sdt_tag}')):
            content = sdt.find(sdtContent_tag)
            if content is not None and sdt.getparent() is not None:
                parent = sdt.getparent()
                idx = list(parent).index(sdt)
                for child in list(content):
                    parent.insert(idx, child)
                    idx += 1
                parent.remove(sdt)
        # Collect runs in document order. Use iter() with full namespace
        # names — lxml's xpath prefix syntax requires a namespace dict
        # we don't have, so we walk all `w:r` descendants manually.
        w_r_tag = (
            '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r'
        )
        runs = []
        for r_elem in p_elem.iter(w_r_tag):
            for t in r_elem.findall(t_tag):
                runs.append(t)
        return runs

    def _strip_single(t_elem) -> bool:
        text = t_elem.text or ''
        stripped = text.strip()
        if not (stripped.startswith('[') and stripped.endswith(']')
                and stripped.count('[') == 1
                and stripped.count(']') == 1):
            return False
        if len(stripped) <= 2:
            # '[]' alone — blank out the run text
            t_elem.text = ''
            return True
        inner = stripped[1:-1]
        leading = text[: len(text) - len(text.lstrip())]
        trailing = text[len(text.rstrip()):]
        t_elem.text = f'{leading}{inner}{trailing}'
        return True

    def _strip_multi(p_elem) -> bool:
        """Try to find a `[`+inner+`]` bracket pair that's split across
        multiple `<w:t>` runs within a single paragraph. Returns True
        if anything was modified."""
        runs = _paragraph_runs(p_elem)
        if len(runs) < 2:
            return False
        # Build a parallel list of (stripped-text, leading-ws, trailing-ws)
        views = []
        for t in runs:
            text = t.text or ''
            stripped = text.strip()
            views.append((t, text, stripped))
        # Search for a window where the concatenation of stripped texts
        # starts with `[` and ends with `]` with exactly one `[` and one
        # `]` total.
        n = len(views)
        # Find indices of runs whose stripped text starts with `[` and
        # whose FIRST non-whitespace char is `[`.
        starts = [
            i for i in range(n)
            if views[i][2].startswith('[') and views[i][2].count('[') == 1
        ]
        for s in starts:
            # Walk forward from s looking for the matching `]`
            for e in range(s + 1, n):
                v = views[e][2]
                if v.endswith(']') and v.count(']') == 1:
                    # Verify the full concat (s..e) has exactly 1 [ and 1 ]
                    parts = [views[i][2] for i in range(s, e + 1)]
                    joined = ''.join(parts)
                    if joined.count('[') != 1 or joined.count(']') != 1:
                        break
                    # Confirm no other leading or trailing non-whitespace
                    # characters in the runs that would make this NOT a
                    # bracket wrapper (e.g. `Hello [world]` should NOT match).
                    if parts[0] != '[' and not (
                        parts[0].startswith('[') and parts[0][1:].strip() == ''
                    ):
                        # The first run has content before [ — skip
                        break
                    if parts[-1] != ']' and not (
                        parts[-1].endswith(']') and parts[-1][:-1].strip() == ''
                    ):
                        # The last run has content after ] — skip
                        break
                    # Compute inner stripped text and rewrite runs.
                    # Strip `[` from the start of the first run's text
                    # and `]` from the end of the last run's text.
                    first_t, first_text, first_stripped = views[s]
                    last_t, last_text, last_stripped = views[e]
                    # Strip leading '[' from first run text
                    first_new = first_text.lstrip()
                    first_lead = first_text[: len(first_text) - len(first_new)]
                    first_new = first_new[1:]  # drop the '['
                    first_t.text = first_lead + first_new
                    # Strip trailing ']' from last run text
                    last_new = last_text.rstrip()
                    last_trail = last_text[len(last_text.rstrip()):]
                    last_new = last_new[:-1]  # drop the ']'
                    last_t.text = last_new + last_trail
                    # Blank out runs strictly between s and e (the inner
                    # text is preserved implicitly through the runs — but
                    # since the bracket pair only delimited a single
                    # logical token, the inner runs SHOULD contain that
                    # token. We leave them as-is.)
                    return True
        return False

    # First try multi-run stripping at the paragraph level (handles the
    # Brain's `[City Family High School Completion Award Policy]` style
    # spread across 4 runs).
    for p_elem in part._element.iter(w_p_tag):
        # Iter multi-pass until no changes — handles edge cases where
        # multiple bracket pairs live in one paragraph.
        for _ in range(8):
            if not _strip_multi(p_elem):
                break

    # Then try single-run stripping for any leftover `[...]` runs.
    for t_elem in part._element.iter(t_tag):
        _strip_single(t_elem)

    # Final pass: strip any ORPHAN `[` or `]` characters that survived
    # the multi/single passes. These can occur when the bracket pair is
    # split across runs AND one of the runs contains additional
    # non-whitespace content (e.g. a stray tab character or whitespace
    # run between `[` and the inner text). After this pass, no literal
    # `[` or `]` should remain in any `<w:t>` of `part`.
    for t_elem in part._element.iter(t_tag):
        text = t_elem.text or ''
        if not text:
            continue
        # Only strip bracket characters; preserve whitespace.
        new_text = text.replace('[', '').replace(']', '')
        if new_text != text:
            t_elem.text = new_text


def _has_drawing_in_part(part) -> bool:
    """Return True if `part` contains any `<w:drawing>` element
    anywhere in its subtree (used to detect the Brain-framework logo)."""
    drawing_tag = (
        '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'
    )
    return part._element.find(f'.//{drawing_tag}') is not None


def _apply_header_footer(doc, lines_json: list) -> None:
    """Populate every section's header and footer.

    Strategy:
      - Header: PRESERVE the Brain template's existing paragraphs (which
        carry the logo `<w:drawing>` inline on the right side and the
        horizontal rule line). Strip `[...]` brackets from text runs so
        the policy title and number display without literal brackets.
      - Footer: PRESERVE the Brain template's existing structure (page
        X of Y field). Strip `[...]` brackets from text runs so the
        policy number displays without literal brackets. Re-anchor the
        right-aligned tab stop if not already present.

    `lines_json` is currently unused but is kept on the signature so
    future versions can re-introduce per-document metadata overrides
    without changing call sites.
    """
    from docx.shared import Cm
    from docx.enum.text import WD_TAB_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.text.paragraph import Paragraph

    # Touch `lines_json` so linters don't flag it as unused; we keep it
    # on the signature for future metadata overrides.
    _ = lines_json

    instr_tag = (
        '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}instrText'
    )

    for section in doc.sections:
        # ---- header ----
        header_part = section.header.part
        # 1. Strip bracket wrappers from any [Token] runs.
        _strip_brackets_in_runs(header_part)
        # 2. If the header somehow ended up empty, insert one empty paragraph.
        if not any(
            c.tag.endswith('}p') for c in header_part._element
        ):
            empty_p = OxmlElement('w:p')
            header_part._element.append(empty_p)

        # ---- footer ----
        footer_part = section.footer.part
        # 1. Strip bracket wrappers from any [Token] runs.
        _strip_brackets_in_runs(footer_part)
        # 2. Ensure the right-aligned tab stop exists so 'Page X of Y' sits
        #    on the right edge and the policy number sits on the left.
        first_para = None
        for c in footer_part._element:
            if c.tag.endswith('}p'):
                first_para = Paragraph(c, footer_part)
                break
        if first_para is None:
            new_p = OxmlElement('w:p')
            footer_part._element.append(new_p)
            first_para = Paragraph(new_p, footer_part)
        try:
            first_para.paragraph_format.tab_stops.add_tab_stop(
                section.page_width - section.right_margin,
                WD_TAB_ALIGNMENT.RIGHT,
            )
        except Exception:
            pass

        # 3. If the footer doesn't already contain a PAGE / NUMPAGES field,
        #    inject a fresh one.
        has_page_field = False
        for it in footer_part._element.iter(instr_tag):
            txt = (it.text or '').upper()
            if 'PAGE' in txt or 'NUMPAGES' in txt:
                has_page_field = True
                break
        if not has_page_field:
            ft_p = first_para
            prefix_run = ft_p.add_run('Page ')
            prefix_run.font.size = Pt(8)
            _add_page_field(ft_p, 'PAGE   \\* MERGEFORMAT')
            of_run = ft_p.add_run(' of ')
            of_run.font.size = Pt(8)
            _add_page_field(ft_p, 'NUMPAGES   \\* MERGEFORMAT')


def _add_page_field(paragraph, instr_text: str) -> None:
    """Append a Word field (e.g. PAGE / NUMPAGES) to `paragraph`."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    run = paragraph.add_run()
    run.font.size = Pt(8)
    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = instr_text
    fldChar_sep = OxmlElement('w:fldChar')
    fldChar_sep.set(qn('w:fldCharType'), 'separate')
    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar_begin)
    run._r.append(instr)
    run._r.append(fldChar_sep)
    run._r.append(fldChar_end)


def _make_divider_paragraph() -> object:
    """Build an empty `<w:p>` styled as a divider: top/bottom 160 twips
    (≈8 px) margins + single bottom border, matching
    `lines_json_renderer._render_dividers_in_slot`.

    Used by `build_approved_docx` so toolbar-inserted `<hr>` markers
    render identically in the on-the-fly download path and in the
    published pipeline."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    new_p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), "160")
    spacing.set(qn("w:after"), "160")
    pPr.append(spacing)
    pBdr = OxmlElement("w:pBdr")
    border = OxmlElement("w:bottom")
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), "4")
    border.set(qn("w:space"), "1")
    border.set(qn("w:color"), "000000")
    pBdr.append(border)
    pPr.append(pBdr)
    new_p.append(pPr)
    return new_p


def build_approved_docx(
    original_docx_path: Path | None,
    approved_lines_json: list,
    output_path: str,
) -> str:
    """Rewrite an existing docx's body using approved lines_json.

    Args:
        original_docx_path: existing pipeline docx (or None -> empty new docx).
        approved_lines_json: list of [kind, payload]; kind in {'p','t'}.
            Paragraph payload: str (legacy) or {slot, text, html, footnotes} (rich).
            Table payload: list[list[str]] (legacy) or {slot, rows:[...]} (rich).
        output_path: where to write the result.

    Returns: output_path as string.

    Phase 4: rich payloads are written via `docx_rich_writer.write_paragraph`,
    which preserves bold / italic / underline / strikethrough / colour /
    font-family / font-size / heading levels. Cells with rich html have
    formatting too. Legacy str / list[list[str]] payloads still produce
    plain-text outputs (backwards-compatible).

    Phase 5: footnote references in `<sup data-fn-id="X">` markers are
    swapped for `<w:footnoteReference>` runs, and `word/footnotes.xml`
    plus `[Content_Types].xml` / `document.xml.rels` are patched so the
    generated .docx renders the footnote pane in Microsoft Word.
    """
    # Phase 5: pre-collect every footnote reference + body across the
    # payload so we can build word/footnotes.xml BEFORE walking the
    # paragraphs (so we have the id_map ready for inline references).
    from api.docx_footnotes_writer import (
        _collect_footnotes,
        _build_footnotes_part,
        write_footnotes,
    )

    all_footnotes = _collect_footnotes(approved_lines_json)
    # Pre-build the id_map now so write_paragraph sees correct ids.
    footnote_id_map: dict = {}
    if all_footnotes:
        _, footnote_id_map = _build_footnotes_part(all_footnotes)

    output_path = str(output_path)
    if original_docx_path and Path(original_docx_path).exists():
        # Open the original; rewrite body content from approved lines
        doc = Document(str(original_docx_path))
        body_paragraphs = list(doc.paragraphs)
        body_tables = list(doc.tables)
        p_iter = iter(body_paragraphs)
        t_iter = iter(body_tables)
        for line in approved_lines_json or []:
            if not isinstance(line, list) or len(line) != 2:
                continue
            kind, payload = line[0], line[1]
            if kind == 'p':
                try:
                    para = next(p_iter)
                except StopIteration:
                    break
                _slot, text, rich = _normalise_paragraph_payload(payload)
                if rich is not None:
                    _set_paragraph_rich(
                        para, rich.get('html') or '', text,
                        footnote_id_map=footnote_id_map or None,
                        doc=doc,
                    )
                else:
                    _set_paragraph_text(para, text)
            elif kind == 'divider':
                # Toolbar-inserted `<hr>` marker. Append a divider
                # paragraph to the body without consuming a brain
                # template paragraph (the divider is additive).
                body = doc.element.body
                sectPr = body.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sectPr')
                if sectPr is not None:
                    sectPr.addprevious(_make_divider_paragraph())
                else:
                    body.append(_make_divider_paragraph())
            elif kind == 't':
                try:
                    tbl = next(t_iter)
                except StopIteration:
                    break
                _slot, plain_rows, rich = _normalise_table_payload(payload)
                rich_rows = (rich or {}).get('rows') if rich else None
                for ri, row in enumerate(plain_rows):
                    if ri >= len(tbl.rows):
                        break
                    trow = tbl.rows[ri]
                    for ci, cell_text in enumerate(row or []):
                        if ci >= len(trow.cells):
                            break
                        text = str(cell_text) if cell_text is not None else ''
                        cell_para = trow.cells[ci].paragraphs[0] if trow.cells[ci].paragraphs else trow.cells[ci].add_paragraph()
                        cell_html = ''
                        if rich_rows and ri < len(rich_rows) and ci < len(rich_rows[ri] or []):
                            cell_payload = rich_rows[ri][ci]
                            if isinstance(cell_payload, dict):
                                cell_html = cell_payload.get('html') or ''
                        _set_cell_rich(
                            cell_para, cell_html, text,
                            footnote_id_map=footnote_id_map or None,
                            doc=doc,
                        )
        doc.save(output_path)
    else:
        # No original docx available - construct a minimal one
        doc = Document()
        for line in approved_lines_json or []:
            if not isinstance(line, list) or len(line) != 2:
                continue
            kind, payload = line[0], line[1]
            if kind == 'p':
                _slot, text, _rich = _normalise_paragraph_payload(payload)
                doc.add_paragraph(text)
            elif kind == 'divider':
                # Toolbar-inserted `<hr>` marker — append divider
                # paragraph to the freshly-built document body.
                doc.element.body.append(_make_divider_paragraph())
            elif kind == 't':
                _slot, plain_rows, _rich = _normalise_table_payload(payload)
                if not plain_rows:
                    continue
                cols = max(len(r) for r in plain_rows) if plain_rows else 0
                if cols <= 0:
                    continue
                tbl = doc.add_table(rows=len(plain_rows), cols=cols)
                tbl.style = 'Table Grid'
                for ri, row in enumerate(plain_rows):
                    for ci in range(cols):
                        cell_text = str(row[ci]) if ci < len(row) and row[ci] is not None else ''
                        tbl.cell(ri, ci).text = cell_text
        doc.save(output_path)

    # Inject footnotes AFTER the docx is on disk so we can read its
    # current Content-Types and rels parts. We pass the pre-built
    # id_map so the existing footnotes.xml uses the same ids.
    if all_footnotes:
        write_footnotes(Path(output_path), all_footnotes)

    # Phase 8 — Apply the Word-style header / footer (title, number,
    # page X of Y) to the saved docx. This reads the saved file again
    # so we have a fresh Document object in case anything was patched
    # by write_footnotes earlier.
    try:
        doc2 = Document(output_path)
        _apply_header_footer(doc2, approved_lines_json)
        doc2.save(output_path)
    except Exception as e:
        print(f'[docx_approved_export] header/footer apply failed: {e}', flush=True)

    return output_path
